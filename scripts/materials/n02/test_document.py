import datetime as dt
import os
import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

from materials.shared.ledger_sheet import select_ledger_sheet
from materials.shared.office_word import save_word_document
from materials.shared.word_sections import (
    clone_paragraph_with_text,
    clone_table_with_data,
    find_heading,
    replace_between,
)


TASK_NAME_BOX = (42, 52, 805, 75)
STRATEGY_BOX = (1028, 52, 1110, 75)
SUCCESS_BOX = (1132, 52, 1185, 75)
LAUNCH_BOX = (1420, 52, 1550, 75)
TASK_NAME_POS = (45, 55)
STRATEGY_POS = (1034, 55)
SUCCESS_POS = (1135, 55)
LAUNCH_POS = (1425, 55)
MIN_FONT_SIZE = 12
WHITE = (255, 255, 255)
BLUE = (25, 128, 255)
TEXT = (51, 51, 51)
RED = (255, 77, 79)


@dataclass(frozen=True)
class N02Task:
    task_name: str
    strategy: str
    strategy_text: str
    success_failure: str
    launch_time: object
    launch_time_text: str
    test_time_text: str
    source_department: str
    table_name: str


def _safe_name(value):
    text = re.sub(r"[\\/:*?\"<>|]+", "_", str(value or "").strip())
    text = re.sub(r"\s+", "_", text)
    return text or "task"


def _load_font(size, chinese=False):
    windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    fonts = windir / "Fonts"
    candidates = [
        "msyh.ttc",
        "msyhbd.ttc",
        "simhei.ttf",
        "simsun.ttc",
        "arial.ttf",
    ]
    if not chinese:
        candidates = [
            "arial.ttf",
            "calibri.ttf",
            "segoeui.ttf",
            "msyh.ttc",
            "simhei.ttf",
            "simsun.ttc",
        ]
    for name in candidates:
        font_path = fonts / name
        if font_path.exists():
            try:
                return ImageFont.truetype(str(font_path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def _erase_text_pixels(image, box):
    pixels = image.load()
    x1, y1, x2, y2 = box
    for y in range(y1, y2):
        for x in range(x1, x2):
            r, g, b = pixels[x, y]
            blueish = b > 135 and b > r + 12 and r < 235
            reddish = r > 155 and g < 155 and b < 155 and r > g + 25
            darkish = r < 165 and g < 165 and b < 165
            gray_text = max(r, g, b) < 210 and max(r, g, b) - min(r, g, b) < 35
            if blueish or reddish or darkish or gray_text:
                pixels[x, y] = WHITE


def _contains_cjk(text):
    return any("\u4e00" <= char <= "\u9fff" for char in str(text or ""))


def _draw_success_failure(draw, text, font):
    value = str(text or "").strip()
    if not value:
        return
    x, y = SUCCESS_POS
    if "/" not in value:
        draw.text((x, y), value, font=font, fill=TEXT)
        return
    success, fail = value.split("/", 1)
    draw.text((x, y), success, font=font, fill=BLUE)
    x += int(round(draw.textlength(success, font=font)))
    draw.text((x, y), "/", font=font, fill=TEXT)
    x += int(round(draw.textlength("/", font=font)))
    draw.text((x, y), fail, font=font, fill=RED)


def normalize_execution_strategy(value):
    text = str(value or "").strip()
    if not text:
        return ""
    daily = re.fullmatch(r"(\d{1,2}):(\d{2})", text)
    if daily:
        return f"每天{int(daily.group(1)):02d}点"
    monthly = re.fullmatch(r"(\d{1,2})\s*,\s*(\d{1,2}):(\d{2})", text)
    if monthly:
        return f"每月{int(monthly.group(1))}号的{int(monthly.group(2)):02d}点"
    period = re.fullmatch(r"(\d+)(hour|min)", text, flags=re.IGNORECASE)
    if period:
        return f"每{period.group(1)}{period.group(2).lower()}"
    return text


def _format_launch_time_text(value):
    if isinstance(value, dt.datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, dt.date):
        return dt.datetime.combine(value, dt.time()).strftime("%Y-%m-%d %H:%M:%S")
    text = str(value or "").strip()
    if not text:
        return ""
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
        return f"{text} 00:00:00"
    return text


def _format_test_time(value):
    if isinstance(value, dt.datetime):
        return value.strftime("%Y%m%d")
    if isinstance(value, dt.date):
        return value.strftime("%Y%m%d")
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return dt.datetime.fromisoformat(text.replace("/", "-")).strftime("%Y%m%d")
    except ValueError:
        return text[:10].replace("-", "")[:8]


def _select_value(row, index):
    if index is None or index >= len(row):
        return ""
    return row[index]


def read_n02_tasks(excel_path):
    workbook = load_workbook(excel_path, data_only=True, read_only=False)
    sheet = select_ledger_sheet(workbook, required_headers=("任务名称", "执行策略", "成功/失败", "上线时间"))
    headers = [str(cell.value).strip() if cell.value is not None else "" for cell in sheet[1]]
    header_index = {header: idx for idx, header in enumerate(headers)}
    required = ("任务名称", "执行策略", "成功/失败", "上线时间", "来源委办", "表名")
    missing = [name for name in required if name not in header_index]
    if missing:
        raise ValueError(f"台账清单缺少必要列: {', '.join(missing)}")

    tasks = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        task_name = str(_select_value(row, header_index["任务名称"]) or "").strip()
        if not task_name:
            continue
        raw_strategy = _select_value(row, header_index["执行策略"])
        launch_time = _select_value(row, header_index["上线时间"])
        tasks.append(
            N02Task(
                task_name=task_name,
                strategy=str(raw_strategy or "").strip(),
                strategy_text=normalize_execution_strategy(raw_strategy),
                success_failure=str(_select_value(row, header_index["成功/失败"]) or "").strip(),
                launch_time=launch_time,
                launch_time_text=_format_launch_time_text(launch_time),
                test_time_text=_format_test_time(launch_time),
                source_department=str(_select_value(row, header_index["来源委办"]) or "").strip(),
                table_name=str(_select_value(row, header_index["表名"]) or "").strip(),
            )
        )

    workbook.close()
    if not tasks:
        raise ValueError(f"未找到匹配数据: 台账清单={excel_path}")
    return tasks


def render_task_screenshot(task, template_image_path, output_path):
    image = Image.open(template_image_path).convert("RGB")
    draw = ImageDraw.Draw(image)
    for box in (TASK_NAME_BOX, STRATEGY_BOX, SUCCESS_BOX, LAUNCH_BOX):
        _erase_text_pixels(image, box)

    font_ascii = _load_font(MIN_FONT_SIZE, chinese=False)
    font_cn = _load_font(MIN_FONT_SIZE, chinese=True)
    task_font = font_cn if _contains_cjk(task.task_name) else font_ascii
    draw.text(TASK_NAME_POS, task.task_name, font=task_font, fill=BLUE)
    draw.text(STRATEGY_POS, task.strategy_text, font=font_cn, fill=TEXT)
    _draw_success_failure(draw, task.success_failure, font_ascii)
    draw.text(LAUNCH_POS, task.launch_time_text, font=font_ascii, fill=TEXT)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)
    return str(output)


def _set_cell_text(cell, text):
    prototype = deepcopy(cell.paragraphs[0]._p)
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(clone_paragraph_with_text(prototype, text))


def _replace_image(cell, image_path, width):
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[-1]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=width)


def build_n02_test_document(excel_path, service_dir, template_path, output_path):
    tasks = read_n02_tasks(excel_path)
    document = Document(template_path)

    test_range_heading = find_heading(document, "Heading 1", "测试范围")
    test_environment_heading = find_heading(document, "Heading 1", "测试环境")
    test_case_heading = find_heading(document, "Heading 1", "测试案例")
    test_conclusion_heading = find_heading(document, "Heading 1", "测试结论")
    attachment_heading = find_heading(document, "Heading 1", "附件：测试结果截图")

    scope_elements = []
    current = test_range_heading._p.getnext()
    while current is not None and current is not test_environment_heading._p:
        scope_elements.append(current)
        current = current.getnext()
    scope_paragraph = next((item for item in scope_elements if item.tag.endswith("}p")), None)
    scope_table = next((item for item in scope_elements if item.tag.endswith("}tbl")), None)
    if scope_paragraph is None or scope_table is None:
        raise ValueError("模板中未找到测试范围的段落或表格原型")

    case_elements = []
    current = test_case_heading._p.getnext()
    while current is not None and current is not test_conclusion_heading._p:
        case_elements.append(current)
        current = current.getnext()
    case_title_prototype = next((item for item in case_elements if item.tag.endswith("}p")), None)
    case_table_prototype = next((item for item in case_elements if item.tag.endswith("}tbl")), None)
    conclusion_elements = []
    current = test_conclusion_heading._p.getnext()
    while current is not None and current is not attachment_heading._p:
        conclusion_elements.append(current)
        current = current.getnext()
    conclusion_paragraph = next((item for item in conclusion_elements if item.tag.endswith("}p")), None)
    if case_title_prototype is None or case_table_prototype is None or conclusion_paragraph is None:
        raise ValueError("模板中未找到测试案例或测试结论原型")

    template_image = Path(template_path).with_name("图片1.png")
    if not template_image.exists():
        raise FileNotFoundError(f"未找到截图模板图片: {template_image}")

    screenshot_width = document.inline_shapes[0].width if document.inline_shapes else None
    if screenshot_width is None:
        raise ValueError("模板中未找到截图原型")

    with tempfile.TemporaryDirectory(prefix="document_filler_n02_") as temp_dir:
        temp_dir_path = Path(temp_dir)
        scope_text = f"针对服务期内新增的{len(tasks)}个归集任务进行测试，涉及{len({task.source_department for task in tasks})}个委办。"
        conclusion_text = f"通过对上述{len(tasks)}个归集任务测试，任务均能正常运行，测试通过。"

        scope_table_data = [[str(index), task.task_name, task.source_department, task.table_name] for index, task in enumerate(tasks, start=1)]
        scope_table_element = clone_table_with_data(
            scope_table,
            ["序号", "任务名", "来源委办", "表名"],
            scope_table_data,
        )

        case_elements_out = []
        for index, task in enumerate(tasks, start=1):
            screenshot_path = temp_dir_path / f"{index:02d}_{_safe_name(task.task_name)}.png"
            render_task_screenshot(task, template_image, screenshot_path)
            case_title = clone_paragraph_with_text(case_title_prototype, task.task_name)
            case_table_xml = deepcopy(case_table_prototype)
            case_table = Table(case_table_xml, document)
            _set_cell_text(case_table.rows[7].cells[1], "栾希旺")
            _set_cell_text(case_table.rows[7].cells[3], task.test_time_text)
            _replace_image(case_table.rows[6].cells[1], screenshot_path, screenshot_width)
            case_elements_out.extend([case_title, case_table_xml])

        replace_between(
            test_range_heading._p,
            test_environment_heading._p,
            [clone_paragraph_with_text(scope_paragraph, scope_text), scope_table_element],
        )
        replace_between(
            test_case_heading._p,
            test_conclusion_heading._p,
            case_elements_out,
        )
        replace_between(
            test_conclusion_heading._p,
            attachment_heading._p,
            [clone_paragraph_with_text(conclusion_paragraph, conclusion_text)],
        )

        work_dir = temp_dir_path / "work"
        return save_word_document(document, output_path, work_dir)
