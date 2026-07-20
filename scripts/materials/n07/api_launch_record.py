import subprocess
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu, Inches

from materials.n07.api_code_doc import (
    _docx_rels,
    _extract_image,
    _image_rids,
    _image_width_emu,
    _iter_docx_blocks,
    _safe_name,
)
from materials.n07.api_requirement import _normalize_heading
from materials.shared.embedded_docx import extract_embedded_docx_by_work_order
from materials.shared.ledger import read_api_work_orders
from materials.shared.office_word import (
    convert_docx_to_legacy_doc,
    convert_legacy_doc_template,
    escape_powershell_string,
    save_word_document,
)
from materials.shared.word_sections import (
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    find_heading,
    replace_after,
    replace_between,
    update_toc_via_com,
)


LAUNCH_IMAGE_LABEL = "共享接口命名规范性"


@dataclass(frozen=True)
class ApiLaunchRecordPrototypes:
    list_table: object
    interface_heading: object
    image_paragraph: object


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _paragraph_style_id(element):
    properties = element.find(qn("w:pPr"))
    style = properties.find(qn("w:pStyle")) if properties is not None else None
    return style.get(qn("w:val")) if style is not None else None


def _first_heading_with_style(document, elements, style_name):
    style_id = document.styles[style_name].style_id
    return next(
        (
            item
            for item in elements
            if item.tag == qn("w:p") and _paragraph_style_id(item) == style_id
        ),
        None,
    )


def _elements_after(paragraph):
    elements = []
    current = paragraph._p.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        elements.append(current)
        current = current.getnext()
    return elements


def _capture_template_prototypes(document, list_heading, record_heading):
    list_elements = elements_between(list_heading._p, record_heading._p)
    list_table = next((item for item in list_elements if item.tag == qn("w:tbl")), None)
    record_elements = _elements_after(record_heading)
    interface_heading = _first_heading_with_style(document, record_elements, "Heading 2")
    image_paragraph = next(
        (item for item in record_elements if item.tag == qn("w:p") and _image_rids(item)),
        None,
    )
    if image_paragraph is None and interface_heading is not None:
        start = record_elements.index(interface_heading) + 1
        image_paragraph = next(
            (item for item in record_elements[start:] if item.tag == qn("w:p")),
            None,
        )
    return ApiLaunchRecordPrototypes(
        list_table=_require_prototype(list_table, "API作业上线清单表格"),
        interface_heading=_require_prototype(interface_heading, "API作业上线记录接口标题"),
        image_paragraph=_require_prototype(image_paragraph, "API作业上线截图段落"),
    )


def parse_api_launch_record_images(report_path, interfaces, work_dir):
    target_names = {_normalize_heading(item.chinese_name): item.chinese_name for item in interfaces}
    sections = {item.chinese_name: [] for item in interfaces}
    blocks = list(_iter_docx_blocks(report_path))
    image_dir = Path(work_dir) / "images"

    with zipfile.ZipFile(report_path, "r") as archive:
        relationships = _docx_rels(archive)
        in_test_content = False
        current_interface = ""
        current_label = ""
        for block in blocks:
            raw_text = str(block.get("text") or "").strip()
            normalized = _normalize_heading(raw_text.rstrip("：:"))
            if not in_test_content:
                if normalized == "测试内容":
                    in_test_content = True
                continue
            if normalized == "测试结论" or raw_text.startswith("附 工单截图"):
                break
            if normalized in target_names:
                current_interface = target_names[normalized]
                current_label = ""
                continue
            if current_interface and normalized == LAUNCH_IMAGE_LABEL:
                current_label = normalized
                continue
            if current_interface and current_label == LAUNCH_IMAGE_LABEL and block.get("images"):
                for rid in block["images"]:
                    sections[current_interface].append(
                        _extract_image(
                            archive,
                            relationships,
                            rid,
                            image_dir,
                            _safe_name(current_interface),
                        )
                    )
                continue
            if raw_text:
                current_label = ""

    missing = [interface.chinese_name for interface in interfaces if not sections[interface.chinese_name]]
    if missing:
        raise ValueError(f"自测报告中缺少作业上线截图: {', '.join(missing)}")
    return sections


def _build_list_table(orders, prototypes):
    rows = []
    index = 1
    for order in orders:
        for interface in order.interfaces:
            rows.append(
                [
                    str(index),
                    interface.english_name,
                    interface.chinese_name,
                    "上海市大数据中心",
                ]
            )
            index += 1
    return clone_table_with_data(
        prototypes.list_table,
        ["序号", "接口代码", "接口名称", "责任委办"],
        rows,
    )


def _image_paragraph_element(document, prototype, image_path):
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    width = _image_width_emu(prototype)
    if width:
        paragraph.add_run().add_picture(str(image_path), width=Emu(width))
    else:
        paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element.getparent().remove(element)
    return element


def _build_record_elements(document, orders, report_images_by_order, prototypes):
    elements = []
    for order in orders:
        for interface in order.interfaces:
            elements.append(
                clone_paragraph_with_text(prototypes.interface_heading, interface.chinese_name)
            )
            images = report_images_by_order[order.work_order_no][interface.chinese_name]
            elements.append(_image_paragraph_element(document, prototypes.image_paragraph, images[0]))
    return elements


def _convert_docx_to_legacy_doc(docx_path, output_path):
    return convert_docx_to_legacy_doc(docx_path, output_path, runner=subprocess.run)


def _escape_powershell_string(value):
    return escape_powershell_string(value)


def _convert_legacy_doc_template(template_path, work_dir):
    return convert_legacy_doc_template(template_path, work_dir, runner=subprocess.run)


def _save_document(document, output_path, work_dir):
    return save_word_document(
        document,
        output_path,
        work_dir,
        legacy_converter=_convert_docx_to_legacy_doc,
        toc_updater=update_toc_via_com,
    )


def build_api_launch_record_document(excel_path, service_dir, template_path, output_path):
    orders = read_api_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_api_launch_record_") as temp_dir:
        reports = extract_embedded_docx_by_work_order(
            excel_path,
            orders,
            "自测报告附件",
            Path(temp_dir) / "reports",
        )
        report_images_by_order = {}
        for order in orders:
            order.self_report_path = reports[order.work_order_no]
            report_images_by_order[order.work_order_no] = parse_api_launch_record_images(
                order.self_report_path,
                order.interfaces,
                Path(temp_dir) / order.work_order_no,
            )

        converted_template = _convert_legacy_doc_template(template_path, Path(temp_dir) / "template")
        document = Document(converted_template)
        list_heading = find_heading(document, "Heading 1", "API作业上线清单")
        record_heading = find_heading(document, "Heading 1", "API作业上线记录")
        prototypes = _capture_template_prototypes(document, list_heading, record_heading)

        replace_between(list_heading._p, record_heading._p, [_build_list_table(orders, prototypes)])
        replace_after(
            record_heading._p,
            _build_record_elements(document, orders, report_images_by_order, prototypes),
        )

        return _save_document(document, output_path, Path(temp_dir) / "output")
