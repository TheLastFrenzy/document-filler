import re
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from materials.n07.api_code_doc import (
    _docx_rels,
    _extract_image,
    _iter_docx_blocks,
    _safe_name,
)
from materials.n07.api_requirement import (
    NORMALIZED_PARAMETER_HEADERS,
    _normalize_heading,
    normalized_parameter_rows,
    parse_api_report,
)
from materials.shared.embedded_docx import extract_embedded_docx_by_work_order
from materials.shared.ledger import read_api_work_orders
from materials.shared.office_word import (
    convert_legacy_doc_template,
    escape_powershell_string,
)
from materials.shared.word_sections import (
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    replace_between,
    update_toc_via_com,
)


TEST_RESULT_LABEL = "测试结果"
TEST_RESULT_DESCRIPTION = "测试结果与预期结果一致，测试通过。"


@dataclass(frozen=True)
class ApiTestReportPrototypes:
    list_table: object
    interface_heading: object
    input_heading: object
    output_heading: object
    result_heading: object
    input_parameter_table: object
    output_parameter_table: object
    result_body: object


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _paragraph_style_id(element):
    properties = element.find(qn("w:pPr"))
    style = properties.find(qn("w:pStyle")) if properties is not None else None
    return style.get(qn("w:val")) if style is not None else None


def _first_paragraph_with_style(document, elements, style_name, text=None):
    style_id = document.styles[style_name].style_id
    for element in elements:
        if element.tag != qn("w:p") or _paragraph_style_id(element) != style_id:
            continue
        if text is None or _normalize_heading(element_text(element)) == _normalize_heading(text):
            return element
    return None


def _first_table_after(elements, paragraph):
    if paragraph is None:
        return None
    start = elements.index(paragraph) + 1
    return next((item for item in elements[start:] if item.tag == qn("w:tbl")), None)


def _first_body_after(elements, paragraph):
    if paragraph is None:
        return None
    start = elements.index(paragraph) + 1
    return next(
        (
            item
            for item in elements[start:]
            if item.tag == qn("w:p")
            and _paragraph_style_id(item) != _paragraph_style_id(paragraph)
        ),
        None,
    )


def _find_heading(document, style_name, text):
    target = _normalize_heading(text)
    for paragraph in document.paragraphs:
        if paragraph.style.name == style_name and _normalize_heading(paragraph.text) == target:
            return paragraph
    raise ValueError(f"模板中未找到{text}章节")


def _capture_template_prototypes(document, list_heading, content_heading, conclusion_heading):
    list_elements = elements_between(list_heading._p, content_heading._p)
    content_elements = elements_between(content_heading._p, conclusion_heading._p)
    list_table = next((item for item in list_elements if item.tag == qn("w:tbl")), None)
    interface_heading = _first_paragraph_with_style(document, content_elements, "Heading 2")
    input_heading = _first_paragraph_with_style(document, content_elements, "Heading 3", "输入参数")
    output_heading = _first_paragraph_with_style(document, content_elements, "Heading 3", "输出参数")
    result_heading = _first_paragraph_with_style(document, content_elements, "Heading 3", "测试结果")
    return ApiTestReportPrototypes(
        list_table=_require_prototype(list_table, "API测试清单表格"),
        interface_heading=_require_prototype(interface_heading, "API测试内容接口标题"),
        input_heading=_require_prototype(input_heading, "输入参数标题"),
        output_heading=_require_prototype(output_heading, "输出参数标题"),
        result_heading=_require_prototype(result_heading, "测试结果标题"),
        input_parameter_table=_require_prototype(
            _first_table_after(content_elements, input_heading), "输入参数表格"
        ),
        output_parameter_table=_require_prototype(
            _first_table_after(content_elements, output_heading), "输出参数表格"
        ),
        result_body=_require_prototype(
            _first_body_after(content_elements, result_heading), "测试结果截图段落"
        ),
    )


def parse_api_test_result_images(report_path, interfaces, work_dir):
    target_names = {_normalize_heading(item.chinese_name): item.chinese_name for item in interfaces}
    sections = {item.chinese_name: [] for item in interfaces}
    blocks = list(_iter_docx_blocks(report_path))
    image_dir = Path(work_dir) / "images"

    with zipfile.ZipFile(report_path, "r") as archive:
        relationships = _docx_rels(archive)
        in_test_content = False
        current_interface = ""
        waiting_for_result_image = False
        for block in blocks:
            raw_text = str(block.get("text") or "").strip()
            normalized = _normalize_heading(raw_text.rstrip("：:"))
            if not in_test_content:
                if normalized == "测试内容":
                    in_test_content = True
                continue
            if normalized == "测试结论" or raw_text.startswith("附 工单截图"):
                break
            if normalized in target_names:
                current_interface = target_names[normalized]
                waiting_for_result_image = False
                continue
            if current_interface and normalized == TEST_RESULT_LABEL:
                waiting_for_result_image = True
                for rid in block.get("images") or []:
                    sections[current_interface].append(
                        _extract_image(
                            archive,
                            relationships,
                            rid,
                            image_dir,
                            _safe_name(current_interface),
                        )
                    )
                continue
            if current_interface and waiting_for_result_image and block.get("images"):
                for rid in block["images"]:
                    sections[current_interface].append(
                        _extract_image(
                            archive,
                            relationships,
                            rid,
                            image_dir,
                            _safe_name(current_interface),
                        )
                    )
                continue
            if raw_text:
                waiting_for_result_image = False

    missing = [interface.chinese_name for interface in interfaces if not sections[interface.chinese_name]]
    if missing:
        raise ValueError(f"自测报告中缺少接口测试结果截图: {', '.join(missing)}")
    return sections


def _build_list_table(orders, prototypes):
    rows = []
    index = 1
    for order in orders:
        for interface in order.interfaces:
            rows.append(
                [
                    str(index),
                    interface.english_name,
                    interface.chinese_name,
                    "上海市大数据中心",
                ]
            )
            index += 1
    return clone_table_with_data(
        prototypes.list_table,
        ["序号", "接口代码", "接口名称", "责任委办"],
        rows,
    )


def _image_paragraph_element(document, prototype, image_path):
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element.getparent().remove(element)
    return element


def _replace_table_property(table, reference, property_name):
    properties = table.find(qn("w:tblPr"))
    reference_properties = reference.find(qn("w:tblPr"))
    if properties is None or reference_properties is None:
        return
    existing = properties.find(qn(property_name))
    if existing is not None:
        properties.remove(existing)
    reference_value = reference_properties.find(qn(property_name))
    if reference_value is not None:
        properties.append(deepcopy(reference_value))


def _grid_widths(table):
    grid = table.find(qn("w:tblGrid"))
    if grid is None:
        return []
    return [int(column.get(qn("w:w"), "0")) for column in grid.findall(qn("w:gridCol"))]


def _match_parameter_table_layout(table, reference):
    source_widths = _grid_widths(table)
    target_widths = _grid_widths(reference)
    source_total = sum(source_widths)
    target_total = sum(target_widths)
    if not source_widths or source_total <= 0 or target_total <= 0:
        return table

    scaled_widths = [round(width * target_total / source_total) for width in source_widths]
    scaled_widths[-1] += target_total - sum(scaled_widths)
    for property_name in ("w:tblW", "w:tblInd", "w:tblLayout"):
        _replace_table_property(table, reference, property_name)

    grid = table.find(qn("w:tblGrid"))
    for column, width in zip(grid.findall(qn("w:gridCol")), scaled_widths):
        column.set(qn("w:w"), str(width))

    for row in table.findall(qn("w:tr")):
        logical_index = 0
        for cell in row.findall(qn("w:tc")):
            properties = cell.find(qn("w:tcPr"))
            span_node = properties.find(qn("w:gridSpan")) if properties is not None else None
            span = int(span_node.get(qn("w:val"), "1")) if span_node is not None else 1
            if properties is None:
                properties = deepcopy(reference.find(".//" + qn("w:tcPr")))
                cell.insert(0, properties)
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = deepcopy(reference.find(".//" + qn("w:tcW")))
                properties.insert(0, cell_width)
            width = sum(scaled_widths[logical_index : logical_index + span])
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            logical_index += span
    return table


def _build_content_elements(document, orders, report_images_by_order, prototypes):
    elements = []
    for order in orders:
        for interface in order.interfaces:
            elements.append(
                clone_paragraph_with_text(
                    prototypes.interface_heading,
                    f"{interface.chinese_name} {interface.english_name}",
                )
            )
            elements.append(clone_paragraph_with_text(prototypes.input_heading, "输入参数"))
            input_table = clone_table_with_data(
                prototypes.input_parameter_table,
                NORMALIZED_PARAMETER_HEADERS,
                normalized_parameter_rows(interface.input_groups),
            )
            elements.append(
                _match_parameter_table_layout(input_table, prototypes.output_parameter_table)
            )
            elements.append(clone_paragraph_with_text(prototypes.output_heading, "输出参数"))
            elements.append(
                clone_table_with_data(
                    prototypes.output_parameter_table,
                    NORMALIZED_PARAMETER_HEADERS,
                    normalized_parameter_rows(interface.output_groups),
                )
            )
            elements.append(clone_paragraph_with_text(prototypes.result_heading, "测试结果"))
            images = report_images_by_order[order.work_order_no][interface.chinese_name]
            elements.append(clone_paragraph_with_text(prototypes.result_body, TEST_RESULT_DESCRIPTION))
            elements.append(_image_paragraph_element(document, prototypes.result_body, images[0]))
    return elements


def _escape_powershell_string(value):
    return escape_powershell_string(value)


def _convert_legacy_doc_template(template_path, work_dir):
    return convert_legacy_doc_template(template_path, work_dir, runner=subprocess.run)


def build_api_test_report_document(excel_path, service_dir, template_path, output_path):
    orders = read_api_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_api_test_report_") as temp_dir:
        reports = extract_embedded_docx_by_work_order(
            excel_path,
            orders,
            "自测报告附件",
            Path(temp_dir) / "reports",
        )
        report_images_by_order = {}
        for order in orders:
            order.self_report_path = reports[order.work_order_no]
            parse_api_report(order.self_report_path, order.interfaces)
            report_images_by_order[order.work_order_no] = parse_api_test_result_images(
                order.self_report_path,
                order.interfaces,
                Path(temp_dir) / order.work_order_no,
            )

        converted_template = _convert_legacy_doc_template(template_path, Path(temp_dir) / "template")
        document = Document(converted_template)
        list_heading = _find_heading(document, "Heading 1", "API测试清单")
        content_heading = _find_heading(document, "Heading 1", "API测试内容")
        conclusion_heading = _find_heading(document, "Heading 1", "测试结论")
        prototypes = _capture_template_prototypes(
            document,
            list_heading,
            content_heading,
            conclusion_heading,
        )

        replace_between(list_heading._p, content_heading._p, [_build_list_table(orders, prototypes)])
        replace_between(
            content_heading._p,
            conclusion_heading._p,
            _build_content_elements(document, orders, report_images_by_order, prototypes),
        )

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)

    update_toc_via_com(output)
    return str(output)
