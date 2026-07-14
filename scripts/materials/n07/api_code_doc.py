import posixpath
import re
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu

from materials.n07.api_requirement import _normalize_heading
from materials.shared.embedded_docx import extract_embedded_docx_by_work_order
from materials.shared.ledger import read_api_work_orders
from materials.shared.word_sections import (
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    find_heading,
    replace_after,
    replace_between,
    update_toc_via_com,
)


IMAGE_LABELS = ("共享接口命名规范性", "共享任务开发代码检查")
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


@dataclass(frozen=True)
class ApiCodePrototypes:
    list_table: object
    interface_heading: object
    naming_image_paragraph: object
    code_image_paragraph: object


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _paragraph_style_id(element):
    properties = element.find(qn("w:pPr"))
    style = properties.find(qn("w:pStyle")) if properties is not None else None
    return style.get(qn("w:val")) if style is not None else None


def _first_heading_with_style(document, elements, style_name):
    style_id = document.styles[style_name].style_id
    return next(
        (
            item
            for item in elements
            if item.tag == qn("w:p") and _paragraph_style_id(item) == style_id
        ),
        None,
    )


def _image_rids(element):
    rids = []
    for blip in element.findall(".//a:blip", NS):
        rid = blip.attrib.get(f"{{{NS['r']}}}embed")
        if rid:
            rids.append(rid)
    for image_data in element.findall(".//v:imagedata", NS):
        rid = image_data.attrib.get(f"{{{NS['r']}}}id")
        if rid:
            rids.append(rid)
    return rids


def _capture_template_prototypes(document, list_heading, detail_heading):
    list_elements = elements_between(list_heading._p, detail_heading._p)
    list_table = next((item for item in list_elements if item.tag == qn("w:tbl")), None)

    detail_elements = []
    current = detail_heading._p.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        detail_elements.append(current)
        current = current.getnext()

    interface_heading = _first_heading_with_style(document, detail_elements, "Heading 2")
    image_paragraphs = [
        item for item in detail_elements if item.tag == qn("w:p") and _image_rids(item)
    ]
    naming_image = image_paragraphs[0] if image_paragraphs else None
    code_image = image_paragraphs[1] if len(image_paragraphs) > 1 else None
    return ApiCodePrototypes(
        list_table=_require_prototype(list_table, "API接口开发列表表格"),
        interface_heading=_require_prototype(interface_heading, "接口模型明细标题"),
        naming_image_paragraph=_require_prototype(naming_image, "共享接口命名规范性截图段落"),
        code_image_paragraph=_require_prototype(code_image, "共享任务开发代码检查截图段落"),
    )


def _docx_rels(archive):
    root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    return {
        element.attrib["Id"]: element.attrib["Target"]
        for element in root
        if "Id" in element.attrib and "Target" in element.attrib
    }


def _block_text(element):
    return "".join(node.text or "" for node in element.findall(".//w:t", NS)).strip()


def _iter_docx_blocks(path):
    with zipfile.ZipFile(path, "r") as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    body = root.find("w:body", NS)
    if body is None:
        return
    for child in list(body):
        if child.tag == f"{{{NS['w']}}}p":
            yield {"type": "paragraph", "text": _block_text(child), "images": _image_rids(child)}
        elif child.tag == f"{{{NS['w']}}}tbl":
            text_values = []
            images = []
            for paragraph in child.findall(".//w:p", NS):
                text = _block_text(paragraph)
                if text:
                    text_values.append(text)
                images.extend(_image_rids(paragraph))
            yield {"type": "table", "text": " | ".join(text_values), "images": images}


def _image_target_path(target):
    if target.startswith("/"):
        return target.lstrip("/")
    return posixpath.normpath("word/" + target).lstrip("/")


def _extract_image(archive, relationships, rid, image_dir, prefix):
    target = relationships.get(rid)
    if not target:
        raise KeyError(f"图片关系 {rid} 不存在")
    zip_path = _image_target_path(target)
    suffix = Path(target).suffix or ".png"
    image_dir.mkdir(parents=True, exist_ok=True)
    output = image_dir / f"{prefix}_{rid}{suffix}"
    if not output.exists():
        output.write_bytes(archive.read(zip_path))
    return output


def _safe_name(value):
    return re.sub(r"[^0-9A-Za-z_\u4e00-\u9fff.-]+", "_", value).strip("_") or "api"


def parse_api_code_report_images(report_path, interfaces, work_dir):
    target_names = {_normalize_heading(item.chinese_name): item.chinese_name for item in interfaces}
    sections = {item.chinese_name: {label: [] for label in IMAGE_LABELS} for item in interfaces}
    blocks = list(_iter_docx_blocks(report_path))
    image_dir = Path(work_dir) / "images"

    with zipfile.ZipFile(report_path, "r") as archive:
        relationships = _docx_rels(archive)
        in_test_content = False
        current_interface = ""
        current_label = ""
        for block in blocks:
            raw_text = str(block.get("text") or "").strip()
            normalized = _normalize_heading(raw_text.rstrip("：:"))
            if not in_test_content:
                if normalized == "测试内容":
                    in_test_content = True
                continue
            if normalized == "测试结论" or raw_text.startswith("附 工单截图"):
                break
            if normalized in target_names:
                current_interface = target_names[normalized]
                current_label = ""
                continue
            if raw_text and normalized.endswith("接口") and not block.get("images"):
                current_interface = ""
                current_label = ""
                continue
            if current_interface and normalized in IMAGE_LABELS:
                current_label = normalized
                continue
            if current_interface and current_label and block.get("images"):
                for rid in block["images"]:
                    sections[current_interface][current_label].append(
                        _extract_image(
                            archive,
                            relationships,
                            rid,
                            image_dir,
                            _safe_name(current_interface),
                        )
                    )
                continue
            if raw_text:
                current_label = ""

    missing = []
    for interface in interfaces:
        for label in IMAGE_LABELS:
            if not sections[interface.chinese_name][label]:
                missing.append(f"{interface.chinese_name}/{label}")
    if missing:
        raise ValueError(f"自测报告中缺少接口开发代码截图: {', '.join(missing)}")
    return sections


def _image_width_emu(prototype):
    extent = prototype.find(".//wp:extent", NS)
    if extent is None:
        return None
    value = extent.attrib.get("cx")
    return int(value) if value and value.isdigit() else None


def _image_paragraph_element(document, prototype, image_path):
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    width = _image_width_emu(prototype)
    if width:
        paragraph.add_run().add_picture(str(image_path), width=Emu(width))
    else:
        paragraph.add_run().add_picture(str(image_path))
    element.getparent().remove(element)
    return element


def _build_list_table(orders, prototypes):
    rows = []
    index = 1
    for order in orders:
        for interface in order.interfaces:
            rows.append(
                [
                    str(index),
                    interface.english_name,
                    interface.chinese_name,
                    "上海市大数据中心",
                ]
            )
            index += 1
    return clone_table_with_data(
        prototypes.list_table,
        ["序号", "接口代码", "接口名称", "责任委办"],
        rows,
    )


def _build_detail_elements(document, orders, report_images_by_order, prototypes):
    elements = []
    for order in orders:
        for interface in order.interfaces:
            elements.append(
                clone_paragraph_with_text(prototypes.interface_heading, interface.chinese_name)
            )
            images = report_images_by_order[order.work_order_no][interface.chinese_name]
            elements.append(
                _image_paragraph_element(
                    document,
                    prototypes.naming_image_paragraph,
                    images["共享接口命名规范性"][0],
                )
            )
            elements.append(
                _image_paragraph_element(
                    document,
                    prototypes.code_image_paragraph,
                    images["共享任务开发代码检查"][0],
                )
            )
    return elements


def build_api_code_document(excel_path, service_dir, template_path, output_path):
    orders = read_api_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_api_code_") as temp_dir:
        reports = extract_embedded_docx_by_work_order(
            excel_path,
            orders,
            "自测报告附件",
            Path(temp_dir) / "reports",
        )
        report_images_by_order = {}
        for order in orders:
            order.self_report_path = reports[order.work_order_no]
            report_images_by_order[order.work_order_no] = parse_api_code_report_images(
                order.self_report_path,
                order.interfaces,
                Path(temp_dir) / order.work_order_no,
            )

        document = Document(template_path)
        list_heading = find_heading(document, "Heading 1", "API接口开发列表")
        detail_heading = find_heading(document, "Heading 1", "接口模型明细")
        prototypes = _capture_template_prototypes(document, list_heading, detail_heading)

        replace_between(list_heading._p, detail_heading._p, [_build_list_table(orders, prototypes)])
        replace_after(
            detail_heading._p,
            _build_detail_elements(document, orders, report_images_by_order, prototypes),
        )

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
    update_toc_via_com(output)
    return str(output)
