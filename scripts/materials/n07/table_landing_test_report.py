import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from materials.n07.table_landing import (
    extract_ledger_images_by_row,
    read_table_landing_work_orders,
)
from materials.shared.office_word import (
    convert_legacy_doc_template as _convert_legacy_doc_template,
    save_word_document as _save_document,
)
from materials.shared.word_sections import (
    clone_element,
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    find_heading,
    replace_after,
    replace_between,
)


LAUNCH_IMAGE_COLUMNS = ("上线交付截图1", "上线交付截图2")
TEST_RESULT_DESCRIPTION = "测试结果与预期结果一致，测试通过。"


@dataclass(frozen=True)
class TableLandingTestItem:
    source_row: int
    source_table: str
    dispatch_description: str
    business_scene: str
    launch_images: tuple[bytes, ...]


@dataclass(frozen=True)
class TableLandingTestReportPrototypes:
    list_table: object
    item_heading: object
    item_body: object
    image_paragraph: object
    result_paragraph: object
    conclusion_table: object
    conclusion_elements: tuple[object, ...]


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _paragraph_style_id(element):
    properties = element.find(qn("w:pPr"))
    style = properties.find(qn("w:pStyle")) if properties is not None else None
    return style.get(qn("w:val")) if style is not None else None


def _first_table(elements, column_count=None):
    for element in elements:
        if element.tag != qn("w:tbl"):
            continue
        if column_count is None:
            return element
        grid = element.find(qn("w:tblGrid"))
        if grid is not None and len(grid.findall(qn("w:gridCol"))) == column_count:
            return element
        first_row = element.find(qn("w:tr"))
        if first_row is not None and len(first_row.findall(qn("w:tc"))) == column_count:
            return element
    return None


def _first_paragraph_with_style(document, elements, style_name):
    style_id = document.styles[style_name].style_id
    return next(
        (
            item
            for item in elements
            if item.tag == qn("w:p")
            and _paragraph_style_id(item) == style_id
            and element_text(item)
        ),
        None,
    )


def _paragraph_has_image(element):
    return bool(list(element.iter(qn("w:drawing"))) or list(element.iter(qn("w:pict"))))


def _first_nonempty_paragraph(elements):
    return next(
        (
            item
            for item in elements
            if item.tag == qn("w:p") and element_text(item) and not _paragraph_has_image(item)
        ),
        None,
    )


def _capture_template_prototypes(document, list_heading, content_heading, conclusion_heading):
    list_elements = elements_between(list_heading._p, content_heading._p)
    list_table = _first_table(list_elements, 3)

    content_elements = elements_between(content_heading._p, conclusion_heading._p)
    item_heading = _first_paragraph_with_style(document, content_elements, "Heading 2")
    start = content_elements.index(item_heading) + 1 if item_heading is not None else 0
    content_after_heading = content_elements[start:]
    item_body = _first_nonempty_paragraph(content_after_heading)
    image_paragraph = next(
        (
            item
            for item in content_after_heading
            if item.tag == qn("w:p") and _paragraph_has_image(item)
        ),
        None,
    )
    if image_paragraph is None:
        image_paragraph = next(
            (item for item in content_after_heading if item.tag == qn("w:p")),
            None,
        )
    result_paragraph = next(
        (
            item
            for item in content_after_heading
            if item.tag == qn("w:p")
            and (
                element_text(item).startswith("测试结论")
                or TEST_RESULT_DESCRIPTION in element_text(item)
            )
        ),
        None,
    )

    conclusion_elements = []
    current = conclusion_heading._p.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        conclusion_elements.append(current)
        current = current.getnext()
    conclusion_table = _first_table(conclusion_elements, 6)

    return TableLandingTestReportPrototypes(
        list_table=_require_prototype(list_table, "库表落地测试清单表格"),
        item_heading=_require_prototype(item_heading, "库表落地测试内容标题"),
        item_body=_require_prototype(item_body, "库表落地测试内容正文"),
        image_paragraph=_require_prototype(image_paragraph, "库表落地测试内容截图段落"),
        result_paragraph=_require_prototype(result_paragraph, "库表落地测试结论段落"),
        conclusion_table=_require_prototype(conclusion_table, "测试结论表格"),
        conclusion_elements=tuple(conclusion_elements),
    )


def _safe_filename(value):
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", str(value or "")).strip(" ._") or "image"


def _image_suffix(image_bytes):
    if image_bytes.startswith(b"\xff\xd8"):
        return ".jpg"
    if image_bytes.startswith(b"GIF"):
        return ".gif"
    return ".png"


def _image_paragraph_element(document, prototype, image_bytes, image_path):
    image_path.write_bytes(image_bytes)
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element.getparent().remove(element)
    return element


def _task_for_source(order, index):
    return order.tasks[index] if index < len(order.tasks) else None


def _build_test_items(excel_path, orders):
    source_rows = [row for order in orders for row in order.source_rows]
    image_map = extract_ledger_images_by_row(excel_path, source_rows, LAUNCH_IMAGE_COLUMNS)
    items = []
    for order in orders:
        for index, source_table in enumerate(order.source_tables):
            source_row = order.source_rows[index] if index < len(order.source_rows) else order.source_rows[-1]
            task = _task_for_source(order, index)
            images = tuple(image_map.get((source_row, column)) for column in LAUNCH_IMAGE_COLUMNS)
            missing = [
                column
                for column, image_bytes in zip(LAUNCH_IMAGE_COLUMNS, images)
                if not image_bytes
            ]
            if missing:
                raise ValueError(
                    f"台账第{source_row}行缺少库表落地上线交付截图: {', '.join(missing)}"
                )
            items.append(
                TableLandingTestItem(
                    source_row=source_row,
                    source_table=source_table,
                    dispatch_description=(
                        task.dispatch_description if task and task.dispatch_description else source_table
                    ),
                    business_scene=task.business_scene if task and task.business_scene else source_table,
                    launch_images=images,
                )
            )
    return items


def _build_list_table(items, prototypes):
    rows = [
        [str(index), item.source_table, item.dispatch_description]
        for index, item in enumerate(items, start=1)
    ]
    return clone_table_with_data(
        prototypes.list_table,
        ["序号", "调度名", "调度中文名"],
        rows,
    )


def _build_content_elements(document, items, prototypes, work_dir):
    elements = []
    image_dir = Path(work_dir) / "launch_images"
    image_dir.mkdir(parents=True, exist_ok=True)
    for index, item in enumerate(items, start=1):
        elements.append(clone_paragraph_with_text(prototypes.item_heading, item.source_table))
        elements.append(clone_paragraph_with_text(prototypes.item_body, item.source_table))
        for image_index, image_bytes in enumerate(item.launch_images, start=1):
            suffix = _image_suffix(image_bytes)
            image_name = f"{index}_{image_index}_{_safe_filename(item.source_table)}{suffix}"
            elements.append(
                _image_paragraph_element(
                    document,
                    prototypes.image_paragraph,
                    image_bytes,
                    image_dir / image_name,
                )
            )
        elements.append(
            clone_paragraph_with_text(
                prototypes.result_paragraph,
                f"测试结论：{TEST_RESULT_DESCRIPTION}",
            )
        )
    return elements


def _build_conclusion_table(items, prototypes):
    rows = [
        [
            str(index),
            item.source_table,
            item.business_scene,
            "程序调度",
            "DACP",
            "通过",
        ]
        for index, item in enumerate(items, start=1)
    ]
    return clone_table_with_data(
        prototypes.conclusion_table,
        ["序号", "测试调度清单", "业务场景", "测试方法", "测试工具", "测试结论"],
        rows,
    )


def _build_conclusion_elements(items, prototypes):
    elements = []
    replaced = False
    for element in prototypes.conclusion_elements:
        if not replaced and element is prototypes.conclusion_table:
            elements.append(_build_conclusion_table(items, prototypes))
            replaced = True
            continue
        elements.append(clone_element(element))
    if not replaced:
        elements.append(_build_conclusion_table(items, prototypes))
    return elements


def build_table_landing_test_report_document(excel_path, service_dir, template_path, output_path):
    orders = read_table_landing_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_table_landing_test_report_") as temp_dir:
        items = _build_test_items(excel_path, orders)
        converted_template = _convert_legacy_doc_template(template_path, Path(temp_dir) / "template")
        document = Document(converted_template)
        list_heading = find_heading(document, "Heading 1", "库表落地测试清单")
        content_heading = find_heading(document, "Heading 1", "库表落地测试内容")
        conclusion_heading = find_heading(document, "Heading 1", "测试结论")
        prototypes = _capture_template_prototypes(
            document,
            list_heading,
            content_heading,
            conclusion_heading,
        )

        replace_between(list_heading._p, content_heading._p, [_build_list_table(items, prototypes)])
        replace_between(
            content_heading._p,
            conclusion_heading._p,
            _build_content_elements(document, items, prototypes, Path(temp_dir)),
        )
        replace_after(conclusion_heading._p, _build_conclusion_elements(items, prototypes))

        return _save_document(document, output_path, Path(temp_dir) / "output")
