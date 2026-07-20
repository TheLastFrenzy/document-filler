import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from materials.n07.table_landing import read_table_landing_work_orders
from materials.shared.office_word import (
    convert_legacy_doc_template as _convert_legacy_doc_template,
    save_word_document as _save_document,
)
from materials.shared.word_sections import (
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    find_heading,
    replace_after,
    replace_between,
)


VERIFY_RESULT_TEXT = "期望记录条数与验证记录条数一致，验证通过；"


@dataclass(frozen=True)
class TableLandingShareRecordPrototypes:
    list_table: object
    item_heading: object
    expected_count: object
    source_image: object
    actual_count: object
    target_image: object
    verify_result: object


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _paragraph_style_id(element):
    properties = element.find(qn("w:pPr"))
    style = properties.find(qn("w:pStyle")) if properties is not None else None
    return style.get(qn("w:val")) if style is not None else None


def _first_table(elements, column_count=None):
    for element in elements:
        if element.tag != qn("w:tbl"):
            continue
        if column_count is None:
            return element
        grid = element.find(qn("w:tblGrid"))
        if grid is not None and len(grid.findall(qn("w:gridCol"))) == column_count:
            return element
        first_row = element.find(qn("w:tr"))
        if first_row is not None and len(first_row.findall(qn("w:tc"))) == column_count:
            return element
    return None


def _paragraph_has_image(element):
    return bool(list(element.iter(qn("w:drawing"))) or list(element.iter(qn("w:pict"))))


def _first_paragraph_with_style(document, elements, style_name):
    style_id = document.styles[style_name].style_id
    return next(
        (
            item
            for item in elements
            if item.tag == qn("w:p")
            and _paragraph_style_id(item) == style_id
            and element_text(item)
        ),
        None,
    )


def _paragraph_with_prefix(elements, prefix, start=0):
    for index in range(start, len(elements)):
        element = elements[index]
        if element.tag == qn("w:p") and element_text(element).startswith(prefix):
            return element, index
    return None, -1


def _first_image_paragraph(elements, start=0):
    for index in range(start, len(elements)):
        element = elements[index]
        if element.tag == qn("w:p") and _paragraph_has_image(element):
            return element, index
    for index in range(start, len(elements)):
        element = elements[index]
        if element.tag == qn("w:p"):
            return element, index
    return None, -1


def _capture_template_prototypes(document, list_heading, record_heading):
    list_elements = elements_between(list_heading._p, record_heading._p)
    list_table = _first_table(list_elements, 4)

    record_elements = []
    current = record_heading._p.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        record_elements.append(current)
        current = current.getnext()

    item_heading = _first_paragraph_with_style(document, record_elements, "Heading 2")
    start = record_elements.index(item_heading) + 1 if item_heading is not None else 0
    expected_count, expected_index = _paragraph_with_prefix(record_elements, "期望记录条数", start)
    source_image, source_index = _first_image_paragraph(record_elements, expected_index + 1)
    actual_count, actual_index = _paragraph_with_prefix(record_elements, "验证记录条数", source_index + 1)
    target_image, target_index = _first_image_paragraph(record_elements, actual_index + 1)
    verify_result, _result_index = _paragraph_with_prefix(record_elements, "验证结果", target_index + 1)

    return TableLandingShareRecordPrototypes(
        list_table=_require_prototype(list_table, "数据库表落地验证清单表格"),
        item_heading=_require_prototype(item_heading, "数据库表落地验证记录标题"),
        expected_count=_require_prototype(expected_count, "期望记录条数段落"),
        source_image=_require_prototype(source_image, "源表数据量截图段落"),
        actual_count=_require_prototype(actual_count, "验证记录条数段落"),
        target_image=_require_prototype(target_image, "目标表数据量截图段落"),
        verify_result=_require_prototype(verify_result, "验证结果段落"),
    )


def _source_table_for_task(order, index):
    if index < len(order.source_tables):
        return order.source_tables[index]
    return order.source_tables[-1] if order.source_tables else ""


def _dispatch_description_for_task(order, index, task):
    if index < len(order.dispatch_descriptions) and order.dispatch_descriptions[index]:
        return order.dispatch_descriptions[index]
    return task.dispatch_description or _source_table_for_task(order, index)


def _all_order_tasks(orders):
    for order in orders:
        for index, task in enumerate(order.tasks):
            yield order, index, task


def _build_list_table(orders, prototypes):
    rows = [
        [
            str(index),
            _source_table_for_task(order, task_index),
            _dispatch_description_for_task(order, task_index, task),
            task.business_scene,
        ]
        for index, (order, task_index, task) in enumerate(_all_order_tasks(orders), start=1)
    ]
    return clone_table_with_data(
        prototypes.list_table,
        ["序号", "调度名", "调度中文名", "场景名称"],
        rows,
    )


def _safe_filename(value):
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", str(value or "")).strip(" ._") or "image"


def _image_suffix(image_bytes):
    if image_bytes.startswith(b"\xff\xd8"):
        return ".jpg"
    if image_bytes.startswith(b"GIF"):
        return ".gif"
    return ".png"


def _image_paragraph_element(document, prototype, image_bytes, image_path):
    image_path.write_bytes(image_bytes)
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element.getparent().remove(element)
    return element


def _build_record_elements(document, orders, prototypes, work_dir):
    elements = []
    image_dir = Path(work_dir) / "share_record_images"
    image_dir.mkdir(parents=True, exist_ok=True)
    for index, (order, task_index, task) in enumerate(_all_order_tasks(orders), start=1):
        source_table = _source_table_for_task(order, task_index)
        elements.append(clone_paragraph_with_text(prototypes.item_heading, source_table))
        elements.append(
            clone_paragraph_with_text(prototypes.expected_count, f"期望记录条数：{task.landing_data_count}")
        )
        if not task.source_volume_image:
            raise ValueError(f"工单{order.work_order_no}的{source_table}缺少源表数据量截图")
        if not task.target_volume_image:
            raise ValueError(f"工单{order.work_order_no}的{source_table}缺少目标表数据量截图")
        source_suffix = _image_suffix(task.source_volume_image)
        target_suffix = _image_suffix(task.target_volume_image)
        safe_name = _safe_filename(source_table)
        elements.append(
            _image_paragraph_element(
                document,
                prototypes.source_image,
                task.source_volume_image,
                image_dir / f"{index}_source_{safe_name}{source_suffix}",
            )
        )
        elements.append(
            clone_paragraph_with_text(prototypes.actual_count, f"验证记录条数：{task.landing_data_count}")
        )
        elements.append(
            _image_paragraph_element(
                document,
                prototypes.target_image,
                task.target_volume_image,
                image_dir / f"{index}_target_{safe_name}{target_suffix}",
            )
        )
        elements.append(
            clone_paragraph_with_text(prototypes.verify_result, f"验证结果：{VERIFY_RESULT_TEXT}")
        )
    return elements


def build_table_landing_share_record_document(excel_path, service_dir, template_path, output_path):
    orders = read_table_landing_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_table_landing_share_record_") as temp_dir:
        converted_template = _convert_legacy_doc_template(template_path, Path(temp_dir) / "template")
        document = Document(converted_template)
        list_heading = find_heading(document, "Heading 1", "数据库表落地验证清单")
        record_heading = find_heading(document, "Heading 1", "数据库表落地验证记录")
        prototypes = _capture_template_prototypes(document, list_heading, record_heading)

        replace_between(list_heading._p, record_heading._p, [_build_list_table(orders, prototypes)])
        replace_after(record_heading._p, _build_record_elements(document, orders, prototypes, Path(temp_dir)))

        return _save_document(document, output_path, Path(temp_dir) / "output")
