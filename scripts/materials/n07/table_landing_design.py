import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from materials.n07.table_landing import read_table_landing_work_orders
from materials.shared.office_word import (
    convert_legacy_doc_template as _convert_legacy_doc_template,
    save_word_document as _save_document,
)
from materials.shared.word_sections import (
    clone_element,
    clone_paragraph_with_text,
    clone_table_with_data,
    element_text,
    elements_between,
    find_heading,
    replace_after,
    replace_between,
)


@dataclass(frozen=True)
class TableLandingDesignPrototypes:
    logic_table: object
    target_name: object
    target_image: object
    detail_table: object
    push_table: object
    push_elements: tuple[object, ...]


def _require_prototype(value, name):
    if value is None:
        raise ValueError(f"模板中未找到{name}格式原型")
    return value


def _first_table(elements, column_count=None):
    for element in elements:
        if element.tag != qn("w:tbl"):
            continue
        if column_count is None:
            return element
        grid = element.find(qn("w:tblGrid"))
        if grid is not None and len(grid.findall(qn("w:gridCol"))) == column_count:
            return element
        first_row = element.find(qn("w:tr"))
        if first_row is not None and len(first_row.findall(qn("w:tc"))) == column_count:
            return element
    return None


def _first_paragraph_after(elements, start_index):
    for index in range(start_index, len(elements)):
        element = elements[index]
        if element.tag == qn("w:p"):
            return element
    return None


def _paragraph_has_image(element):
    return bool(list(element.iter(qn("w:drawing"))) or list(element.iter(qn("w:pict"))))


def _capture_template_prototypes(logic_heading, detail_heading, push_heading):
    logic_elements = elements_between(logic_heading._p, detail_heading._p)
    logic_table = _first_table(logic_elements, 6)
    logic_table_index = logic_elements.index(logic_table) if logic_table is not None else -1
    target_name = next(
        (
            item
            for item in logic_elements[logic_table_index + 1 :]
            if item.tag == qn("w:p") and element_text(item)
        ),
        None,
    )
    target_image = next(
        (
            item
            for item in logic_elements[logic_table_index + 1 :]
            if item.tag == qn("w:p") and _paragraph_has_image(item)
        ),
        None,
    )
    if target_image is None:
        target_image = _first_paragraph_after(logic_elements, logic_table_index + 1)

    detail_elements = elements_between(detail_heading._p, push_heading._p)
    detail_table = _first_table(detail_elements, 6)

    push_elements = []
    current = push_heading._p.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        push_elements.append(current)
        current = current.getnext()
    push_table = _first_table(push_elements, 4)

    return TableLandingDesignPrototypes(
        logic_table=_require_prototype(logic_table, "业务逻辑说明表格"),
        target_name=_require_prototype(target_name, "目标表名段落"),
        target_image=_require_prototype(target_image, "目标表数据量截图段落"),
        detail_table=_require_prototype(detail_table, "库表说明表格"),
        push_table=_require_prototype(push_table, "推送频率表格"),
        push_elements=tuple(push_elements),
    )


def _source_table_for_task(order, index):
    if index < len(order.source_tables):
        return order.source_tables[index]
    return order.source_tables[-1] if order.source_tables else ""


def _all_order_tasks(orders):
    for order in orders:
        for index, task in enumerate(order.tasks):
            yield order, index, task


def _build_logic_table(orders, prototypes):
    rows = [
        [
            str(index),
            task.landing_database,
            order.database_type,
            task.landing_table,
            order.update_cycle,
            task.landing_data_count,
        ]
        for index, (order, _task_index, task) in enumerate(_all_order_tasks(orders), start=1)
    ]
    return clone_table_with_data(
        prototypes.logic_table,
        ["序号", "目标库", "目标库类型", "目标表", "频率", "规格（记录数）"],
        rows,
    )


def _image_paragraph_element(document, prototype, image_bytes, image_path):
    image_path.write_bytes(image_bytes)
    paragraph = document.add_paragraph()
    element = paragraph._p
    properties = element.find(qn("w:pPr"))
    if properties is not None:
        element.remove(properties)
    prototype_properties = prototype.find(qn("w:pPr"))
    if prototype_properties is not None:
        element.insert(0, deepcopy(prototype_properties))
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element.getparent().remove(element)
    return element


def _build_target_volume_elements(document, orders, prototypes, work_dir):
    elements = []
    image_dir = Path(work_dir) / "target_volume_images"
    image_dir.mkdir(parents=True, exist_ok=True)
    for index, (order, _task_index, task) in enumerate(_all_order_tasks(orders), start=1):
        elements.append(
            clone_paragraph_with_text(prototypes.target_name, f"目标表名：{task.landing_table}")
        )
        if not task.target_volume_image:
            raise ValueError(f"工单{order.work_order_no}的{task.landing_table}缺少目标表数据量截图")
        elements.append(
            _image_paragraph_element(
                document,
                prototypes.target_image,
                task.target_volume_image,
                image_dir / f"{index}_{task.landing_table}.png",
            )
        )
    return elements


def _build_detail_table(orders, prototypes):
    rows = [
        [
            str(index),
            task.landing_database,
            task.landing_table,
            order.target_user,
            _source_table_for_task(order, task_index),
            task.business_scene,
        ]
        for index, (order, task_index, task) in enumerate(_all_order_tasks(orders), start=1)
    ]
    return clone_table_with_data(
        prototypes.detail_table,
        ["序号", "落地库", "落地表", "对象用户", "库表来源", "业务场景"],
        rows,
    )


def _build_push_table(orders, prototypes):
    rows = [
        [
            str(index),
            _source_table_for_task(order, task_index),
            task.dispatch_description,
            order.update_cycle,
        ]
        for index, (order, task_index, task) in enumerate(_all_order_tasks(orders), start=1)
    ]
    return clone_table_with_data(
        prototypes.push_table,
        ["序号", "调度名", "调度中文名", "推送频率"],
        rows,
    )


def _build_push_elements(orders, prototypes):
    elements = []
    replaced = False
    for element in prototypes.push_elements:
        if not replaced and element is prototypes.push_table:
            elements.append(_build_push_table(orders, prototypes))
            replaced = True
            continue
        elements.append(clone_element(element))
    if not replaced:
        elements.append(_build_push_table(orders, prototypes))
    return elements


def build_table_landing_design_document(excel_path, service_dir, template_path, output_path):
    orders = read_table_landing_work_orders(excel_path, service_dir)
    with tempfile.TemporaryDirectory(prefix="document_filler_table_landing_design_") as temp_dir:
        converted_template = _convert_legacy_doc_template(template_path, Path(temp_dir) / "template")
        document = Document(converted_template)
        logic_heading = find_heading(document, "Heading 1", "业务逻辑说明")
        detail_heading = find_heading(document, "Heading 1", "库表说明")
        push_heading = find_heading(document, "Heading 1", "推送方式")
        prototypes = _capture_template_prototypes(logic_heading, detail_heading, push_heading)

        replace_between(
            logic_heading._p,
            detail_heading._p,
            [_build_logic_table(orders, prototypes)]
            + _build_target_volume_elements(document, orders, prototypes, Path(temp_dir)),
        )
        replace_between(detail_heading._p, push_heading._p, [_build_detail_table(orders, prototypes)])
        replace_after(push_heading._p, _build_push_elements(orders, prototypes))

        return _save_document(document, output_path, Path(temp_dir) / "output")
