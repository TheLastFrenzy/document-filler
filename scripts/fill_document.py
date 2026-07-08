#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验收文档自动填充脚本 - 支持多种材料类型
用法:
  01-需求文档:
    python fill_document.py --service-dir "N08-数据报表服务" --material-type "01-数据报表_需求文档" --excel "台账清单.xlsx" --template "模板.docx" --output "输出.docx"
  02-设计文档:
    python fill_document.py --service-dir "N08-数据报表服务" --material-type "02-数据报表_设计文档" --excel "台账清单.xlsx" --template "模板.docx" --catalog "数据目录数据.xlsx" --output "输出.docx"
"""

import argparse, re, sys, os, subprocess, io, zipfile, tempfile, importlib.util, html, json, ntpath
from pathlib import Path

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        if hasattr(sys.stdout, "buffer"):
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# ── Dependencies ──
def ensure_module(name, pip_name=None):
    package_name = pip_name or name
    if importlib.util.find_spec(name) is None:
        raise ImportError(
            f"缺少 Python 依赖: {package_name}。请先在 document-filler 目录运行: "
            f"{sys.executable} -m pip install -r requirements.txt"
        )

ensure_module("openpyxl")
ensure_module("docx", "python-docx")
ensure_module("olefile")
ensure_module("PIL", "pillow")
ensure_module("fitz", "pymupdf")

import openpyxl
import olefile
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import struct

HEADER_BG = "F1F1F1"

MATERIAL_OUTPUT_EXTENSIONS = {
    "01-数据报表_需求文档": ".docx",
    "01-数据统计分析_需求文档": ".docx",
    "02-数据报表_设计文档": ".docx",
    "02-数据统计分析_设计文档": ".docx",
    "03-数据报表_上线记录": ".docx",
    "03-数据统计分析_测试文档": ".docx",
    "04-数据统计分析_结果表及使用说明": ".xlsx",
}


def resolve_output_path(output_path, material_type):
    output = Path(output_path)
    if output.exists() and output.is_dir():
        extension = MATERIAL_OUTPUT_EXTENSIONS.get(material_type, "")
        return str(output / f"{material_type}{extension}")
    return str(output)

# ══════════════════════════════════════════════════════════════
# Shared Helpers
# ══════════════════════════════════════════════════════════════

_STATS_RESULT_BUILDER = None

LEDGER_COLUMN_ALIASES = {
    "工单内容": ("工单内容", "工单标题"),
    "报表统计次数": ("报表统计次数", "程序数"),
    "业务说明": ("业务说明", "工单描述"),
    "业务描述": ("业务描述", "工单描述", "业务说明"),
    "统计分析结果表清单": ("统计分析结果表清单", "结果表清单"),
}


def ledger_aliases(canonical_name):
    return LEDGER_COLUMN_ALIASES.get(canonical_name, (canonical_name,))


def ledger_header_index(headers, canonical_name):
    for alias in ledger_aliases(canonical_name):
        if alias in headers:
            return headers.index(alias)
    return -1


def normalize_ledger_record(record):
    normalized = dict(record)
    for canonical_name, aliases in LEDGER_COLUMN_ALIASES.items():
        if str(normalized.get(canonical_name, "") or "").strip():
            continue
        for alias in aliases:
            value = str(normalized.get(alias, "") or "").strip()
            if value:
                normalized[canonical_name] = value
                break
        normalized.setdefault(canonical_name, "")
    return normalized


def read_excel(excel_path, service_dir):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    rows = []
    for r in range(2, ws.max_row + 1):
        rd = {}
        for c in range(1, ws.max_column + 1):
            v = ws.cell(row=r, column=c).value
            rd[headers[c - 1]] = str(v).strip() if v is not None else ""
        rd = normalize_ledger_record(rd)
        if rd.get("服务目录") == service_dir:
            rd["_row"] = r
            rd["_vml_row"] = r - 1
            rd["_vml_col"] = headers.index("业务逻辑") if "业务逻辑" in headers else -1
            rows.append(rd)
    if not rows:
        raise ValueError(f"未找到匹配数据: 服务目录={service_dir}")
    return rows


def excel_column_numbers(excel_path, header_names):
    try:
        wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
        wanted = set(header_names)
        wb.close()
        return [idx + 1 for idx, header in enumerate(headers) if header in wanted]
    except Exception:
        return []


def load_stats_result_builder():
    global _STATS_RESULT_BUILDER
    if _STATS_RESULT_BUILDER is not None:
        return _STATS_RESULT_BUILDER
    script_path = Path(__file__).resolve().parent / "build_stats_result_usage_workbook.py"
    spec = importlib.util.spec_from_file_location("_document_filler_stats_result_builder", script_path)
    module = importlib.util.module_from_spec(spec)
    if spec.loader is None:
        raise ImportError(f"无法加载统计分析结果表解析脚本: {script_path}")
    spec.loader.exec_module(module)
    _STATS_RESULT_BUILDER = module
    return module


def add_solid_borders(tblPr):
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def xml_safe(t):
    if not t: return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', t)


STATS_LOGIC_TERM_REPLACEMENTS = (
    ("数据筛选清洗", "数据范围确认"),
    ("筛选清洗", "范围确认"),
    ("清洗加工后的数据", "整理后的数据"),
    ("清洗加工后", "整理后"),
    ("清洗加工", "整理"),
    ("标准化抽取", "按字段口径整理"),
    ("数据抽取", "数据取用"),
    ("抽取", "取用"),
    ("质量检查", "口径核对"),
    ("清洗", "整理"),
    ("加密", ""),
)


def sanitize_stats_logic_text(value):
    text = str(value or "")
    for old, new in STATS_LOGIC_TERM_REPLACEMENTS:
        text = text.replace(old, new)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"（\s*）", "", text)
    text = re.sub(r"：\s*，", "：", text)
    text = re.sub(r"，{2,}", "，", text)
    return text.strip()


def mp(text, style_id, indent=None, word_wrap=False):
    p = OxmlElement("w:p")
    pr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pr.append(ps)
    if indent:
        i = OxmlElement("w:ind")
        i.set(qn("w:firstLine"), str(indent))
        pr.append(i)
    if word_wrap:
        ww = OxmlElement("w:wordWrap")
        ww.set(qn("w:val"), "1")
        pr.append(ww)
    p.append(pr)
    r = OxmlElement("w:r")
    tt = OxmlElement("w:t")
    tt.text = text
    tt.set(qn("xml:space"), "preserve")
    r.append(tt)
    p.append(r)
    return p


DATA_REPORT_CATALOG_COL = "02-数据报表_设计文档-数据来源库表清单对应数据目录代码"
DATA_REPORT_TEMPLATE_FRAGMENTS = (
    "围绕业务说明中列明的统计口径和数据目录",
    "报表需覆盖主要统计对象、关键字段、数据量、更新时间、空值情况和样例数据等内容",
    "需按工单主题和报表内容命名",
    "材料应包含统计结果、字段说明和必要的口径说明",
    "统计范围依据交付物附件整理",
)


def compact_spaces(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def unique_list(values, limit=None):
    result = []
    for value in values:
        value = str(value or "").strip()
        if value and value not in result:
            result.append(value)
            if limit and len(result) >= limit:
                break
    return result


def final_cn_sentence(text):
    parts = [part.strip() for part in re.split(r"(?<=[。；;])\s*", str(text or "")) if part.strip()]
    return parts[-1] if parts else ""


def text_variant_index(row, catalog_codes, size, salt=""):
    seed = "|".join(
        [
            str(row.get("工单号", "")),
            str(row.get("工单内容", "")),
            str(row.get("业务说明", ""))[:80],
            "、".join(unique_list(catalog_codes, limit=5)),
            salt,
        ]
    )
    return sum(ord(ch) for ch in seed) % size if size else 0


def extract_data_report_codes(text, include_plain_numbers=True, limit=None):
    patterns = [r"[A-Z0-9]{5,}/\d{6}", r"\b[A-Z]{1,4}\d{3,}/\d{6}\b"]
    if include_plain_numbers:
        patterns.append(r"(?<![A-Za-z0-9])\d{15}(?![A-Za-z0-9])")
    values = []
    for pattern in patterns:
        values.extend(re.findall(pattern, str(text or "")))
    return unique_list(values, limit)


def join_catalog_codes(codes, fallback="相关目录", limit=3):
    values = unique_list(codes)
    selected = values[:limit]
    if not selected:
        return fallback
    suffix = "等目录" if len(values) > len(selected) else "目录"
    return "、".join(selected) + suffix


def infer_data_report_subject(row):
    for key in ("工单内容", "业务说明", "需求描述"):
        value = compact_spaces(row.get(key, ""))
        if value:
            sentence = re.split(r"(?<=[。；;])\s*", value)[0]
            sentence = re.sub(r"^根据[^，。；;]{0,40}[，,]", "", sentence)
            return sentence[:90]
    return "本次数据报表需求"


def report_names_from_business_text(row, limit=3):
    biz = str(row.get("业务说明", "") or "")
    match = re.search(r"本次工作拟产出(?:以下)?\d+份报表成果", biz)
    if match:
        names = parse_report_names(biz[match.end():])
        if names:
            return display_name_list(names, "相关报表成果", limit=limit)
    attachment_names = collect_attachment_report_names(row.get("_attachment_names") or [])
    if attachment_names:
        return display_name_list(attachment_names, "相关报表成果", limit=limit)
    labels = []
    for program in row.get("_programs") or []:
        label = compact_spaces(program.get("program_cn", ""))
        if label and label not in labels:
            labels.append(label)
    if labels:
        return display_name_list(labels, "相关报表成果", limit=limit)
    parsed_names = [cn or en for cn, en in parse_report_list(row.get("统计分析结果表清单", ""))]
    if parsed_names:
        return display_name_list(parsed_names, "相关报表成果", limit=limit)
    return "报表成果"


def deliverable_report_names_from_row(row, fallback="本次交付材料", limit=3):
    names = collect_attachment_report_names(row.get("_attachment_names") or [])
    if names:
        return display_name_list(names, fallback, limit=limit)
    raw_values = []
    for key in ("交付物", "交付附件"):
        value = row.get(key)
        if value:
            raw_values.extend(re.split(r"[\n；;]+", str(value)))
    names = collect_attachment_report_names(raw_values)
    if names:
        return display_name_list(names, fallback, limit=limit)
    return fallback


DELIVERY_SUFFIX_CATEGORY = {
    ".xlsx": "excel",
    ".xlsm": "excel",
    ".xls": "excel",
    ".csv": "excel",
    ".docx": "word",
    ".doc": "word",
    ".pdf": "pdf",
    ".zip": "archive",
    ".rar": "archive",
    ".7z": "archive",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".txt": "text",
}


def _delivery_attachment_values(row):
    values = []
    for key in ("_attachment_files", "_attachment_names"):
        raw = row.get(key)
        if isinstance(raw, (list, tuple, set)):
            values.extend(str(item) for item in raw if item)
        elif raw:
            values.append(str(raw))
    for key in ("交付物", "交付附件"):
        raw = row.get(key)
        if raw:
            values.append(str(raw))
    return values


def _is_generated_internal_archive(value, suffix):
    if suffix not in {".zip", ".rar", ".7z"}:
        return False
    stem = Path(str(value or "")).stem
    generated_stems = {"Ole10Native_embedded", "Workbook_embedded", "Workbook", "package"}
    return stem in generated_stems or bool(re.fullmatch(r"deliverable_row\d+", stem, flags=re.I))


def delivery_attachment_suffixes(row):
    suffixes = []
    suffix_pattern = re.compile(r"\.(xlsx|xlsm|xls|csv|docx|doc|pdf|zip|rar|7z|png|jpe?g|txt)\b", re.I)
    for value in _delivery_attachment_values(row):
        suffix = Path(value).suffix.lower()
        if suffix in DELIVERY_SUFFIX_CATEGORY and suffix not in suffixes and not _is_generated_internal_archive(value, suffix):
            suffixes.append(suffix)
        for match in suffix_pattern.finditer(value):
            suffix = "." + match.group(1).lower()
            if suffix in DELIVERY_SUFFIX_CATEGORY and suffix not in suffixes and not _is_generated_internal_archive(value, suffix):
                suffixes.append(suffix)
    return suffixes


def delivery_format_categories(row):
    categories = []
    order = ["excel", "word", "pdf", "archive", "image", "text"]
    for suffix in delivery_attachment_suffixes(row):
        category = DELIVERY_SUFFIX_CATEGORY.get(suffix)
        if category and category not in categories:
            categories.append(category)
    return [category for category in order if category in categories]


def delivery_naming_sentence(row):
    subject_text = "".join(str(row.get(key, "") or "") for key in ("工单内容", "业务说明", "交付物"))
    if "接口" in subject_text:
        naming_basis = "工单编号、接口编号、统计时间段和报表主题"
    else:
        naming_basis = "工单编号、统计时间段和报表主题"
    return f"文件命名应对应{naming_basis}，保留可追溯线索。"


def delivery_format_sentence(row):
    categories = delivery_format_categories(row)
    naming = delivery_naming_sentence(row)
    if not categories:
        return f"交付材料按实际附件内容整理，{naming}"
    phrases = []
    if "excel" in categories:
        phrases.append("Excel电子表格报表")
    if "word" in categories:
        phrases.append("Word文档说明材料")
    if "pdf" in categories:
        phrases.append("PDF版定稿或签收材料")
    if "archive" in categories:
        phrases.append("保留原目录结构的压缩包附件")
    if "image" in categories:
        phrases.append("图片类截图或佐证材料")
    if "text" in categories:
        phrases.append("文本类补充清单")
    return f"交付成果包括{'、'.join(phrases)}，{naming}"


def catalog_resource_labels(catalog_codes, catalog_context=None, limit=3):
    labels = []
    catalog_context = catalog_context or {}
    for code in unique_list(catalog_codes):
        info = catalog_context.get(code, {})
        name = compact_spaces(info.get("资源名称", ""))
        label = f"{name}（{code}）" if name else code
        if label and label not in labels:
            labels.append(label)
        if len(labels) >= limit:
            break
    if not labels:
        return "相关数据目录"
    suffix = "等目录" if len(unique_list(catalog_codes)) > len(labels) else "目录"
    return "、".join(labels) + suffix


def catalog_field_labels(catalog_codes, catalog_context=None, limit=6):
    fields = []
    catalog_context = catalog_context or {}
    for code in unique_list(catalog_codes):
        for field in catalog_context.get(code, {}).get("字段", []):
            field = compact_spaces(field)
            if field and field not in fields:
                fields.append(field)
            if len(fields) >= limit:
                return fields
    return fields


def program_field_labels(row, limit=5):
    fields = []
    for program in row.get("_programs") or []:
        field_comments = program.get("field_comments")
        if field_comments is None and program.get("xml"):
            try:
                field_comments = extract_indicator_field_comments_from_program(program)
            except Exception:
                field_comments = []
        for field in field_comments or []:
            field = compact_spaces(field)
            if field and field not in fields:
                fields.append(field)
            if len(fields) >= limit:
                return fields
    return fields


def data_report_requirement_field_hint(row, catalog_codes, catalog_context=None, limit=6):
    fields = []
    for field in catalog_field_labels(catalog_codes, catalog_context, limit=limit):
        if field not in fields:
            fields.append(field)
    for field in program_field_labels(row, limit=limit):
        if field not in fields:
            fields.append(field)
        if len(fields) >= limit:
            break
    if fields:
        return "、".join(fields[:limit])
    return "目录代码、资源名称、统计结果"


def build_data_report_catalog_context(catalog_path, data_rows):
    if not catalog_path:
        return {}
    if not os.path.exists(catalog_path):
        raise FileNotFoundError(f"数据目录数据文件不存在: {catalog_path}")
    all_codes = []
    for row in data_rows:
        for code in extract_data_report_codes(row.get(DATA_REPORT_CATALOG_COL, ""), include_plain_numbers=False):
            if code not in all_codes:
                all_codes.append(code)
    if not all_codes:
        return {}
    rmap, fmap = load_catalog_data(catalog_path, all_codes)
    context = {}
    for code in all_codes:
        info = rmap.get(code, {})
        fields = []
        for item in fmap.get(code, []):
            name = compact_spaces(item.get("数据项名称", ""))
            if name and name not in fields:
                fields.append(name)
        context[code] = {
            "资源名称": compact_spaces(info.get("资源名称", "")),
            "资源编码": compact_spaces(info.get("资源编码", "")),
            "字段": fields,
        }
    return context


def _ensure_partial_catalog_mentions_use_etc(data_req, catalog_codes):
    text = str(data_req or "")
    if not text or not catalog_codes:
        return text
    mentioned = [code for code in catalog_codes if code in text]
    if not mentioned or len(mentioned) >= len(catalog_codes):
        return text
    mention_start = min(text.find(code) for code in mentioned if code in text)
    mention_end = max(text.find(code) + len(code) for code in mentioned if code in text)
    nearby = text[mention_start : min(len(text), mention_end + 8)]
    if "等" in nearby:
        return text
    code_pattern = r"\s*[、,，]\s*".join(re.escape(code) for code in mentioned)
    pattern = re.compile(f"({code_pattern})(?:\\s*目录)?")
    return pattern.sub(lambda match: match.group(1) + "等目录", text, count=1)


def _clean_delivery_text(delivery, catalog_codes):
    text = re.sub(r"[【\[].*?[】\]]", "", str(delivery or ""))
    if catalog_codes:
        for code in extract_data_report_codes(text, include_plain_numbers=True):
            if code not in catalog_codes:
                text = text.replace(code, "")
    replacements = [
        ("验收时抽查", "验收时以"),
        ("抽查", "确认"),
        ("复核", "对照"),
        ("检查", "确认"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"(及|和|与|、)+并", "并", text)
    text = re.sub(r"(及|和|与|、)+作为", "作为", text)
    text = re.sub(r"，+", "，", text).strip("，。；;、 ")
    return text + "。" if text and not text.endswith(("。", "；", ";")) else text


def infer_data_report_requirement(row, catalog_codes, catalog_context=None):
    subject = infer_data_report_subject(row)
    resource_hint = catalog_resource_labels(catalog_codes, catalog_context, limit=3)
    field_hint = data_report_requirement_field_hint(row, catalog_codes, catalog_context, limit=6)
    report_hint = report_names_from_business_text(row, limit=3)
    endings = [
        "涉及多张目录的，报表中应保留目录代码和资源名称，后续核对口径时能直接追到来源。",
        "同一报表跨目录取数时，应写清目录来源和字段口径，避免验收时只看到汇总值。",
        "统计周期、筛选范围和更新时间要能在材料中找到依据，便于业务方复核。",
        "对口径调整或字段缺失的情况，文档中应写明处理方式，避免后续重复确认。",
    ]
    ending = endings[text_variant_index(row, catalog_codes, len(endings), "requirement")]
    return (
        f"本次数据需求基于{resource_hint}，围绕{subject}业务场景，梳理{field_hint}等关键字段，"
        f"明确统计范围和结果口径，形成{report_hint}。数据内容应覆盖资源名称、目录代码、字段列、"
        f"统计结果和更新时间等可核验信息。{ending}"
    )


def infer_data_report_delivery(row, catalog_codes, catalog_context=None):
    resource_hint = catalog_resource_labels(catalog_codes, catalog_context, limit=3)
    field_hint = data_report_requirement_field_hint(row, catalog_codes, catalog_context, limit=6)
    subject = infer_data_report_subject(row)
    report_hint = report_names_from_business_text(row, limit=3)
    program_hint = data_report_program_labels(row, limit=3)
    endings = [
        "文件归档时按工单和报表主题归类，后续查找口径时不需要再反查原始附件。",
        "材料提交后应能直接对应到工单、目录和报表成果，便于验收归档。",
        "业务方拿到材料后可按统计周期继续更新，不需要重新整理目录来源。",
        "涉及接口口径调整的，文件名和说明中应保留版本线索，便于后续管理。",
    ]
    ending = endings[text_variant_index(row, catalog_codes, len(endings), "delivery")]
    return (
        f"交付材料包括{report_hint}、涉及数据目录清单、字段口径说明和必要附件。{delivery_format_sentence(row)}"
        f"验收时对照{resource_hint}中的{field_hint}等字段、{program_hint}处理结果与{subject}业务说明，确认统计范围、"
        f"字段口径、结果呈现和附件内容一致。{ending}"
    )


def ensure_cn_period(text):
    text = compact_spaces(text)
    if not text:
        return ""
    return text if text.endswith(("。", "；", ";")) else text + "。"


def first_sentence_from_fields(row, field_names, fallback="", limit=120):
    for field_name in field_names:
        value = compact_spaces(row.get(field_name, ""))
        if not value:
            continue
        sentence = re.split(r"(?<=[。；;])\s*", value)[0]
        return sentence[:limit]
    return fallback


def clean_design_basis_sentence(text):
    text = compact_spaces(text)
    if not text:
        return ""
    list_like_patterns = [r"包括但不限于\s*\d", r"报表链接", r'"系', r"下发部门确认是否保留"]
    if any(re.search(pattern, text) for pattern in list_like_patterns) or text.count("、") >= 8:
        return ""
    text = re.sub(r"^根据附件[，,]?", "", text)
    replacements = [
        ("切实", ""),
        ("精准、全面的", ""),
        ("精准、全面", ""),
        ("充分展现", "反映"),
        ("现需", "需"),
        ("并输出一次性报表", "形成一次性报表"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"为会议提供的数据支撑", "为会议提供数据支撑", text)
    text = re.sub(r"\s+", "", text).strip("，,；; ")
    return ensure_cn_period(text) if text else ""


def data_report_business_focus(row):
    text = "".join(
        compact_spaces(row.get(key, ""))
        for key in ("工单内容", "业务说明", "业务描述", "统计分析结果表清单")
    )
    focus_rules = [
        ("工作汇报", "阶段性工作汇报材料整理"),
        ("便捷共享", "便捷共享目录统计和会议材料准备"),
        ("涉企", "涉企数据资源高价值目录梳理"),
        ("婚介码", "婚介码服务目录信息核对"),
        ("社会救助", "社会救助事项办理和人员信息比对"),
        ("残疾人", "残疾人助学、个体工商户事项办理核对"),
        ("特殊困难老人", "特殊困难老人关心关爱名单排查"),
        ("低收入人口", "低收入人口动态监测"),
        ("养老迁入迁出", "养老服务对象迁入迁出比对"),
        ("下发量", "目录下发量口径调整和统计核验"),
    ]
    for keyword, focus in focus_rules:
        if keyword in text:
            return focus
    return infer_data_report_subject(row)


def display_name_list(values, fallback, limit=3):
    names = []
    for value in values or []:
        name = compact_spaces(Path(str(value)).stem if "." in str(value) else value)
        if name and name not in names:
            names.append(name)
            if len(names) >= limit:
                break
    if not names:
        return fallback
    suffix = "等" if len(unique_list(values)) > len(names) else ""
    return "、".join(names) + suffix


def data_report_program_labels(row, limit=3):
    labels = []
    for program in row.get("_programs") or []:
        cn = compact_spaces(program.get("program_cn", ""))
        en = compact_spaces(program.get("program_en", ""))
        if cn and en:
            label = f"{cn}（{en}）"
        else:
            label = cn or en
        if label and label not in labels:
            labels.append(label)
    if not labels:
        for cn, en in parse_report_list(row.get("统计分析结果表清单", "")):
            label = f"{cn}（{en}）" if cn and en else cn or en
            if label and label not in labels:
                labels.append(label)
    return display_name_list(labels, "相关报表程序", limit=limit)


def data_report_field_comment_hint(row, limit=4):
    comments = []
    for program in row.get("_programs") or []:
        field_comments = program.get("field_comments")
        if field_comments is None and program.get("xml"):
            try:
                field_comments = extract_indicator_field_comments_from_program(program)
            except Exception:
                field_comments = []
        for comment in field_comments or []:
            text = compact_spaces(comment)
            if text and text not in comments:
                comments.append(text)
                if len(comments) >= limit:
                    break
        if len(comments) >= limit:
            break
    if not comments:
        return ""
    suffix = "等字段" if len(comments) >= limit else "字段"
    return f"，重点字段包括{'、'.join(comments)}{suffix}"


def infer_data_report_content_description(row, catalog_codes):
    subject = infer_data_report_subject(row)
    code_hint = join_catalog_codes(catalog_codes, "相关数据目录", limit=2)
    basis = clean_design_basis_sentence(first_sentence_from_fields(row, ("业务说明", "业务描述"), ""))
    if basis:
        return ensure_cn_period(f"{basis}本工单围绕{code_hint}梳理{subject}涉及的统计口径、字段范围和交付附件")
    return ensure_cn_period(f"本工单围绕{subject}和{code_hint}梳理报表统计口径、字段范围和交付附件")


def infer_data_report_business_scene(row, catalog_codes, catalog_context=None):
    subject = infer_data_report_subject(row)
    basis = clean_design_basis_sentence(first_sentence_from_fields(row, ("业务描述", "业务说明"), ""))
    focus = data_report_business_focus(row)
    resource_hint = catalog_resource_labels(catalog_codes, catalog_context, limit=2)
    variants = [
        f"业务侧需要一份口径清楚的统计清单，用于{focus}，并能和{resource_hint}对应起来。",
        f"{focus}需要把{resource_hint}中的关键数据放到同一套统计口径下，便于业务人员直接核对明细和汇总结果。",
        f"该报表面向{focus}，重点解决{subject}中统计对象分散、字段口径不易对齐的问题。",
        f"业务人员可直接看到{focus}涉及的对象范围、口径变化和结果差异，减少后续重复整理。",
    ]
    sentence = variants[text_variant_index(row, catalog_codes, len(variants), "design_scene")]
    if basis:
        return ensure_cn_period(basis + sentence)
    return ensure_cn_period(sentence)


def infer_data_report_content(row, catalog_codes, catalog_context=None):
    subject = infer_data_report_subject(row)
    resource_hint = catalog_resource_labels(catalog_codes, catalog_context, limit=3)
    field_hint = data_report_requirement_field_hint(row, catalog_codes, catalog_context, limit=6)
    report_hint = deliverable_report_names_from_row(row, fallback="本次交付报表", limit=2)
    focus = data_report_business_focus(row)
    variants = [
        f"数据内容以{resource_hint}为主要来源，保留{field_hint}等字段，并按{focus}口径整理为{report_hint}。",
        f"本工单取用{resource_hint}中的{field_hint}等信息，结果侧覆盖{report_hint}需要的明细、汇总值和更新时间。",
        f"数据范围覆盖{resource_hint}中的{field_hint}等信息，统计结果需能对应{subject}的业务口径和{report_hint}的表内字段。",
        f"设计时需把{resource_hint}、{field_hint}等字段和{focus}的统计口径放在同一条链路里，便于后续核对来源和结果。",
    ]
    return ensure_cn_period(variants[text_variant_index(row, catalog_codes, len(variants), "design_content")])


def infer_data_report_result_form(row):
    report_names = deliverable_report_names_from_row(row, fallback="本次交付材料", limit=3)
    return ensure_cn_period(f"最终形成{report_names}，随文档保留统计结果、字段口径、数据来源清单和必要截图材料")


def infer_data_report_processing_logic(row, catalog_codes):
    subject = infer_data_report_subject(row)
    code_hint = join_catalog_codes(catalog_codes, "相关数据目录", limit=2)
    program_hint = data_report_program_labels(row)
    field_hint = data_report_field_comment_hint(row)
    return ensure_cn_period(sanitize_stats_logic_text(
        f"根据{subject}的统计口径，读取{code_hint}并通过{program_hint}完成数据抽取、字段整理、统计汇总和结果写入{field_hint}"
    ))


def _is_template_like(text):
    value = str(text or "")
    return not compact_spaces(value) or any(fragment in value for fragment in DATA_REPORT_TEMPLATE_FRAGMENTS)


def normalize_data_report_text_fields(row, catalog_context=None):
    normalized = dict(row)
    catalog_codes = extract_data_report_codes(
        normalized.get(DATA_REPORT_CATALOG_COL, ""),
        include_plain_numbers=False,
    )

    data_req = normalized.get("数据需求", "")
    if _is_template_like(data_req):
        data_req = infer_data_report_requirement(normalized, catalog_codes, catalog_context)
    data_req = _ensure_partial_catalog_mentions_use_etc(data_req, catalog_codes)
    normalized["数据需求"] = data_req

    delivery = normalized.get("交付要求", "")
    if _is_template_like(delivery):
        delivery = infer_data_report_delivery(normalized, catalog_codes, catalog_context)
    normalized["交付要求"] = _clean_delivery_text(delivery, catalog_codes)

    data_content = normalized.get("数据内容", "")
    if _is_template_like(data_content):
        normalized["数据内容"] = infer_data_report_content(normalized, catalog_codes, catalog_context)

    content_desc = normalized.get("内容描述", "")
    if _is_template_like(content_desc):
        normalized["内容描述"] = infer_data_report_content_description(normalized, catalog_codes)

    business_scene = normalized.get("业务场景", "")
    if _is_template_like(business_scene):
        normalized["业务场景"] = infer_data_report_business_scene(normalized, catalog_codes, catalog_context)

    result_form = normalized.get("结果形式", "")
    if _is_template_like(result_form):
        normalized["结果形式"] = infer_data_report_result_form(normalized)

    processing_logic = normalized.get("数据处理逻辑", "")
    if _is_template_like(processing_logic):
        processing_logic = infer_data_report_processing_logic(normalized, catalog_codes)
    normalized["数据处理逻辑"] = sanitize_stats_logic_text(processing_logic)

    return normalized


def quote_ps_path(path):
    return str(path).replace("'", "''")


def crop_and_normalize_image(image_path, min_width=1400, min_height=520):
    image_path = Path(image_path)
    if not image_path.exists():
        return False
    try:
        image = Image.open(image_path)
    except Exception:
        return False
    if image.mode in {"RGBA", "LA"} or "transparency" in image.info:
        background = Image.new("RGBA", image.size, "white")
        background.alpha_composite(image.convert("RGBA"))
        image = background.convert("RGB")
    else:
        image = image.convert("RGB")
    mask = image.point(lambda value: 255 if value < 245 else 0)
    bbox = mask.getbbox()
    if bbox:
        left, top, right, bottom = bbox
        padding = 18
        image = image.crop((
            max(0, left - padding),
            max(0, top - padding),
            min(image.width, right + padding),
            min(image.height, bottom + padding),
        ))
    scale = max(min_width / image.width, min_height / image.height, 1)
    if scale > 1:
        image = image.resize((int(image.width * scale), int(image.height * scale)), Image.Resampling.LANCZOS)
    image.save(image_path)
    return image_path.exists() and image_path.stat().st_size > 0


def image_bytes_for_docx(payload):
    try:
        image = Image.open(io.BytesIO(payload))
        if image.mode in {"RGBA", "LA"} or "transparency" in image.info:
            background = Image.new("RGBA", image.size, "white")
            background.alpha_composite(image.convert("RGBA"))
            image = background.convert("RGB")
        else:
            image = image.convert("RGB")
        out = io.BytesIO()
        image.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def select_excel_preview_range(excel_path):
    wb = openpyxl.load_workbook(excel_path, data_only=True, read_only=True)
    try:
        best = None
        for ws in wb.worksheets:
            first_row = first_col = None
            last_row = last_col = 0
            non_empty = 0
            for r_idx, row in enumerate(ws.iter_rows(), start=1):
                for c_idx, cell in enumerate(row, start=1):
                    value = cell.value
                    if value is None or str(value).strip() == "":
                        continue
                    first_row = r_idx if first_row is None else min(first_row, r_idx)
                    first_col = c_idx if first_col is None else min(first_col, c_idx)
                    last_row = max(last_row, r_idx)
                    last_col = max(last_col, c_idx)
                    non_empty += 1
            if first_row is None:
                continue
            rows = max(1, last_row - first_row + 1)
            cols = max(1, last_col - first_col + 1)
            score = non_empty * 12 + min(rows, 60) * 2 + min(cols, 20)
            if best is None or score > best[0]:
                best = (score, ws.title, first_row, first_col, last_row, last_col)
        if best is None:
            return None
        _, sheet_name, first_row, first_col, last_row, last_col = best
        used_rows = max(1, last_row - first_row + 1)
        used_cols = max(1, last_col - first_col + 1)
        row_cap = min(used_rows, 32)
        if used_rows <= 8:
            col_cap = min(used_cols, 16)
        elif used_cols <= 8:
            col_cap = used_cols
        else:
            col_cap = min(used_cols, 9)
        return sheet_name, first_row, first_col, first_row + row_cap - 1, first_col + col_cap - 1
    finally:
        wb.close()


def excel_range_to_png(source, target):
    source = Path(source)
    target = Path(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    selected = select_excel_preview_range(source)
    if not selected:
        return False
    sheet_name, first_row, first_col, last_row, last_col = selected
    ps = f"""
$ErrorActionPreference = 'Stop'
$excel = $null
$wb = $null
$chartObj = $null
try {{
  $excel = New-Object -ComObject Excel.Application
  $excel.Visible = $false
  $excel.DisplayAlerts = $false
  $wb = $excel.Workbooks.Open('{quote_ps_path(source)}', 0, $true)
  $ws = $wb.Worksheets.Item('{quote_ps_path(sheet_name)}')
  $ws.Activate()
  if ($excel.ActiveWindow -ne $null) {{ $excel.ActiveWindow.Zoom = 180 }}
  $range = $ws.Range($ws.Cells({first_row}, {first_col}), $ws.Cells({last_row}, {last_col}))
  $range.CopyPicture(1, 2)
  Start-Sleep -Milliseconds 700
  $chartWidth = [Math]::Max(1000, [Math]::Min(2400, $range.Width * 2.4))
  $chartHeight = [Math]::Max(520, [Math]::Min(1600, $range.Height * 2.4))
  $chartObj = $ws.ChartObjects().Add($range.Left, $range.Top, $chartWidth, $chartHeight)
  $chart = $chartObj.Chart
  $chart.ChartArea.Format.Line.Visible = 0
  $chart.ChartArea.Format.Fill.Visible = -1
  $chart.ChartArea.Format.Fill.ForeColor.RGB = 16777215
  $chart.Paste()
  if ($chart.Shapes.Count -gt 0) {{
    $shape = $chart.Shapes.Item(1)
    $shape.LockAspectRatio = -1
    $scale = [Math]::Min(($chartWidth - 6) / $shape.Width, ($chartHeight - 6) / $shape.Height)
    if ($scale -gt 1) {{ $shape.Width = $shape.Width * $scale }}
    $shape.Left = 3
    $shape.Top = 3
  }}
  $chart.Export('{quote_ps_path(target)}', 'PNG')
}} finally {{
  if ($chartObj -ne $null) {{ $chartObj.Delete() }}
  if ($wb -ne $null) {{ $wb.Close($false) }}
  if ($excel -ne $null) {{ $excel.Quit() }}
}}
"""
    proc = subprocess.run(["powershell", "-NoProfile", "-Command", ps], capture_output=True, text=True, timeout=120)
    if proc.returncode != 0 or not target.exists() or target.stat().st_size == 0:
        return False
    return crop_and_normalize_image(target)


def excel_com_used_range_to_png(source, target):
    source = Path(source)
    target = Path(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    ps = f"""
$ErrorActionPreference = 'Stop'
$excel = $null
$wb = $null
$chartObj = $null
try {{
  $excel = New-Object -ComObject Excel.Application
  $excel.Visible = $false
  $excel.DisplayAlerts = $false
  $wb = $excel.Workbooks.Open('{quote_ps_path(source)}', 0, $true)
  $bestWs = $null
  $bestScore = -1
  foreach ($sheet in $wb.Worksheets) {{
    $used = $sheet.UsedRange
    if ($used -eq $null) {{ continue }}
    $rows = [int]$used.Rows.Count
    $cols = [int]$used.Columns.Count
    if ($rows -lt 1 -or $cols -lt 1) {{ continue }}
    $score = $rows * $cols
    if ($score -gt $bestScore) {{
      $bestScore = $score
      $bestWs = $sheet
    }}
  }}
  if ($bestWs -eq $null) {{ throw 'no used range' }}
  $bestWs.Activate()
  if ($excel.ActiveWindow -ne $null) {{ $excel.ActiveWindow.Zoom = 180 }}
  $used = $bestWs.UsedRange
  $firstRow = [int]$used.Row
  $firstCol = [int]$used.Column
  $rows = [int]$used.Rows.Count
  $cols = [int]$used.Columns.Count
  $rowCap = [Math]::Min($rows, 32)
  if ($rows -le 8) {{
    $colCap = [Math]::Min($cols, 16)
  }} elseif ($cols -le 8) {{
    $colCap = $cols
  }} else {{
    $colCap = [Math]::Min($cols, 9)
  }}
  $range = $bestWs.Range($bestWs.Cells($firstRow, $firstCol), $bestWs.Cells($firstRow + $rowCap - 1, $firstCol + $colCap - 1))
  $range.CopyPicture(1, 2)
  Start-Sleep -Milliseconds 700
  $chartWidth = [Math]::Max(1200, [Math]::Min(2600, $range.Width * 2.6))
  $chartHeight = [Math]::Max(560, [Math]::Min(1700, $range.Height * 2.6))
  $chartObj = $bestWs.ChartObjects().Add($range.Left, $range.Top, $chartWidth, $chartHeight)
  $chart = $chartObj.Chart
  $chart.ChartArea.Format.Line.Visible = 0
  $chart.ChartArea.Format.Fill.Visible = -1
  $chart.ChartArea.Format.Fill.ForeColor.RGB = 16777215
  $chart.Paste()
  if ($chart.Shapes.Count -gt 0) {{
    $shape = $chart.Shapes.Item(1)
    $shape.LockAspectRatio = -1
    $scale = [Math]::Min(($chartWidth - 6) / $shape.Width, ($chartHeight - 6) / $shape.Height)
    if ($scale -gt 1) {{ $shape.Width = $shape.Width * $scale }}
    $shape.Left = 3
    $shape.Top = 3
  }}
  $chart.Export('{quote_ps_path(target)}', 'PNG')
}} finally {{
  if ($chartObj -ne $null) {{ $chartObj.Delete() }}
  if ($wb -ne $null) {{ $wb.Close($false) }}
  if ($excel -ne $null) {{ $excel.Quit() }}
}}
"""
    try:
        proc = subprocess.run(["powershell", "-NoProfile", "-Command", ps], capture_output=True, text=True, timeout=120)
    except Exception:
        return False
    if proc.returncode != 0 or not target.exists() or target.stat().st_size == 0:
        return False
    return crop_and_normalize_image(target)


def office_export_to_pdf(source, target):
    source = Path(source)
    target = Path(target)
    suffix = source.suffix.lower()
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    if suffix in {".docx", ".doc"}:
        ps = f"""
$ErrorActionPreference = 'Stop'
$word = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open('{quote_ps_path(source)}')
  $doc.SaveAs([ref]'{quote_ps_path(target)}', [ref]17)
  $doc.Close($false)
}} finally {{
  if ($word -ne $null) {{ $word.Quit() }}
}}
"""
    elif suffix in {".xlsx", ".xls"}:
        ps = f"""
$ErrorActionPreference = 'Stop'
$excel = $null
try {{
  $excel = New-Object -ComObject Excel.Application
  $excel.Visible = $false
  $excel.DisplayAlerts = $false
  $wb = $excel.Workbooks.Open('{quote_ps_path(source)}')
  foreach ($ws in $wb.Worksheets) {{ $ws.PageSetup.Zoom = $false; $ws.PageSetup.FitToPagesWide = 1; $ws.PageSetup.FitToPagesTall = $false }}
  $wb.ExportAsFixedFormat(0, '{quote_ps_path(target)}')
  $wb.Close($false)
}} finally {{
  if ($excel -ne $null) {{ $excel.Quit() }}
}}
"""
    else:
        return False
    proc = subprocess.run(["powershell", "-NoProfile", "-Command", ps], capture_output=True, text=True, timeout=120)
    return office_export_created_output(proc.returncode, target)


def office_export_created_output(returncode, target):
    target = Path(target)
    if target.exists() and target.stat().st_size > 0:
        return True
    return returncode == 0


def render_pdf_first_page(pdf_path, image_path):
    import fitz

    pdf_path = Path(pdf_path)
    image_path = Path(image_path)
    image_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        pdf = fitz.open(str(pdf_path))
        if len(pdf) == 0:
            pdf.close()
            return False
        page = pdf[0]
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.8, 2.8), alpha=False)
        pixmap.save(str(image_path))
        pdf.close()
    except Exception:
        return False
    return crop_and_normalize_image(image_path)


def choose_screenshot_source(files):
    preferred = {".xlsx": 0, ".xls": 0, ".pdf": 1, ".docx": 2, ".doc": 2}
    candidates = [Path(path) for path in files if Path(path).suffix.lower() in preferred and Path(path).exists()]
    if not candidates:
        return None
    candidates.sort(key=lambda path: (preferred[path.suffix.lower()], -(path.stat().st_size if path.exists() else 0)))
    return candidates[0]


def generate_attachment_screenshot_bytes(files, work_dir, row_number):
    source = choose_screenshot_source(files)
    if not source:
        return None
    work_dir = Path(work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)
    image_path = work_dir / f"row{int(row_number):02d}_attachment_screenshot.png"
    suffix = source.suffix.lower()
    if suffix == ".xlsx" and excel_range_to_png(source, image_path):
        return image_path.read_bytes()
    if suffix in {".xlsx", ".xls"} and excel_com_used_range_to_png(source, image_path):
        return image_path.read_bytes()
    pdf_path = source if suffix == ".pdf" else work_dir / f"row{int(row_number):02d}_{safe_file_name(source.stem)}.pdf"
    if suffix != ".pdf" and not office_export_to_pdf(source, pdf_path):
        return None
    if pdf_path.exists() and render_pdf_first_page(pdf_path, image_path):
        return image_path.read_bytes()
    return None


def safe_file_name(value):
    value = re.sub(r'[\x00-\x1f<>:"/\\|?*]+', "_", str(value or ""))
    return value.strip(" ._") or "attachment"


def _payload_output_name(base, suffix):
    base = safe_file_name(base)
    current_suffix = Path(base).suffix.lower()
    if current_suffix:
        return base
    return f"{base}{suffix}"


def _valid_deliverable_file_name(value):
    name = ntpath.basename(str(value or "").replace("\\", "/")).strip()
    suffix = Path(name).suffix.lower()
    if not name or suffix not in DELIVERY_SUFFIX_CATEGORY:
        return ""
    if any(ch in name for ch in "\x00\r\n\t"):
        return ""
    return name


def _decode_ole_native_string(raw):
    for encoding in ("gbk", "utf-8", "latin1"):
        try:
            text = raw.decode(encoding).strip("\x00 ")
        except Exception:
            continue
        if text:
            return text
    return ""


def extract_ole10native_filename(payload):
    candidates = []
    for match in re.finditer(rb"[\x20-\xff]{4,}\x00", payload[:4096]):
        text = _decode_ole_native_string(match.group(0).rstrip(b"\x00"))
        name = _valid_deliverable_file_name(text)
        if name and name not in candidates:
            candidates.append(name)
    return candidates[0] if candidates else ""


def _office_payload_suggested_name(payload):
    suffix = zip_office_suffix(payload)
    if suffix != ".xlsx":
        return ""
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            workbook_xml = zf.read("xl/workbook.xml").decode("utf-8", errors="replace")
    except Exception:
        return ""
    generic = {"sheet", "sheet1", "sheet2", "sheet3", "工作表", "工作表1", "工作表2"}
    for raw_name in re.findall(r'<sheet[^>]*name="([^"]+)"', workbook_xml):
        name = html.unescape(raw_name).strip()
        if name and name.lower() not in generic:
            return name
    return ""


def zip_office_suffix(payload):
    suffix = ".zip"
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            names = zf.namelist()
        if "[Content_Types].xml" in names and any(name.startswith("word/") for name in names):
            suffix = ".docx"
        elif "[Content_Types].xml" in names and any(name.startswith("xl/") for name in names):
            suffix = ".xlsx"
        elif "[Content_Types].xml" in names and any(name.startswith("ppt/") for name in names):
            suffix = ".pptx"
    except Exception:
        suffix = ".zip"
    return suffix


def write_detected_payload(row_dir, base, payload, files):
    if payload.startswith(b"PK\x03\x04"):
        suffix = zip_office_suffix(payload)
        target = row_dir / _payload_output_name(base, suffix)
        target.write_bytes(payload)
        files.append(target)
    elif payload.startswith(b"%PDF"):
        target = row_dir / _payload_output_name(base, ".pdf")
        target.write_bytes(payload)
        files.append(target)


def extract_ole_native_payload(payload):
    for signature in (b"PK\x03\x04", b"%PDF"):
        index = payload.find(signature)
        if index >= 0:
            return payload[index:]
    return None


def expand_zip_payloads(path, row_dir, files):
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                suffix = Path(name).suffix.lower()
                if suffix not in {".docx", ".xlsx", ".xls", ".pdf"}:
                    continue
                target = row_dir / "expanded" / safe_file_name(name)
                target.parent.mkdir(parents=True, exist_ok=True)
                target.write_bytes(zf.read(name))
                files.append(target)
    except Exception:
        return


def _sheet_rel_target(target):
    if target.startswith("../"):
        return "xl/" + target[3:]
    if target.startswith("/"):
        return target.lstrip("/")
    if target.startswith("xl/"):
        return target
    return "xl/worksheets/" + target


def extract_deliverable_attachments(excel_path, row_numbers, extract_dir, col_numbers=None):
    row_numbers = {int(row) for row in row_numbers}
    col_numbers = {int(col) for col in col_numbers} if col_numbers else None
    extract_dir = Path(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)
    result = {row: [] for row in row_numbers}
    try:
        with zipfile.ZipFile(excel_path, "r") as zf:
            names = set(zf.namelist())
            required = {"xl/worksheets/sheet1.xml", "xl/worksheets/_rels/sheet1.xml.rels", "xl/drawings/vmlDrawing1.vml"}
            if not required.issubset(names):
                return result
            sheet = zf.read("xl/worksheets/sheet1.xml").decode("utf-8", errors="replace")
            rels = zf.read("xl/worksheets/_rels/sheet1.xml.rels").decode("utf-8", errors="replace")
            vml = zf.read("xl/drawings/vmlDrawing1.vml").decode("utf-8", errors="replace")
            rel_map = {
                match.group(1): _sheet_rel_target(match.group(3))
                for match in re.finditer(r'<Relationship[^>]*Id="(rId\d+)"[^>]*Type="([^"]+)"[^>]*Target="([^"]+)"', rels)
            }
            anchors = {}
            for match in re.finditer(r"<v:shape\b(.*?)</v:shape>", vml, re.DOTALL):
                body = match.group(1)
                id_match = re.search(r'id="([^"]+)"', body)
                anchor_match = re.search(r"<x:Anchor>([^<]+)</x:Anchor>", body)
                if not id_match or not anchor_match:
                    continue
                parts = [int(part.strip()) for part in anchor_match.group(1).split(",")]
                anchors[id_match.group(1)] = (parts[2] + 1, parts[0] + 1)
            seen = set()
            for match in re.finditer(r'<oleObject[^>]*shapeId="(\d+)"[^>]*r:id="(rId\d+)"', sheet):
                shape_id, rid = match.groups()
                if (shape_id, rid) in seen or rid not in rel_map:
                    continue
                seen.add((shape_id, rid))
                row = col = None
                for shape, anchor in anchors.items():
                    if shape.endswith("_s" + shape_id):
                        row, col = anchor
                        break
                if row not in row_numbers or (col_numbers and col not in col_numbers):
                    continue
                target = rel_map[rid]
                try:
                    data = zf.read(target)
                except Exception:
                    continue
                row_dir = extract_dir / f"row{row:02d}"
                row_dir.mkdir(parents=True, exist_ok=True)
                raw_path = row_dir / Path(target).name
                raw_path.write_bytes(data)
                files = result.setdefault(row, [])
                if data.startswith(b"PK\x03\x04"):
                    suffix = zip_office_suffix(data)
                    suggested_name = _office_payload_suggested_name(data) or f"deliverable_row{row:02d}"
                    final = row_dir / _payload_output_name(suggested_name, suffix)
                    raw_path.replace(final)
                    files.append(final)
                    if suffix == ".zip":
                        expand_zip_payloads(final, row_dir, files)
                elif data.startswith(bytes.fromhex("d0cf11e0a1b11ae1")):
                    try:
                        ole = olefile.OleFileIO(str(raw_path))
                        for stream in ole.listdir():
                            try:
                                payload = ole.openstream(stream).read()
                            except Exception:
                                continue
                            stream_name = safe_file_name("_".join(stream))
                            native_name = extract_ole10native_filename(payload) if "Ole10Native" in stream_name else ""
                            native_payload = extract_ole_native_payload(payload) if native_name else None
                            if native_name and native_payload:
                                write_detected_payload(row_dir, native_name, native_payload, files)
                                continue
                            if stream_name.lower() == "workbook":
                                workbook_path = row_dir / f"{stream_name}.xls"
                                workbook_path.write_bytes(data)
                                files.append(workbook_path)
                            write_detected_payload(row_dir, stream_name, payload, files)
                            nested = extract_ole_native_payload(payload)
                            if nested and nested != payload:
                                suggested_name = native_name or _office_payload_suggested_name(nested) or f"{stream_name}_embedded"
                                write_detected_payload(row_dir, suggested_name, nested, files)
                        ole.close()
                    except Exception:
                        pass
                    for zip_path in list(row_dir.glob("*.zip")):
                        expand_zip_payloads(zip_path, row_dir, files)
    except Exception:
        return result
    return result


def build_attachment_previews(excel_path, row_numbers):
    with tempfile.TemporaryDirectory(prefix="document_filler_attachments_") as temp_dir:
        temp = Path(temp_dir)
        attachments = extract_deliverable_attachments(
            excel_path,
            row_numbers,
            temp / "attachments",
            excel_column_numbers(excel_path, ["交付物"]) or None,
        )
        previews = {}
        for row, files in attachments.items():
            payload = generate_attachment_screenshot_bytes(files, temp / "previews", row)
            if payload:
                previews[row] = payload
        return previews


def attach_delivery_files_for_requirement(excel_path, data_rows):
    if not data_rows or not os.path.exists(excel_path):
        return data_rows
    row_nums = set()
    for rd in data_rows:
        for row in rd.get("_row_numbers") or [rd.get("_row")]:
            try:
                row_nums.add(int(row))
            except Exception:
                continue
    if not row_nums:
        return data_rows
    deliverable_cols = excel_column_numbers(excel_path, ["交付物"])
    if not deliverable_cols:
        return data_rows
    with tempfile.TemporaryDirectory(prefix="document_filler_requirement_attachments_") as temp_dir:
        attachments = extract_deliverable_attachments(
            excel_path,
            row_nums,
            Path(temp_dir) / "attachments",
            deliverable_cols,
        )
        for rd in data_rows:
            files = []
            for row in rd.get("_row_numbers") or [rd.get("_row")]:
                try:
                    files.extend(attachments.get(int(row), []))
                except Exception:
                    continue
            if not files:
                continue
            existing_files = list(rd.get("_attachment_files") or [])
            seen_files = {str(item) for item in existing_files}
            for file in files:
                if str(file) not in seen_files:
                    existing_files.append(file)
                    seen_files.add(str(file))
            rd["_attachment_files"] = existing_files
            existing_names = list(rd.get("_attachment_names") or [])
            seen_names = {str(item) for item in existing_names}
            for file in files:
                name = Path(str(file)).name
                if name and name not in seen_names:
                    existing_names.append(name)
                    seen_names.add(name)
            rd["_attachment_names"] = existing_names
    return data_rows


def build_launch_identifier_line(row):
    return f"需求编号：{row.get('需求单号', '')}\t对应工单编号：{row.get('工单号', '')}"


def build_launch_requirement_description(row):
    for field_name in ("需求描述", "业务说明", "业务描述"):
        value = compact_spaces(row.get(field_name, ""))
        if value:
            return value
    subject = infer_data_report_subject(row)
    return ensure_cn_period(f"本次上线围绕{subject}完成数据报表产出、交付和使用记录归档")


def launch_image_columns_for_row(row, cell_imgs, columns):
    row_number = row.get("_row")
    return [column for column in columns if (row_number, column) in cell_imgs]


def make_cell_oxml(text, grid_span=None, bold=False, bg=None, align=None, width=None):
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    if width:
        tcw = OxmlElement("w:tcW")
        tcw.set(qn("w:w"), str(width))
        tcw.set(qn("w:type"), "dxa")
        tcPr.append(tcw)
    if grid_span:
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(grid_span))
        tcPr.append(gs)
    if bg:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)
    tc.append(tcPr)
    p = OxmlElement("w:p")
    if align:
        pp = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        pp.append(jc)
        p.append(pp)
    r = OxmlElement("w:r")
    if bold:
        rp = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rp.append(b)
        r.append(rp)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc

def make_table_oxml(headers, rows_data, col_widths=None):
    tbl = OxmlElement("w:tbl")
    tp = OxmlElement("w:tblPr")
    ts = OxmlElement("w:tblStyle")
    ts.set(qn("w:val"), "TableGrid")
    tp.append(ts)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    tp.append(tw)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tp.append(layout)
    add_solid_borders(tp)
    tbl.append(tp)
    ncols = len(headers)
    if col_widths is None:
        col_widths = [str(9000 // ncols)] * ncols
    tg = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), w)
        tg.append(gc)
    tbl.append(tg)
    hdr = OxmlElement("w:tr")
    for idx, h in enumerate(headers):
        hdr.append(make_cell_oxml(h, bold=True, bg=HEADER_BG, width=col_widths[idx]))
    tbl.append(hdr)
    for rd in rows_data:
        tr = OxmlElement("w:tr")
        for idx, val in enumerate(rd):
            width = col_widths[idx] if idx < len(col_widths) else None
            tr.append(make_cell_oxml(val, width=width))
        tbl.append(tr)
    return tbl

def insert_after_oxml(ref_elem, new_elems):
    parent = ref_elem.getparent()
    children = list(parent)
    for i, child in enumerate(children):
        if child is ref_elem:
            for j, elem in enumerate(new_elems):
                parent.insert(i + 1 + j, elem)
            return

def parse_report_names(text_after_marker):
    lines = text_after_marker.strip().split("\n")
    names = []
    for line in lines:
        line = line.strip()
        if line and len(line) > 1:
            line = re.sub(r"^[\d]+[、．.）\)]\s*", "", line)
            line = line.strip("，。；,.;")
            if len(line) > 1:
                names.append(line)
    return names

def update_toc_via_com(doc_path):
    if sys.platform != "win32":
        print("非Windows环境，跳过TOC更新，请手动在Word中更新域。")
        return
    ps_script = f'''
$word = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open("{doc_path}")
    $tocCount = $doc.TablesOfContents.Count
    for ($i = $tocCount; $i -ge 1; $i--) {{ $doc.TablesOfContents.Item($i).Delete() }}
    for ($i = 1; $i -le $doc.Paragraphs.Count; $i++) {{
        if ($doc.Paragraphs.Item($i).Range.Text -match "文档介绍" -or $doc.Paragraphs.Item($i).Range.Text -match "需求来源") {{
            if ($i -gt 1) {{
                $tocRange = $doc.Paragraphs.Item($i - 1).Range
                $toc = $doc.TablesOfContents.Add($tocRange, $true, 1, 3, $false, "", $true, $true)
                $toc.Update()
            }}
            break
        }}
    }}
    $doc.Save()
    $doc.Close()
    Write-Host "TOC updated"
}} catch {{
    Write-Host "TOC error: $_"
}} finally {{
    if ($word) {{ try {{ $word.Quit() }} catch {{}} }}
}}
'''
    try:
        result = subprocess.run(
            ["powershell", "-ExecutionPolicy", "Bypass", "-Command", ps_script],
            capture_output=True, text=True, timeout=120
        )
        if "TOC updated" in result.stdout:
            print("TOC 已更新")
        else:
            print(f"TOC: {result.stdout.strip()[:200]}")
    except subprocess.TimeoutExpired:
        print("TOC 超时，请手动在Word中更新域")
    except Exception as e:
        print(f"TOC 失败: {e}，请手动在Word中更新域")

# ══════════════════════════════════════════════════════════════
# 01-数据报表_需求文档 Fill Logic
# ══════════════════════════════════════════════════════════════


def extract_ole_by_cell(excel_path):
    import olefile as _ole
    result = {}
    with zipfile.ZipFile(excel_path, 'r') as zf:
        vml = zf.read('xl/drawings/vmlDrawing1.vml').decode('utf-8')
        sx = zf.read('xl/worksheets/sheet1.xml').decode('utf-8')
        rx = zf.read('xl/worksheets/_rels/sheet1.xml.rels').decode('utf-8')
        vml_shapes = {}
        for m in re.finditer(r'<v:shape\b(.*?)</v:shape>', vml, re.DOTALL):
            body = m.group(1); id_m = re.search(r'id="([^"]+)"', body)
            am = re.search(r'<x:Anchor>([^<]+)</x:Anchor>', body)
            if id_m and am:
                parts = am.group(1).split(',')
                vml_shapes[id_m.group(1)] = {'row': int(parts[2].strip()), 'col': int(parts[0].strip()), 'ole': 'o:ole="t"' in body}
        ole_map = {}
        for m in re.finditer(r'<oleObject[^>]*shapeId="(\d+)"[^>]*r:id="(rId\d+)"', sx):
            sn = m.group(1)
            if sn not in ole_map: ole_map[sn] = m.group(2)
        rid2f = {}
        for m in re.finditer(r'<Relationship[^>]*Id="(rId\d+)"[^>]*oleObject[^>]*Target="([^"]+)"', rx):
            rid2f[m.group(1)] = m.group(2).split('/')[-1]
        for sn, rid in ole_map.items():
            if rid not in rid2f: continue
            for vid, vs in vml_shapes.items():
                if vs.get('ole') and vid.endswith('_s' + sn):
                    fn = rid2f[rid]; vr = vs['row']; vc = vs['col']
                    try: od = zf.read(f'xl/embeddings/{fn}')
                    except: continue
                    tf = tempfile.NamedTemporaryFile(suffix='.bin', delete=False)
                    tf.write(od); tf.close()
                    ole = _ole.OleFileIO(tf.name)
                    wd = ole.openstream('WordDocument').read()
                    ole.close(); os.unlink(tf.name)
                    chunks = []; i = 0
                    while i < len(wd) - 1:
                        cc = struct.unpack_from('<H', wd, i)[0]
                        if 0x20 <= cc <= 0xFFFF and cc != 0xFFFE:
                            si = i
                            while i < len(wd) - 1:
                                cc2 = struct.unpack_from('<H', wd, i)[0]
                                if cc2 in (0x000D, 0x0007, 0x0000): break
                                i += 2
                            try:
                                text = wd[si:i].decode('utf-16-le', errors='replace')
                                if len(text.strip()) > 2: chunks.append(text)
                            except: pass
                        i += 2
                    raw = ''.join(chunks)
                    clean = re.sub(r'[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffefa-zA-Z0-9\s\.\,\;\:\!\?\-\+\=\(\)\[\]\{\}\/\\\@\#\%\&\*_\u201c\u201d\u2018\u2019\u3001\u3002\u2014\u2026\n\r\t]', '', raw)
                    clean = re.sub(r'\n{3,}', '\n\n', clean)
                    logic = clean
                    m2 = re.search(r'需求处理逻辑(.+)', clean, re.DOTALL)
                    if m2: logic = m2.group(1)
                    em2 = re.search(r'(需求验收标准|完成时间|资源保证|风险评估)', logic)
                    if em2: logic = logic[:em2.start()]
                    result[(vr, vc)] = logic.strip()
                    break
    return result
def compute_stats(data_rows):
    unique_reqs = len(set(r["需求单号"] for r in data_rows if r["需求单号"]))
    unique_wos = len(set(r["工单号"] for r in data_rows if r["工单号"]))
    total_cnt = sum(int(r["报表统计次数"]) for r in data_rows if r["报表统计次数"].isdigit())
    return unique_reqs, unique_wos, total_cnt


def summarize_biz_logic(text):
    if not text: return [("1", "详见附件需求规格说明书")]
    steps = []; text = text.strip()
    markers = list(re.finditer(r'(?:^|(?<=[\n\u3002\u3001])\s*)(\d+)[.\u3001)\uff09]\s*', text))
    if len(markers) >= 2:
        for i, m in enumerate(markers):
            start = m.end(); end = markers[i+1].start() if i+1 < len(markers) else len(text)
            cp = text[start:end].strip()
            if len(cp) > 3: steps.append((str(i+1), xml_safe(cp[:250])))
    else:
        segs = re.split(r'(?<=[。、])\s*', text); sn = 1; cur = ""
        for seg in segs:
            seg = seg.strip()
            if not seg: continue
            if len(cur) + len(seg) > 200 and cur:
                steps.append((str(sn), xml_safe(cur[:250]))); sn += 1; cur = seg
            else:
                cur = (cur + "。" + seg) if cur else seg
        if cur: steps.append((str(sn), xml_safe(cur[:250])))
    if not steps: steps = [("1", xml_safe(text[:200]))]
    return steps[:15]


def parse_report_list(text):
    if not text: return []
    items = []
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line: continue
        parts = line.rsplit(None, 1) if ' ' in line else [line, '']
        cn = parts[0].strip(); en = parts[1].strip() if len(parts) > 1 else ''
        if cn: items.append((cn, en))
    return items


def norm_table_name(name):
    text = (name or "").strip()
    if "." in text:
        text = text.split(".")[-1]
    return re.sub(r"[^A-Za-z0-9_]", "", text).upper()


def merged_value_getter(ws):
    merged_values = {}
    for merged_range in ws.merged_cells.ranges:
        value = ws.cell(merged_range.min_row, merged_range.min_col).value
        for row in range(merged_range.min_row, merged_range.max_row + 1):
            for col in range(merged_range.min_col, merged_range.max_col + 1):
                merged_values[(row, col)] = value

    def get(row, col):
        value = ws.cell(row, col).value
        if value is None:
            value = merged_values.get((row, col))
        return str(value).strip() if value is not None else ""

    return get


def read_stats_requirement_groups(excel_path, service_dir):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    if "服务目录" not in headers:
        raise ValueError("台账清单缺少「服务目录」列")
    get = merged_value_getter(ws)
    service_col = headers.index("服务目录") + 1
    logic_col = headers.index("业务逻辑") if "业务逻辑" in headers else -1
    groups = []
    by_key = {}

    for row in range(2, ws.max_row + 1):
        if get(row, service_col) != service_dir:
            continue
        record = {header: get(row, idx + 1) for idx, header in enumerate(headers) if header}
        record = normalize_ledger_record(record)
        key = record.get("工单号") or f"{record.get('需求单号', '')}|{record.get('工单内容', '')}|{row}"
        if key not in by_key:
            record["_row"] = row
            record["_vml_row"] = row - 1
            record["_vml_col"] = logic_col
            record["results"] = []
            by_key[key] = record
            groups.append(record)

        result_text = record.get("统计分析结果表清单", "")
        for result in parse_report_list(result_text):
            if result not in by_key[key]["results"]:
                by_key[key]["results"].append(result)

    for group in groups:
        group["统计分析结果表清单"] = "\n".join(
            f"{cn} {en}".strip() for cn, en in group.get("results", [])
        )

    if not groups:
        raise ValueError(f"未找到匹配数据: 服务目录={service_dir}")
    return groups


def read_data_report_design_groups(excel_path, service_dir):
    wb = openpyxl.load_workbook(excel_path, data_only=False)
    ws = wb.active
    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    if "服务目录" not in headers:
        raise ValueError("台账清单缺少「服务目录」列")
    get = merged_value_getter(ws)
    service_col = headers.index("服务目录") + 1
    result_col_index = ledger_header_index(headers, "统计分析结果表清单")
    result_col = result_col_index + 1 if result_col_index >= 0 else None
    xml_col = headers.index("程序XML文本") + 1 if "程序XML文本" in headers else None
    logic_col = headers.index("数据处理逻辑") if "数据处理逻辑" in headers else -1
    groups = []
    by_key = {}

    for row in range(2, ws.max_row + 1):
        if get(row, service_col) != service_dir:
            continue
        record = {header: get(row, idx + 1) for idx, header in enumerate(headers) if header}
        record = normalize_ledger_record(record)
        key = record.get("工单号") or f"{record.get('需求单号', '')}|{record.get('工单内容', '')}|{row}"
        if key not in by_key:
            record["_row"] = row
            record["_vml_row"] = row - 1
            record["_vml_col"] = logic_col
            record["_row_numbers"] = []
            record["_programs"] = []
            by_key[key] = record
            groups.append(record)

        group = by_key[key]
        group["_row_numbers"].append(row)
        program_text = ""
        if result_col:
            program_text = str(ws.cell(row, result_col).value or "").strip()
            if not program_text:
                program_text = get(row, result_col)
        xml_text = ""
        if xml_col:
            xml_text = str(ws.cell(row, xml_col).value or "").strip()
            if not xml_text:
                xml_text = get(row, xml_col)
        for program_cn, program_en in parse_report_list(program_text):
            program = {
                "program_cn": program_cn,
                "program_en": program_en,
                "xml": xml_text,
                "row": row,
            }
            if program not in group["_programs"]:
                group["_programs"].append(program)

    for group in groups:
        group["统计分析结果表清单"] = "\n".join(
            f"{item['program_cn']} {item['program_en']}".strip()
            for item in group.get("_programs", [])
        )
        group["_row_numbers"] = sorted(set(group.get("_row_numbers", [group["_row"]])))

    if not groups:
        raise ValueError(f"未找到匹配数据: 服务目录={service_dir}")
    return groups


def resolve_stats_relation_workbook(template_path, relation_hint=None, output_path=None):
    candidates = []
    if relation_hint:
        candidates.append(Path(relation_hint))
    if output_path:
        candidates.append(Path(output_path).with_name("04-数据统计分析_结果表及使用说明.xlsx"))
    template = Path(template_path)
    candidates.append(template.with_name("04-数据统计分析_结果表及使用说明.xlsx"))
    candidates.append(template.parent / "04-数据统计分析_结果表及使用说明.xlsx")

    seen = set()
    for candidate in candidates:
        if not candidate or candidate in seen or not candidate.exists():
            continue
        seen.add(candidate)
        try:
            wb = openpyxl.load_workbook(candidate, read_only=True, data_only=True)
            if "2、表融合关系" in wb.sheetnames:
                wb.close()
                return str(candidate)
            wb.close()
        except Exception:
            continue
    return None


def _parse_relation_title(text):
    value = str(text or "").strip()
    if not value:
        return ""
    match = re.search(r"([A-Za-z][A-Za-z0-9_$.]{2,})\s*$", value)
    return norm_table_name(match.group(1)) if match else ""


def load_stats_relation_descriptions(relation_path):
    wb = openpyxl.load_workbook(relation_path, data_only=True)
    if "2、表融合关系" not in wb.sheetnames:
        wb.close()
        return {}
    ws = wb["2、表融合关系"]
    descriptions = {}
    current_result = ""
    for row in range(1, ws.max_row + 1):
        first_value = ws.cell(row, 1).value
        parsed_result = _parse_relation_title(first_value)
        if parsed_result:
            current_result = parsed_result
        for col in range(1, ws.max_column + 1):
            label = str(ws.cell(row, col).value or "").strip()
            if label == "文字描述" and current_result:
                desc = ws.cell(row, col + 1).value if col + 1 <= ws.max_column else ""
                if desc:
                    descriptions[current_result] = sanitize_stats_logic_text(desc)
                break
    wb.close()
    return descriptions


def _clean_logic_step(text):
    text = sanitize_stats_logic_text(xml_safe(str(text or "")))
    text = re.sub(r"^\s*\d+\s*[\.、．)\uff09]?\s*", "", text)
    text = re.sub(r"[\t ]+", " ", text)
    return text.strip(" \r\n\t；;。")


def split_logic_description(text):
    text = str(text or "").strip()
    if not text:
        return []
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    markers = list(re.finditer(r"(?:^|\n)\s*(\d+)\s*(?:[\.、．)\uff09]|\t|\s)\s*", normalized))
    steps = []
    if len(markers) >= 2:
        for idx, marker in enumerate(markers):
            start = marker.end()
            end = markers[idx + 1].start() if idx + 1 < len(markers) else len(normalized)
            piece = _clean_logic_step(normalized[start:end])
            if piece:
                steps.append(piece)
    else:
        for line in normalized.split("\n"):
            piece = _clean_logic_step(line)
            if piece:
                steps.append(piece)

    if len(steps) <= 1:
        source = steps[0] if steps else _clean_logic_step(normalized)
        steps = [
            _clean_logic_step(piece)
            for piece in re.split(r"(?<=[。；;])\s*", source)
            if _clean_logic_step(piece)
        ]

    return steps[:15]


def build_stats_requirement_business_logic_steps(results, descriptions):
    if not results:
        return [("1", "未匹配到统计分析结果表，请手动补充业务逻辑说明。")]

    all_steps = []
    for cn, en in results:
        desc = descriptions.get(norm_table_name(en), "")
        pieces = split_logic_description(desc)
        table_label = f"{cn}（{en}）" if en else cn
        if not pieces:
            all_steps.append(f"{table_label}：未在结果表及使用说明中匹配到文字描述，请手动补充。")
        elif len(results) == 1:
            all_steps.extend(pieces)
        else:
            merged = "；".join(piece.rstrip("。；;") for piece in pieces)
            all_steps.append(f"{table_label}：{merged}。")

    return [(str(idx), xml_safe(sanitize_stats_logic_text(step)[:500])) for idx, step in enumerate(all_steps, start=1)]


def make_biz_table(header_text, rows_data):
    tbl = OxmlElement("w:tbl"); tp = OxmlElement("w:tblPr")
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), "5000"); tw.set(qn("w:type"), "pct"); tp.append(tw)
    add_solid_borders(tp); tbl.append(tp)
    tg = OxmlElement("w:tblGrid")
    tg.append(OxmlElement("w:gridCol")); tg[-1].set(qn("w:w"), "900")
    tg.append(OxmlElement("w:gridCol")); tg[-1].set(qn("w:w"), "8100"); tbl.append(tg)
    tr0 = OxmlElement("w:tr")
    tr0.append(make_cell_oxml(header_text, grid_span=2, bold=True, bg=HEADER_BG, align="center"))
    tbl.append(tr0)
    tr1 = OxmlElement("w:tr")
    tr1.append(make_cell_oxml("步骤", bold=True, bg=HEADER_BG))
    tr1.append(make_cell_oxml("说明", bold=True, bg=HEADER_BG))
    tbl.append(tr1)
    for sn, st in rows_data:
        tr = OxmlElement("w:tr")
        tr.append(make_cell_oxml(sn))
        tr.append(make_cell_oxml(st))
        tbl.append(tr)
    return tbl


def make_stats_source_table(headers, rows_data):
    tbl = OxmlElement("w:tbl"); tp = OxmlElement("w:tblPr")
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), "5000"); tw.set(qn("w:type"), "pct"); tp.append(tw)
    add_solid_borders(tp); tbl.append(tp)
    n = len(headers)
    tg = OxmlElement("w:tblGrid")
    for _ in range(n):
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(9000 // n)); tg.append(gc)
    tbl.append(tg)
    tr = OxmlElement("w:tr")
    for h in headers:
        tr.append(make_cell_oxml(h, bold=True, bg=HEADER_BG))
    tbl.append(tr)
    for rd in rows_data:
        tr = OxmlElement("w:tr")
        for v in rd: tr.append(make_cell_oxml(v))
        tbl.append(tr)
    return tbl
def fill_requirement_doc(data_rows, template_path, output_path, catalog_context=None):
    """填充 01-数据报表_需求文档"""
    req_count, wo_count, total_reports = compute_stats(data_rows)
    print(f"结果: {len(data_rows)} 条, 需求单={req_count}, 工单={wo_count}, 报表={total_reports}")

    doc = Document(template_path)
    body = doc.element.body

    # ── Update count description ──
    count_text = f"服务周期内，共有{req_count}张需求单，{wo_count}张工单涉及{total_reports}次数据报表服务。具体需求单、工单和产出如下表："
    for p in doc.paragraphs:
        if p.style.name == "Body Text" and "服务周期内" in p.text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = count_text
            break

    # ── Table 2 rebuild ──
    t2 = doc.tables[2]
    tbl_elem = t2._tbl
    add_solid_borders(tbl_elem.find(qn('w:tblPr')))
    for tr_elem in list(tbl_elem.findall(qn("w:tr")))[1:]:
        tbl_elem.remove(tr_elem)
    for idx, rd in enumerate(data_rows):
        tr = OxmlElement("w:tr")
        for val in [str(idx + 1), rd["需求单号"], rd["工单号"], rd["工单内容"], rd["报表统计次数"]]:
            tr.append(make_cell_oxml(val))
        tbl_elem.append(tr)
    tr = OxmlElement("w:tr")
    tr.append(make_cell_oxml("合计", grid_span=2))
    tr.append(make_cell_oxml(""))
    tr.append(make_cell_oxml(""))
    tr.append(make_cell_oxml(str(total_reports)))
    tbl_elem.append(tr)

    # ── Section 3 ──
    style_ids = {
        "Heading 1": doc.styles["Heading 1"].style_id,
        "Heading 2": doc.styles["Heading 2"].style_id,
        "Heading 3": doc.styles["Heading 3"].style_id,
        "Body Text": doc.styles["Body Text"].style_id,
    }

    h1_content = h1_other = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            if "需求内容" in p.text:
                h1_content = p._element
            elif "其他要求" in p.text:
                h1_other = p._element
                break
    if h1_content is None or h1_other is None:
        raise ValueError("模板中未找到 需求内容 或 其他要求 标题")

    remove = False
    to_remove = []
    for child in list(body):
        if child is h1_content:
            remove = True
            continue
        elif child is h1_other:
            remove = False
            continue
        if remove:
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)

    new_elems = []
    for rd in data_rows:
        rd = normalize_data_report_text_fields(rd, catalog_context)
        gongdan = rd["工单内容"]
        biz = rd["业务说明"]
        data_req = rd["数据需求"]
        delivery = rd["交付要求"]

        new_elems.append(mp(gongdan, style_ids["Heading 2"]))
        new_elems.append(mp("业务说明", style_ids["Heading 3"]))

        biz_match = re.search(r"本次工作拟产出以下(\d+)份报表成果", biz)
        if not biz_match:
            biz_match = re.search(r"本次工作拟产出(\d+)份报表成果", biz)

        if biz_match:
            clean_biz = biz[:biz_match.start()].strip().rstrip("\n\r")
            new_elems.append(mp(clean_biz, style_ids["Body Text"], 480))
            report_names = parse_report_names(biz[biz_match.end():])
            if report_names:
                new_elems.append(mp(
                    f"本次工作拟产出以下{len(report_names)}份报表成果：",
                    style_ids["Body Text"], 480))
                table_rows = [[str(i + 1), name] for i, name in enumerate(report_names)]
                new_elems.append(make_table_oxml(["序号", "名称"], table_rows))
                print(f"  [{gongdan[:30]}] 内嵌表格: {len(report_names)} 行")
        else:
            new_elems.append(mp(biz, style_ids["Body Text"], 480))

        new_elems.append(mp("数据需求", style_ids["Heading 3"]))
        new_elems.append(mp(data_req, style_ids["Body Text"], 480))
        new_elems.append(mp("交付要求", style_ids["Heading 3"]))
        new_elems.append(mp(delivery, style_ids["Body Text"], 480))

    insert_after_oxml(h1_content, new_elems)
    print(f"第3节: {len(new_elems)} 个元素")

    doc.save(output_path)
    print(f"已保存: {output_path}")
    update_toc_via_com(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════
# 02-数据报表_设计文档 Fill Logic
# ══════════════════════════════════════════════════════════════

INDICATOR_HEADERS = ["所属报表名称", "指标名称", "指标定义", "指标业务口径", "指标展示形式", "与其他指标的关联关系", "备注"]
INDICATOR_DATA = [
    ["目录信息", "总条数", "各报表目录的总数据量条数", "无", "文本", "无", "总条数"],
    ["各报表目录详情", "库名", "目录所在库名", "无", "文本", "无", "库名"],
    ["各报表目录详情", "资源表名", "该目录所挂载的资源表名", "无", "文本", "无", "资源表名"],
    ["各报表目录详情", "表注释", "表中文名", "无", "文本", "无", "表注释"],
    ["各报表目录详情", "字段名", "字段英文名", "无", "文本", "无", "字段名"],
    ["各报表目录详情", "字段注释", "字段中文名", "无", "文本", "无", "字段注释"],
    ["各报表目录详情", "总数", "该目录中字段总条数", "无", "文本", "无", "总数"],
    ["各报表目录详情", "空值数", "该目录字段空值数", "无", "文本", "无", "空值数"],
    ["各报表目录详情", "空值率", "该目录字段空值率", "无", "文本", "无", "空值率"],
    ["各报表目录详情", "样例数据", "该字段示例", "无", "文本", "无", "样例数据"],
]

def clean_attachment_report_name(value):
    name = Path(str(value or "")).name.strip()
    stem = Path(name).stem if name else ""
    generated_names = {"Ole10Native_embedded", "Workbook_embedded", "Workbook", "package"}
    if stem in generated_names or re.fullmatch(r"deliverable_row\d+", stem, flags=re.I):
        return ""
    return stem or name


def _attachment_report_entry(value):
    name = Path(str(value or "")).name.strip()
    report_name = clean_attachment_report_name(value)
    if not report_name:
        return "", ""
    return report_name, Path(name).suffix.lower()


def collect_attachment_report_names(files):
    entries = []
    for path in files:
        report_name, suffix = _attachment_report_entry(path)
        if report_name:
            entries.append((report_name, suffix))
    if any(suffix not in {".zip", ".rar", ".7z"} for _name, suffix in entries):
        entries = [entry for entry in entries if entry[1] not in {".zip", ".rar", ".7z"}]
    names = []
    for report_name, _suffix in entries:
        if report_name not in names:
            names.append(report_name)
    return names


def _extract_sqls_from_program_xml(xml_text):
    if not xml_text:
        return []
    try:
        builder = load_stats_result_builder()
        nodes, _edges = builder.parse_xml_program(xml_text)
        return [str(node.get("sql") or "") for node in nodes if node.get("sql")]
    except Exception:
        pass
    text = str(xml_text)
    sqls = []
    for match in re.finditer(r'modelData="([^"]*)"', text):
        raw = html.unescape(match.group(1))
        try:
            value = json.loads(raw)
        except Exception:
            value = {}
        sql = str(value.get("sql") or "") if isinstance(value, dict) else ""
        if sql:
            sqls.append(html.unescape(sql))
    if sqls:
        return sqls
    return [html.unescape(re.sub(r"</?[^>]+>", " ", text))]


def _candidate_sql_targets(sqls, preferred_target):
    builder = load_stats_result_builder()
    targets = []
    preferred = builder.norm_name(preferred_target) if preferred_target else ""
    if preferred:
        targets.append(preferred)
    pattern = re.compile(
        r"create\s+(?:external\s+)?table\s+(?:if\s+not\s+exists\s+)?[`\"]?([A-Za-z_][\w]*(?:\.[A-Za-z_][\w]*)?|\$\{[^}]+\})[`\"]?",
        flags=re.I,
    )
    for sql in sqls:
        for match in pattern.finditer(sql):
            target = builder.norm_name(match.group(1))
            if target and target not in targets:
                targets.append(target)
    return targets


INSERT_FIELD_LABELS = {
    "CATA_CODE": "目录代码",
    "CATA_TITLE": "目录名称",
    "PROVIDER_NAME": "提供方名称",
    "SCHEMA_NAME": "库名",
    "TABLE_NAME": "资源表名",
    "TABLE_COMMENTS": "表注释",
    "CLOUMN_NAME": "字段名",
    "COLUMN_NAME": "字段名",
    "CLOUMN_COMMENTS": "字段注释",
    "COLUMN_COMMENTS": "字段注释",
    "DATA_COUNT": "数据量",
    "EMPTY_VALUE": "空值数",
    "EMPTY_VALUE_RATE": "空值率",
    "IS_PROCESS": "是否加工",
    "UPDATE_FREQUENCY": "更新频率",
    "LAST_UPDATE_TIME": "最后更新时间",
    "SAMPLE_DATA1": "样例数据1",
    "SAMPLE_DATA2": "样例数据2",
    "SAMPLE_DATA3": "样例数据3",
    "SAMPLE_DATA4": "样例数据4",
    "SAMPLE_DATA5": "样例数据5",
}


def split_sql_csv(text):
    items = []
    current = ""
    depth = 0
    quote = None
    for ch in text:
        if quote:
            current += ch
            if ch == quote:
                quote = None
            continue
        if ch in "'\"`":
            quote = ch
            current += ch
            continue
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        elif ch == "," and depth == 0:
            if current.strip():
                items.append(current.strip())
            current = ""
            continue
        current += ch
    if current.strip():
        items.append(current.strip())
    return items


def extract_insert_field_labels(sqls):
    labels = []
    pattern = re.compile(r"insert\s+into\s+[^\s(]+\s*\((.*?)\)\s*values\s*\(", flags=re.I | re.S)
    for sql in sqls:
        for match in pattern.finditer(sql):
            for raw in split_sql_csv(match.group(1)):
                key = re.sub(r"[^A-Za-z0-9_]", "", raw).upper()
                if not key:
                    continue
                label = INSERT_FIELD_LABELS.get(key, raw.strip("`\" "))
                if label and label not in labels:
                    labels.append(label)
    return labels


def extract_indicator_field_comments_from_program(program):
    sqls = _extract_sqls_from_program_xml(program.get("xml", ""))
    if not sqls:
        return []
    builder = load_stats_result_builder()
    fields = []
    for target in _candidate_sql_targets(sqls, program.get("program_en", "")):
        fields = builder.parse_result_fields_from_sql(sqls, target)
        if fields:
            break
    if not fields:
        for sql in sqls:
            fields = builder.parse_result_fields_from_sql([sql], "")
            if fields:
                break
    if not fields:
        return extract_insert_field_labels(sqls)
    comments = []
    for field in fields:
        comment = (
            str(field.get("字段注释") or "").strip()
            or str(field.get("字段中文名") or "").strip()
            or str(field.get("字段英文名") or "").strip()
        )
        if comment and comment not in comments:
            comments.append(comment)
    return comments


def build_data_report_indicator_rows(programs, attachment_names):
    report_names = collect_attachment_report_names(attachment_names)
    if not report_names:
        report_names = ["本次交付材料"]
    rows = []
    for index, program in enumerate(programs or []):
        report_name = report_names[index % len(report_names)]
        field_comments = program.get("field_comments")
        if field_comments is None:
            field_comments = extract_indicator_field_comments_from_program(program)
        if not field_comments:
            field_comments = [program.get("program_cn") or program.get("program_en") or "未解析到字段信息"]
        for comment in field_comments:
            text = str(comment or "").strip()
            if not text:
                continue
            rows.append([report_name, text, text, "无", "文本", "无", text])
    return rows or INDICATOR_DATA


def mk_indicator_table(rows_data=None):
    rows_data = rows_data if rows_data is not None else INDICATOR_DATA
    tbl = OxmlElement("w:tbl")
    tp = OxmlElement("w:tblPr")
    ts = OxmlElement("w:tblStyle")
    ts.set(qn("w:val"), "TableGrid")
    tp.append(ts)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    tp.append(tw)
    add_solid_borders(tp)
    tbl.append(tp)
    n = 7
    tg = OxmlElement("w:tblGrid")
    for _ in range(n):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(9000 // n))
        tg.append(gc)
    tbl.append(tg)
    tr_h = OxmlElement("w:tr")
    for h_text in INDICATOR_HEADERS:
        tr_h.append(make_cell_oxml(h_text, bold=True, bg=HEADER_BG))
    tbl.append(tr_h)
    for rd in rows_data:
        tr = OxmlElement("w:tr")
        for v in rd:
            tr.append(make_cell_oxml(v))
        tbl.append(tr)
    return tbl

def mk_biz_table(items):
    tbl = OxmlElement("w:tbl")
    tp = OxmlElement("w:tblPr")
    ts = OxmlElement("w:tblStyle")
    ts.set(qn("w:val"), "TableGrid")
    tp.append(ts)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    tp.append(tw)
    add_solid_borders(tp)
    tbl.append(tp)
    tg = OxmlElement("w:tblGrid")
    for w in ["2500", "2500"]:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), w)
        tg.append(gc)
    tbl.append(tg)
    for title, content in items:
        tr = OxmlElement("w:tr")
        tr.append(make_cell_oxml(title, bold=True, bg=HEADER_BG))
        tr.append(make_cell_oxml(content))
        tbl.append(tr)
    return tbl

def mk_ds_table(table_title, fields):
    tbl = OxmlElement("w:tbl")
    tp = OxmlElement("w:tblPr")
    ts = OxmlElement("w:tblStyle")
    ts.set(qn("w:val"), "TableGrid")
    tp.append(ts)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    tp.append(tw)
    add_solid_borders(tp)
    tbl.append(tp)
    tg = OxmlElement("w:tblGrid")
    for _ in range(4):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), "2250")
        tg.append(gc)
    tbl.append(tg)
    # Row 0: merged title (name + code)
    tr0 = OxmlElement("w:tr")
    tr0.append(make_cell_oxml(table_title, grid_span=4, bold=True))
    tbl.append(tr0)
    # Row 1: column headers
    tr1 = OxmlElement("w:tr")
    for h_text in ["字段中文名称", "字段英文名称", "数据类型", "备注"]:
        tr1.append(make_cell_oxml(h_text, bold=True, bg=HEADER_BG))
    tbl.append(tr1)
    for f in fields:
        tr = OxmlElement("w:tr")
        for key in ["数据项名称", "英文名称", "数据类型", "备注"]:
            tr.append(make_cell_oxml(f.get(key, "")))
        tbl.append(tr)
    return tbl

def extract_images_from_excel(excel_path):
    """Extract embedded images from Excel, return dict mapping (row,col) -> bytes."""
    try:
        with zipfile.ZipFile(excel_path, "r") as z:
            all_images = [z.read(f) for f in sorted(n for n in z.namelist() if n.startswith("xl/media/image") and n.endswith(".png"))]
    except Exception:
        all_images = []
    wb_img = openpyxl.load_workbook(excel_path)
    ws_img = wb_img.active
    positions = sorted([
        (r, c) for r in range(2, ws_img.max_row + 1)
        for c in range(1, ws_img.max_column + 1)
        if ws_img.cell(row=r, column=c).value and "DISPIMG" in str(ws_img.cell(row=r, column=c).value)
    ])
    return {pos: all_images[i] if i < len(all_images) else None for i, pos in enumerate(positions)}


def extract_images_via_cellimages(excel_path):
    """Parse xl/cellimages.xml (WPS format) to get image bytes by name."""
    import zipfile as _zf
    result = {}
    with _zf.ZipFile(excel_path, 'r') as zf:
        if 'xl/cellimages.xml' not in zf.namelist():
            return result
        ci = zf.read('xl/cellimages.xml').decode('utf-8')
        name_to_rid = {}
        for block in re.findall(r'<etc:cellImage>(.*?)</etc:cellImage>', ci, re.DOTALL):
            nm = re.search(r'name="([^"]+)"', block)
            rm = re.search(r'r:embed="(rId\d+)"', block)
            if nm and rm:
                name_to_rid[nm.group(1)] = rm.group(1)
        rels = zf.read('xl/_rels/cellimages.xml.rels').decode('utf-8')
        rid_to_file = {}
        for m in re.finditer(r'<Relationship[^>]*Id="(rId\d+)"[^>]*Target="([^"]+)"', rels):
            target = m.group(2)
            if target.startswith("../"):
                target = "xl/" + target[3:]
            elif target.startswith("xl/"):
                target = target
            else:
                target = "xl/" + target.lstrip("/")
            rid_to_file[m.group(1)] = target
        for name, rid in name_to_rid.items():
            if rid in rid_to_file:
                try:
                    result[name] = zf.read(rid_to_file[rid])
                except:
                    pass
    return result


def match_images_to_cells(excel_path, img_bytes_by_name, img_cols, row_numbers):
    """Match DISPIMG formulas to images by name for given row numbers."""
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    result = {}
    for r in row_numbers:
        for cn in img_cols:
            if cn not in headers:
                continue
            ci = headers.index(cn) + 1
            val = ws.cell(row=r, column=ci).value
            if val and 'DISPIMG' in str(val):
                m = re.search(r'ID_([A-F0-9]+)', str(val))
                if m:
                    iname = "ID_" + m.group(1)
                    if iname in img_bytes_by_name:
                        result[(r, cn)] = img_bytes_by_name[iname]
    wb.close()
    return result

def load_catalog_data(catalog_path, all_codes):
    """Load resource info and field data from catalog Excel."""
    ensure_module("pandas")
    import pandas as pd
    df_r = pd.read_excel(catalog_path, sheet_name="关联资源信息")
    df_r["数据目录代码"] = df_r["数据目录代码"].astype(str).str.strip()
    rmap = {}
    for _, r in df_r[df_r["数据目录代码"].isin(all_codes)].iterrows():
        rmap[r["数据目录代码"]] = {
            "资源名称": str(r.get("资源名称", "")).strip(),
            "资源编码": str(r.get("资源编码", "")).strip()
        }
    df_i = pd.read_excel(catalog_path, sheet_name="数据项")
    df_i["数据目录代码"] = df_i["数据目录代码"].astype(str).str.strip()
    fmap = {c: [] for c in all_codes}
    for _, r in df_i[df_i["数据目录代码"].isin(all_codes)].iterrows():
        desc = str(r.get("字段描述", "")).strip()
        if not desc or desc.lower() in ("nan", "none"):
            desc = "No"
        fmap[r["数据目录代码"]].append({
            "数据项名称": str(r.get("数据项名称", "")).strip(),
            "英文名称": str(r.get("英文名称", "")).strip(),
            "数据类型": str(r.get("数据类型", "")).strip(),
            "备注": desc
        })
    return rmap, fmap

def fill_design_doc(data_rows, template_path, output_path, catalog_path):
    """Fill 02-数据报表_设计文档"""
    print(f"结果: {len(data_rows)} 条记录")

    # Extract images from ledger
    img_by_pos = extract_images_from_excel(template_path)  # placeholder; we use excel in extract call
    # Actually we need to extract from the original excel passed via --excel, but we already have data_rows.
    # We need to re-load the excel for images. Let us accept excel_path as additional param.
    # For now, we accept that images are extracted from the same excel that was read.
    # But data_rows are already filtered. We need the original excel path.

    # Since data_rows already passed in, we need the excel path separately for images.
    # Let me adjust: fill_design_doc needs excel_path too.
    pass  # Will be restructured below

def fill_design_doc_full(excel_path, data_rows, template_path, output_path, catalog_path):
    """Complete 02-设计文档 fill, with image extraction from excel."""
    # Extract images via cellimages.xml for name-based matching (not positional)
    imgs = extract_images_via_cellimages(excel_path)
    row_nums = set()
    for rd in data_rows:
        row_nums.update(rd.get("_row_numbers") or [rd["_row"]])
    img_cols = ["02-数据报表_设计文档-数据内容截图", "数据处理逻辑"]
    cell_imgs = match_images_to_cells(excel_path, imgs, img_cols, row_nums)
    print(f"匹配到 {len(cell_imgs)} 张图片")
    attachment_previews = build_attachment_previews(excel_path, row_nums)
    if attachment_previews:
        print(f"附件截图: {len(attachment_previews)} 张")

    deliverable_cols = excel_column_numbers(excel_path, ["交付物"]) or None
    with tempfile.TemporaryDirectory(prefix="document_filler_indicator_attachments_") as temp_dir:
        attachments = extract_deliverable_attachments(excel_path, row_nums, Path(temp_dir) / "attachments", deliverable_cols)
        for rd in data_rows:
            files = []
            for row in rd.get("_row_numbers") or [rd["_row"]]:
                files.extend(attachments.get(row, []))
            rd["_attachment_names"] = collect_attachment_report_names(files)

    catalog_context = build_data_report_catalog_context(catalog_path, data_rows)

    def first_group_payload(mapping, rd, col_name=None):
        for row in rd.get("_row_numbers") or [rd["_row"]]:
            key = (row, col_name) if col_name else row
            payload = mapping.get(key)
            if payload:
                return payload
        return None

    for rd in data_rows:
        rd.update(normalize_data_report_text_fields(rd, catalog_context))
        rd["_img_content"] = (
            first_group_payload(cell_imgs, rd, "02-数据报表_设计文档-数据内容截图")
            or first_group_payload(attachment_previews, rd)
        )
        rd["_img_logic"] = first_group_payload(cell_imgs, rd, "数据处理逻辑")
        rd["_logic_is_img"] = "DISPIMG" in rd.get("数据处理逻辑", "")

    # Collect all directory codes
    all_codes = set()
    code_col = "02-数据报表_设计文档-数据来源库表清单对应数据目录代码"
    for rd in data_rows:
        codes_str = rd.get(code_col, "")
        rd["_codes"] = [c.strip() for c in codes_str.split("\n") if c.strip()]
        all_codes.update(rd["_codes"])

    # Load catalog
    rmap, fmap = load_catalog_data(catalog_path, all_codes)

    # Open template and find insertion point
    doc = Document(template_path)
    body = doc.element.body
    all_c = list(body)

    h1d_idx = None
    for i, c in enumerate(all_c):
        tag = c.tag.split("}")[-1]
        if tag == "p":
            for pi, p in enumerate(doc.paragraphs):
                if p._element is c and p.style.name == "Heading 1" and "数据报表设计" in p.text:
                    h1d_idx = i
                    break
        if h1d_idx is not None:
            break
    if h1d_idx is None:
        raise ValueError("模板中未找到「数据报表设计」 Heading 1 标题")

    for i in range(len(all_c) - 1, h1d_idx, -1):
        body.remove(all_c[i])

    S = {
        "H2": doc.styles["Heading 2"].style_id,
        "H3": doc.styles["Heading 3"].style_id,
        "H4": doc.styles["Heading 4"].style_id,
        "BT": doc.styles["Body Text"].style_id,
        "NL": doc.styles["Normal"].style_id,
    }

    print("构建文档内容...")
    for idx, rd in enumerate(data_rows):
        gd = rd["工单内容"]
        codes = rd.get("_codes", [])

        body.append(mp(gd, S["H2"]))
        body.append(mp("业务分析", S["H3"]))
        body.append(mk_biz_table([
            ("内容描述", rd.get("内容描述", "")),
            ("业务场景", rd.get("业务场景", "")),
            ("数据内容", rd.get("数据内容", "")),
            ("结果形式", rd.get("结果形式", "")),
        ]))

        body.append(mp("数据内容", S["H3"]))
        body.append(mp("", S["NL"]))

        body.append(mp("数据来源库表清单", S["H3"]))
        body.append(mp("为完成本次数据报表，需使用到以下数据：", S["BT"], 480))

        if codes:
            for code in codes:
                res = rmap.get(code, {})
                rn = res.get("资源名称", "")
                rc = res.get("资源编码", "")
                h4_title = rc if rc else code
                table_title = f"{rn} {rc}" if rn and rc else code
                body.append(mp(h4_title, S["H4"]))
                fs = fmap.get(code, [])
                if fs:
                    body.append(mk_ds_table(table_title, fs))
                else:
                    body.append(mp("未匹配到此目录的任何表结构信息，请手动补充。", S["BT"], 480))
        else:
            body.append(mp("（无数据来源库表清单）", S["BT"], 480))

        body.append(mp("数据处理逻辑", S["H3"]))
        if rd.get("_logic_is_img") and rd.get("_img_logic"):
            body.append(mp("", S["NL"]))
        else:
            body.append(mp(rd.get("数据处理逻辑", ""), S["BT"], 480))

        body.append(mp("报表指标设计", S["H3"]))
        body.append(mk_indicator_table(build_data_report_indicator_rows(rd.get("_programs", []), rd.get("_attachment_names", []))))

    # Insert images
    h3s = [(i, p) for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 3"]
    for ri, pi in enumerate([i for i, p in h3s if p.text.strip() == "数据内容"]):
        if ri < len(data_rows) and data_rows[ri].get("_img_content") and pi + 1 < len(doc.paragraphs):
            np = doc.paragraphs[pi + 1]
            if np.style.name == "Normal" and not np.text.strip():
                tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                tf.write(data_rows[ri]["_img_content"])
                tf.close()
                for re_e in np._element.findall(qn("w:r")):
                    np._element.remove(re_e)
                np.add_run().add_picture(tf.name, width=Inches(5.5))
                os.unlink(tf.name)

    for ri, pi in enumerate([i for i, p in h3s if p.text.strip() == "数据处理逻辑"]):
        if ri < len(data_rows) and data_rows[ri].get("_logic_is_img") and data_rows[ri].get("_img_logic") and pi + 1 < len(doc.paragraphs):
            np = doc.paragraphs[pi + 1]
            if np.style.name == "Normal" and not np.text.strip():
                tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                tf.write(data_rows[ri]["_img_logic"])
                tf.close()
                for re_e in np._element.findall(qn("w:r")):
                    np._element.remove(re_e)
                np.add_run().add_picture(tf.name, width=Inches(5.5))
                os.unlink(tf.name)

    doc.save(output_path)
    print(f"已保存: {output_path}")
    update_toc_via_com(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════
# 03-数据报表_上线记录 Fill Logic
# ══════════════════════════════════════════════════════════════

def fill_launch_record_doc(excel_path, data_rows, template_path, output_path):
    """Fill 03-数据报表_上线记录 template."""
    img_cols = ["上线交付截图1", "上线交付截图2", "使用记录截图1", "使用记录截图2"]
    print("提取图片...")
    imgs = extract_images_via_cellimages(excel_path)
    row_nums = set(rd["_row"] for rd in data_rows)
    cell_imgs = match_images_to_cells(excel_path, imgs, img_cols, row_nums)
    print(f"匹配到 {len(cell_imgs)} 张图片")

    doc = Document(template_path)
    body = doc.element.body

    S = {}
    for sn in ["Heading 1", "Heading 2", "Heading 3", "Body Text", "Normal"]:
        try:
            S[sn] = doc.styles[sn].style_id
        except:
            S[sn] = sn

    # Find key positions
    src_para = None
    desc_para = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and "需求来源" in p.text:
            src_para = i
            if i + 1 < len(doc.paragraphs):
                desc_para = i + 1
            break
    if src_para is None:
        raise ValueError("模板中未找到'需求来源'章节")

    children = list(body)
    src_body = next(j for j, c in enumerate(children) if c is doc.paragraphs[src_para]._element)

    # Find table after 需求来源
    tbl_body = None
    for j in range(src_body + 1, len(children)):
        if children[j].tag == qn('w:tbl'):
            tbl_body = j
            break

    # Find first 工单 Heading 1
    gd_body = None
    for j in range(tbl_body + 1 if tbl_body else src_body + 1, len(children)):
        if children[j].tag == qn('w:p'):
            for p in doc.paragraphs:
                if p._element is children[j] and p.style.name == "Heading 1" and "上线记录" in p.text:
                    gd_body = j
                    break
        if gd_body:
            break

    # Update description
    ureq = len(set(r["需求单号"] for r in data_rows))
    ugd = len(set(r["工单号"] for r in data_rows))
    trep = sum(int(r.get("报表统计次数", 0) or 0) for r in data_rows)
    nd = f"服务周期内，共有{ureq}张需求单，{ugd}张工单涉及{trep}次数据报表服务。具体需求单、工单和产出如下表："
    dp = doc.paragraphs[desc_para]
    if dp.runs:
        dp.runs[0].text = nd
    for rn in dp.runs[1:]:
        rn.text = ""
    
    # Fill table with solid borders
    tbl = children[tbl_body]
    for tr in tbl.findall(qn('w:tr'))[1:]:
        tbl.remove(tr)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for eb in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(eb)
    for ts in tblPr.findall(qn('w:tblStyle')):
        tblPr.remove(ts)
    add_solid_borders(tblPr)

    for idx, rd in enumerate(data_rows):
        tr = OxmlElement("w:tr")
        tr.append(make_cell_oxml(str(idx + 1)))
        tr.append(make_cell_oxml(rd.get("需求单号", "")))
        tr.append(make_cell_oxml(rd.get("工单号", "")))
        tr.append(make_cell_oxml(rd.get("工单内容", "")))
        tr.append(make_cell_oxml(rd.get("报表统计次数", "")))
        tbl.append(tr)

    # Total row
    trt = OxmlElement("w:tr")
    trt.append(make_cell_oxml("合计", grid_span=2, bold=True))
    trt.append(make_cell_oxml(""))
    trt.append(make_cell_oxml(""))
    trt.append(make_cell_oxml(str(trep), bold=True))
    tbl.append(trt)
    print(f"表格: {len(data_rows)} 行 + 合计")

    # Remove old 工单 content
    if gd_body:
        children = list(body)
        for j in range(len(children) - 1, gd_body - 1, -1):
            body.remove(children[j])

    # Build 工单 sections
    body.append(mp("", S["Normal"]))
    print(f"构建 {len(data_rows)} 个工单章节...")
    for rd in data_rows:
        gd = rd.get("工单内容", "")
        launch_cols = launch_image_columns_for_row(rd, cell_imgs, ["上线交付截图1", "上线交付截图2"])
        usage_cols = launch_image_columns_for_row(rd, cell_imgs, ["使用记录截图1", "使用记录截图2"])
        body.append(mp(f"{gd}的上线记录", S["Heading 1"]))
        body.append(mp("产出说明", S["Heading 2"]))
        body.append(mp(build_launch_identifier_line(rd), S["Body Text"], 480, word_wrap=True))
        body.append(mp(f"需求描述：{build_launch_requirement_description(rd)}", S["Body Text"], 480))
        body.append(mp(f"统计报表：{rd.get('报表统计次数', '')}次。", S["Normal"], 480, word_wrap=True))
        body.append(mp("上线交付截图", S["Heading 2"]))
        if launch_cols:
            for _ in launch_cols:
                body.append(mp("", S["Body Text"]))
        else:
            body.append(mp("截图待补充。", S["Body Text"], 480))
        body.append(mp("使用记录", S["Heading 2"]))
        body.append(mp("使用记录如下：", S["Body Text"]))
        if usage_cols:
            for _ in usage_cols:
                body.append(mp("", S["Body Text"]))
        else:
            body.append(mp("截图待补充。", S["Body Text"], 480))

    doc.save(output_path)

    # Insert images
    print("插入图片...")
    doc2 = Document(output_path)
    h2s = [(i, p) for i, p in enumerate(doc2.paragraphs) if p.style.name == "Heading 2"]
    gi = 0
    for pi, p in h2s:
        if p.text.strip() == "上线交付截图":
            if gi < len(data_rows):
                rn = data_rows[gi]["_row"]
                for o, cn in enumerate(launch_image_columns_for_row(data_rows[gi], cell_imgs, ["上线交付截图1", "上线交付截图2"])):
                    tpi = pi + 1 + o
                    if tpi < len(doc2.paragraphs) and (rn, cn) in cell_imgs:
                        tp = doc2.paragraphs[tpi]
                        image_payload = image_bytes_for_docx(cell_imgs[(rn, cn)])
                        if not image_payload:
                            continue
                        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                        tf.write(image_payload)
                        tf.close()
                        tp.add_run().add_picture(tf.name, width=Inches(5.5))
                        os.unlink(tf.name)
        elif p.text.strip() == "使用记录":
            if gi < len(data_rows):
                rn = data_rows[gi]["_row"]
                for o, cn in enumerate(launch_image_columns_for_row(data_rows[gi], cell_imgs, ["使用记录截图1", "使用记录截图2"])):
                    tpi = pi + 2 + o
                    if tpi < len(doc2.paragraphs) and (rn, cn) in cell_imgs:
                        tp = doc2.paragraphs[tpi]
                        image_payload = image_bytes_for_docx(cell_imgs[(rn, cn)])
                        if not image_payload:
                            continue
                        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                        tf.write(image_payload)
                        tf.close()
                        tp.add_run().add_picture(tf.name, width=Inches(5.5))
                        os.unlink(tf.name)
                gi += 1

    doc2.save(output_path)
    update_toc_via_com(output_path)
    return output_path

# ================================================================================
# 01-数据统计分析_需求文档 Fill Logic
# ================================================================================

def fill_stats_requirement_doc(excel_path, data_rows, template_path, output_path, relation_path=None):
    """填充 01-数据统计分析_需求文档"""
    print(f"数据: {len(data_rows)} 条")

    relation_descriptions = {}
    if relation_path:
        print(f"读取表融合关系文字描述: {relation_path}")
        relation_descriptions = load_stats_relation_descriptions(relation_path)
        print(f"文字描述匹配: {len(relation_descriptions)} 个结果表")

    bts = {}
    if not relation_descriptions:
        print("提取业务逻辑附件...")
        obc = extract_ole_by_cell(excel_path)
        for rd in data_rows:
            k = (rd.get('_vml_row', -1), rd.get('_vml_col', -1))
            if k in obc:
                bts[rd.get("工单内容", "")] = obc[k]
        print(f"附件匹配: {len(bts)} 个")

    doc = Document(template_path)
    body = doc.element.body

    S = {}
    for sn in ["Heading 1", "Heading 2", "Heading 3", "Body Text", "Normal"]:
        try: S[sn] = doc.styles[sn].style_id
        except: S[sn] = sn

    # Find "需求来源" description paragraph to update count
    src_para = desc_para = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 2" and "需求来源" in p.text:
            src_para = i; desc_para = i + 1 if i + 1 < len(doc.paragraphs) else None; break
    if src_para is None:
        raise ValueError("未找到需求来源")

    # Find "需求内容" H1
    content_h1 = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and "需求内容" in p.text:
            content_h1 = i; break

    # Build body children index
    children = list(body)
    src_body = next(j for j, c in enumerate(children) if c is doc.paragraphs[src_para]._element)
    tbl_body = next(j for j in range(src_body + 1, len(children)) if children[j].tag == qn('w:tbl'))
    content_body = next(j for j, c in enumerate(children) if c is doc.paragraphs[content_h1]._element)

    # Find first H2 after 需求内容 (the first old gongdan)
    gd_body = None
    for j in range(content_body + 1, len(children)):
        if children[j].tag == qn('w:p'):
            for p in doc.paragraphs:
                if p._element is children[j] and p.style.name == "Heading 2":
                    gd_body = j; break
        if gd_body is not None: break

    # Find "其他要求" H1
    other_h1_body = None
    for j in range(content_body + 1, len(children)):
        if children[j].tag == qn('w:p'):
            for p in doc.paragraphs:
                if p._element is children[j] and p.style.name == "Heading 1" and "其他要求" in p.text:
                    other_h1_body = j; break
        if other_h1_body is not None: break

    # Update count description
    ureq = len(set(r["需求单号"] for r in data_rows))
    ugd = len(set(r["工单号"] for r in data_rows))
    trep = sum(int(r.get("报表统计次数", 0) or 0) for r in data_rows)
    nd = f"服务周期内，共有{ureq}张需求单，{ugd}张工单涉及{trep}次数据统计分析。具体需求单、工单和产出如下表："
    
    dp = doc.paragraphs[desc_para]
    if dp.runs:
        dp.runs[0].text = nd
    for rn in dp.runs[1:]:
        rn.text = ""

    # Rebuild table
    tbl = children[tbl_body]
    for tr in tbl.findall(qn('w:tr'))[1:]:
        tbl.remove(tr)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for eb in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(eb)
    for ts in tblPr.findall(qn('w:tblStyle')):
        tblPr.remove(ts)
    add_solid_borders(tblPr)

    # Fixed layout + column widths 1:2:2:4:1
    tw_el = tblPr.find(qn('w:tblW'))
    if tw_el is not None:
        tw_el.set(qn('w:type'), 'pct'); tw_el.set(qn('w:w'), '5000')
    for tl in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(tl)
    tl = OxmlElement('w:tblLayout'); tl.set(qn('w:type'), 'fixed'); tblPr.append(tl)
    tg = tbl.find(qn('w:tblGrid'))
    if tg is not None:
        for gc in list(tg):
            tg.remove(gc)
        for w in ["900", "1800", "1800", "3600", "900"]:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), w); tg.append(gc)

    for idx, rd in enumerate(data_rows):
        tr = OxmlElement("w:tr")
        tr.append(make_cell_oxml(str(idx + 1)))
        tr.append(make_cell_oxml(rd.get("需求单号", "")))
        tr.append(make_cell_oxml(rd.get("工单号", "")))
        tr.append(make_cell_oxml(rd.get("工单内容", "")))
        tr.append(make_cell_oxml(rd.get("报表统计次数", "")))
        tbl.append(tr)

    trt = OxmlElement("w:tr")
    trt.append(make_cell_oxml("合计", grid_span=2, bold=True))
    trt.append(make_cell_oxml(""))
    trt.append(make_cell_oxml(""))
    trt.append(make_cell_oxml(str(trep), bold=True))
    tbl.append(trt)
    print(f"表格: {len(data_rows)} 行 + 合计")

    # Remove old gongdan content between first gongdan H2 and 其他要求
    if gd_body is not None:
        children2 = list(body)
        end_idx = other_h1_body if other_h1_body else len(children2)
        for j in range(end_idx - 1, gd_body - 1, -1):
            body.remove(children2[j])

    # Find insertion point: right before "其他要求"
    insert_before = None
    children3 = list(body)
    for j, child in enumerate(children3):
        if child.tag == qn('w:p'):
            for p in doc.paragraphs:
                if p._element is child and p.style.name == "Heading 1" and "其他要求" in p.text:
                    insert_before = child
                    break
        if insert_before is not None: break

    # Build new gongdan content
    print(f"构建 {len(data_rows)} 个工单章节...")
    new_elems = [mp("", S["Normal"])]

    # Need bold mp variant
    def mp_bold(text, style_id, indent=None):
        p = OxmlElement("w:p"); pr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"), style_id); pr.append(ps)
        if indent:
            i = OxmlElement("w:ind"); i.set(qn("w:firstLine"), str(indent)); pr.append(i)
        p.append(pr)
        r = OxmlElement("w:r"); rp = OxmlElement("w:rPr"); rp.append(OxmlElement("w:b")); r.append(rp)
        t = OxmlElement("w:t"); t.text = xml_safe(text); t.set(qn("xml:space"), "preserve")
        r.append(t); p.append(r); return p

    for rd in data_rows:
        gd = rd.get("工单内容", "")
        new_elems.append(mp(xml_safe(gd), S["Heading 2"]))
        new_elems.append(mp("业务描述", S["Heading 3"]))
        new_elems.append(mp(xml_safe(rd.get("业务描述", "")), S["Body Text"], 480))
        ri = rd.get("results") or parse_report_list(rd.get("统计分析结果表清单", "")); nr = len(ri)
        new_elems.append(mp(f"本次工作拟产出{nr}个统计分析结果表。", S["Body Text"]))
        if ri:
            rr = [[str(i + 1), xml_safe(cn), xml_safe(en)] for i, (cn, en) in enumerate(ri)]
            new_elems.append(make_stats_source_table(["序号", "结果中文表名称", "表名"], rr))
        new_elems.append(mp("", S["Body Text"]))
        new_elems.append(mp("业务逻辑", S["Heading 3"]))
        gk = rd.get("工单内容", "")
        if relation_descriptions:
            steps = build_stats_requirement_business_logic_steps(ri, relation_descriptions)
        else:
            biz = bts.get(gk, "")
            steps = summarize_biz_logic(biz) if biz else [("1", "无附件，请手动补充业务逻辑说明。")]
        new_elems.append(make_biz_table("业务逻辑说明", steps))
        new_elems.append(mp("数据加工周期", S["Heading 3"]))
        new_elems.append(mp(f"数据统计分析执行周期：{xml_safe(rd.get('数据统计分析执行周期', ''))}。", S["Body Text"]))
        new_elems.append(mp(f"数据更新要求：{xml_safe(rd.get('数据更新要求', ''))}。", S["Body Text"]))
        new_elems.append(mp(f"数据量对后续运维的特殊要求：{xml_safe(rd.get('数据量对后续运维的特殊要求', ''))}。", S["Body Text"]))

    # Insert new elements at the right position
    if insert_before is not None:
        parent = insert_before.getparent()
        idx = list(parent).index(insert_before)
        for elem in reversed(new_elems):
            parent.insert(idx, elem)
    else:
        for elem in new_elems:
            body.append(elem)

    doc.save(output_path)
    print(f"已保存: {output_path}")
    update_toc_via_com(output_path)
    return output_path


# ================================================================================
# 02-数据统计分析_设计文档 Fill Logic
# ================================================================================

def image_id_from_formula(formula_text):
    match = re.search(r'DISPIMG\("([^"]+)"', str(formula_text or ""))
    return match.group(1) if match else ""


def extract_openpyxl_images_by_cell(xlsx_path, sheet_name):
    """Return standard Excel drawing images keyed by 1-based (row, col)."""
    wb = None
    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=False)
        if sheet_name not in wb.sheetnames:
            return {}
        ws = wb[sheet_name]
        images = {}
        for image in getattr(ws, "_images", []):
            marker = getattr(getattr(image, "anchor", None), "_from", None)
            if marker is None:
                continue
            try:
                images[(marker.row + 1, marker.col + 1)] = image._data()
            except Exception:
                continue
        return images
    finally:
        if wb is not None:
            wb.close()


def load_stats_design_usage_data(relation_path):
    if not relation_path:
        raise ValueError("02-数据统计分析_设计文档需要模板同目录的 04-数据统计分析_结果表及使用说明.xlsx")

    relation_images = extract_openpyxl_images_by_cell(relation_path, "2、表融合关系")
    wb = openpyxl.load_workbook(relation_path, data_only=False)
    source_map = {}
    if "1、数据源表list" in wb.sheetnames:
        ws = wb["1、数据源表list"]
        headers = [str(ws.cell(1, col).value or "").strip() for col in range(1, ws.max_column + 1)]
        for row in range(2, ws.max_row + 1):
            rd = {header: str(ws.cell(row, idx + 1).value or "").strip() for idx, header in enumerate(headers) if header}
            result = norm_table_name(rd.get("数据融合加工表", ""))
            source = norm_table_name(rd.get("资源信息（表名）", ""))
            if result and source:
                source_map.setdefault(result, []).append(rd)

    relation_map = {}
    if "2、表融合关系" in wb.sheetnames:
        ws = wb["2、表融合关系"]
        current = ""
        for row in range(1, ws.max_row + 1):
            title = ws.cell(row, 1).value
            parsed = _parse_relation_title(title)
            if parsed:
                current = parsed
            for col in range(1, ws.max_column + 1):
                label = str(ws.cell(row, col).value or "").strip()
                if label == "文字描述" and current:
                    relation_map.setdefault(current, {})["description"] = sanitize_stats_logic_text(ws.cell(row, col + 1).value)
                    break
                if label == "数据处理流程图" and current:
                    entry = relation_map.setdefault(current, {})
                    entry["image_formula"] = str(ws.cell(row, col + 1).value or "").strip()
                    if (row, col + 1) in relation_images:
                        entry["image_bytes"] = relation_images[(row, col + 1)]
                    break

    detail_map = {}
    if "4、数据统计分析结果表详情" in wb.sheetnames:
        ws = wb["4、数据统计分析结果表详情"]
        current = ""
        row = 1
        while row <= ws.max_row:
            value_b = str(ws.cell(row, 2).value or "").strip()
            if norm_table_name(value_b).startswith("FUSION_"):
                current = norm_table_name(value_b)
                row += 1
                continue
            headers_here = [str(ws.cell(row, col).value or "").strip() for col in range(1, ws.max_column + 1)]
            if current and "字段中文名" in headers_here and "字段英文名" in headers_here:
                header_map = {header: idx + 1 for idx, header in enumerate(headers_here) if header}
                row += 1
                fields = []
                while row <= ws.max_row:
                    first = str(ws.cell(row, 1).value or "").strip()
                    second = str(ws.cell(row, 2).value or "").strip()
                    if re.match(r"4\.\d+", first) or norm_table_name(second).startswith("FUSION_"):
                        row -= 1
                        break
                    if any(str(ws.cell(row, col).value or "").strip() for col in range(1, ws.max_column + 1)):
                        fields.append({header: str(ws.cell(row, col).value or "").strip() for header, col in header_map.items()})
                    row += 1
                detail_map[current] = fields
            row += 1
    wb.close()

    return source_map, relation_map, detail_map, extract_images_via_cellimages(relation_path)


def load_stats_design_catalog(catalog_path, needed_names=None):
    wb = openpyxl.load_workbook(catalog_path, data_only=True, read_only=True)
    if "关联资源信息" not in wb.sheetnames:
        wb.close()
        return {}
    ws = wb["关联资源信息"]
    headers = [str(ws.cell(1, col).value or "").strip() for col in range(1, ws.max_column + 1)]
    rows = {}
    for values in ws.iter_rows(min_row=2, values_only=True):
        rd = {header: str(values[idx] or "").strip() for idx, header in enumerate(headers) if header and idx < len(values)}
        code = norm_table_name(rd.get("资源编码", ""))
        if needed_names and code not in needed_names:
            continue
        if code and code not in rows:
            rows[code] = rd
    wb.close()
    return rows


def make_stats_design_entity_row(table_name, source_rd, resource_info, table_type, seq):
    key = norm_table_name(table_name)
    catalog_rd = resource_info.get(key, {})
    if source_rd:
        directory_code = source_rd.get("资源编目（非必填）") or catalog_rd.get("数据目录代码", "")
        resource_name = source_rd.get("资源名称") or catalog_rd.get("资源名称", "")
        source_name = source_rd.get("资源信息（表名）") or key
    else:
        directory_code = catalog_rd.get("数据目录代码", "")
        resource_name = catalog_rd.get("资源名称", "")
        source_name = key
    return [
        str(seq),
        directory_code,
        resource_name,
        source_name,
        table_type,
        catalog_rd.get("业务数据更新周期", ""),
    ]


def find_heading_paragraph(doc, style_name, contains):
    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_name and contains in paragraph.text:
            return paragraph
    return None


def remove_content_after_heading(body, heading_paragraph):
    found = False
    for child in list(body):
        if child is heading_paragraph._element:
            found = True
            continue
        if found and child.tag != qn("w:sectPr"):
            body.remove(child)


def append_body_element(body, element):
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
            break
    if sect_pr is not None:
        body.insert(list(body).index(sect_pr), element)
    else:
        body.append(element)


def append_image_paragraph(doc, body, image_bytes, missing_text):
    paragraph = doc.add_paragraph("")
    append_body_element(body, paragraph._element)
    if image_bytes:
        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tf.write(image_bytes)
        tf.close()
        paragraph.add_run().add_picture(tf.name, width=Inches(5.8))
        os.unlink(tf.name)
    else:
        paragraph.text = missing_text
    return paragraph


def first_non_empty(mapping, keys, default=""):
    for key in keys:
        value = str(mapping.get(key, "") or "").strip()
        if value:
            return value
    return default


def fill_stats_design_doc(excel_path, data_rows, template_path, output_path, catalog_path, relation_path=None):
    """填充 02-数据统计分析_设计文档，只生成 Word 文档。"""
    if not catalog_path:
        raise ValueError("02-数据统计分析_设计文档需要 --catalog 参数指定数据目录数据路径")
    if not os.path.exists(catalog_path):
        raise FileNotFoundError(f"数据目录数据文件不存在: {catalog_path}")
    if not relation_path:
        relation_path = resolve_stats_relation_workbook(template_path)
    if not relation_path:
        raise FileNotFoundError("未找到 04-数据统计分析_结果表及使用说明.xlsx，请放在模板同目录或通过 --catalog 传入")

    print(f"数据: {len(data_rows)} 个工单")
    programs = []
    for group in data_rows:
        for result_cn, result_en in group.get("results") or parse_report_list(group.get("统计分析结果表清单", "")):
            if result_en:
                item = dict(group)
                item["result_cn"] = result_cn
                item["result_en"] = norm_table_name(result_en)
                programs.append(item)
    print(f"程序: {len(programs)} 个")

    source_map, relation_map, detail_map, image_bytes = load_stats_design_usage_data(relation_path)
    needed_names = {item["result_en"] for item in programs}
    for result_en in list(needed_names):
        for source in source_map.get(result_en, []):
            source_name = norm_table_name(source.get("资源信息（表名）", ""))
            if source_name:
                needed_names.add(source_name)
    resource_info = load_stats_design_catalog(catalog_path, needed_names)

    doc = Document(template_path)
    body = doc.element.body
    design_h1 = find_heading_paragraph(doc, "Heading 1", "数据统计分析设计")
    if design_h1 is None:
        raise ValueError("模板中未找到「数据统计分析设计」章节")

    source_heading = find_heading_paragraph(doc, "Heading 2", "需求来源")
    if source_heading:
        source_index = next(
            (
                idx for idx, paragraph in enumerate(doc.paragraphs)
                if paragraph.style.name == "Heading 2" and "需求来源" in paragraph.text
            ),
            -1,
        )
        if source_index + 1 < len(doc.paragraphs):
            desc = doc.paragraphs[source_index + 1]
            unique_reqs = len({item.get("需求单号", "") for item in programs if item.get("需求单号", "")})
            unique_orders = len({item.get("工单号", "") for item in programs if item.get("工单号", "")})
            desc.text = f"服务周期内，共有{unique_reqs}张需求单，{unique_orders}张工单涉及{len(programs)}次数据统计分析服务。具体需求单、工单和产出如下表："
        if len(doc.tables) >= 3:
            tbl = doc.tables[2]._tbl
            for tr in tbl.findall(qn("w:tr"))[1:]:
                tbl.remove(tr)
            tbl_pr = tbl.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl.insert(0, tbl_pr)
            for border in tbl_pr.findall(qn("w:tblBorders")):
                tbl_pr.remove(border)
            add_solid_borders(tbl_pr)
            for idx, item in enumerate(programs, start=1):
                tr = OxmlElement("w:tr")
                for value in [
                    str(idx),
                    item.get("需求单号", ""),
                    item.get("工单号", ""),
                    item.get("工单内容", ""),
                    item["result_en"],
                ]:
                    tr.append(make_cell_oxml(xml_safe(value)))
                tbl.append(tr)

    remove_content_after_heading(body, design_h1)

    style = {}
    for local_name, style_name in [("H2", "Heading 2"), ("H3", "Heading 3"), ("H4", "Heading 4"), ("BT", "Body Text"), ("NL", "Normal")]:
        try:
            style[local_name] = doc.styles[style_name].style_id
        except Exception:
            style[local_name] = style_name

    append_body_element(body, mp(
        f"服务单服务周期内，形成了{len(programs)}个数据统计分析程序，各程序分析设计和融合加工的具体过程如下：",
        style["NL"],
    ))

    for item in programs:
        result_cn = item["result_cn"]
        result_en = item["result_en"]
        append_body_element(body, mp(xml_safe(result_cn), style["H2"]))

        append_body_element(body, mp("涉及到的实体表", style["H3"]))
        entity_rows = []
        seq = 1
        seen_sources = set()
        for source in source_map.get(result_en, []):
            table_name = source.get("资源信息（表名）", "")
            key = norm_table_name(table_name)
            if not key or key in seen_sources:
                continue
            seen_sources.add(key)
            entity_rows.append(make_stats_design_entity_row(table_name, source, resource_info, "源表", seq))
            seq += 1
        entity_rows.append(make_stats_design_entity_row(result_en, None, resource_info, "目标表", seq))
        append_body_element(body, make_table_oxml(
            ["序号", "数据目录/编码（如有）", "数据目录中文名称（目录名）", "表名", "表类型", "数据更新周期"],
            entity_rows,
            ["500", "1350", "3400", "2450", "800", "1250"],
        ))

        append_body_element(body, mp("数据统计分析设计", style["H3"]))
        append_body_element(body, mp(xml_safe(f"{result_cn} {result_en}"), style["H4"]))
        fields = []
        for field in detail_map.get(result_en, []):
            fields.append([
                field.get("字段中文名", ""),
                field.get("字段英文名", ""),
                field.get("字段类型", ""),
                first_non_empty(field, ["是否为空", "不可为空"]),
                first_non_empty(field, ["主键/外键", "唯一"]),
                field.get("字段注释", "") or "No",
            ])
        if not fields:
            fields = [["未匹配到字段信息", "", "", "", "", "请补充"]]
        append_body_element(body, make_table_oxml(
            ["字段中文名", "字段英文名", "字段类型", "是否为空", "主键/外键", "字段说明"],
            fields,
            ["1700", "1800", "1500", "1100", "1100", "2800"],
        ))

        append_body_element(body, mp("数据处理流程图", style["H3"]))
        relation_data = relation_map.get(result_en, {})
        image_payload = relation_data.get("image_bytes")
        if not image_payload:
            img_id = image_id_from_formula(relation_data.get("image_formula", ""))
            image_payload = image_bytes.get(img_id)
        append_image_paragraph(doc, body, image_payload, "未匹配到数据处理流程图，请补充。")

        append_body_element(body, mp("数据加工逻辑", style["H3"]))
        steps = split_logic_description(relation_map.get(result_en, {}).get("description", ""))
        if not steps:
            steps = ["未匹配到数据加工逻辑说明，请手动补充。"]
        append_body_element(body, make_biz_table(
            "数据统计分析加工逻辑说明",
            [(str(idx), step) for idx, step in enumerate(steps, start=1)],
        ))

    doc.save(output_path)
    print(f"已保存: {output_path}")
    update_toc_via_com(output_path)
    return output_path

# Dispatcher
# ══════════════════════════════════════════════════════════════

def fill_stats_result_usage_workbook(excel_path, service_dir, template_path, output_path, catalog_path):
    """Fill 04-数据统计分析_结果表及使用说明 workbook template."""
    if not catalog_path:
        raise ValueError("04-数据统计分析_结果表及使用说明需要 --catalog 参数指定数据目录数据路径")
    if not os.path.exists(catalog_path):
        raise FileNotFoundError(f"数据目录数据文件不存在: {catalog_path}")
    ensure_module("PIL", "pillow")
    builder_path = os.path.join(os.path.dirname(__file__), "build_stats_result_usage_workbook.py")
    spec = importlib.util.spec_from_file_location("build_stats_result_usage_workbook", builder_path)
    module = importlib.util.module_from_spec(spec)
    if spec.loader is None:
        raise RuntimeError(f"无法加载生成脚本: {builder_path}")
    spec.loader.exec_module(module)
    return module.build_stats_result_usage_workbook(
        ledger_path=excel_path,
        service_dir=service_dir,
        template_path=template_path,
        output_path=output_path,
        catalog_path=catalog_path,
    )


def fill_stats_test_pdf(excel_path, service_dir, template_path, output_path):
    """Fill 03-数据统计分析_测试文档 PDF."""
    ensure_module("PIL", "pillow")
    ensure_module("fitz", "pymupdf")
    ensure_module("reportlab")
    builder_path = os.path.join(os.path.dirname(__file__), "build_stats_test_pdf.py")
    spec = importlib.util.spec_from_file_location("build_stats_test_pdf", builder_path)
    module = importlib.util.module_from_spec(spec)
    if spec.loader is None:
        raise RuntimeError(f"无法加载生成脚本: {builder_path}")
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module.build_stats_test_pdf(
        ledger_path=excel_path,
        service_dir=service_dir,
        template_path=template_path,
        output_path=output_path,
    )


def fill_stats_test_docx(excel_path, service_dir, template_path, output_path):
    """Fill 03-数据统计分析_测试文档 Word document."""
    ensure_module("docx", "python-docx")
    ensure_module("PIL", "pillow")
    ensure_module("olefile")
    builder_path = os.path.join(os.path.dirname(__file__), "build_stats_test_docx.py")
    spec = importlib.util.spec_from_file_location("build_stats_test_docx", builder_path)
    module = importlib.util.module_from_spec(spec)
    if spec.loader is None:
        raise RuntimeError(f"无法加载生成脚本: {builder_path}")
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module.build_stats_test_docx(
        ledger_path=excel_path,
        service_dir=service_dir,
        template_path=template_path,
        output_path=output_path,
    )


def fill_document(excel_path, service_dir, material_type, template_path, output_path, catalog_path=None):
    output_path = resolve_output_path(output_path, material_type)
    print(f"台账清单: {excel_path}")
    print(f"筛选条件: 服务目录={service_dir}")
    print(f"材料类型: {material_type}")

    if material_type == "01-数据报表_需求文档":
        data_rows = read_excel(excel_path, service_dir)
        data_rows = attach_delivery_files_for_requirement(excel_path, data_rows)
        catalog_context = build_data_report_catalog_context(catalog_path, data_rows) if catalog_path else {}
        return fill_requirement_doc(data_rows, template_path, output_path, catalog_context=catalog_context)
    elif material_type == "02-数据报表_设计文档":
        data_rows = read_data_report_design_groups(excel_path, service_dir)
        if not catalog_path:
            raise ValueError("02-设计文档需要 --catalog 参数指定数据目录数据路径")
        if not os.path.exists(catalog_path):
            raise FileNotFoundError(f"数据目录数据文件不存在: {catalog_path}")
        return fill_design_doc_full(excel_path, data_rows, template_path, output_path, catalog_path)
    elif material_type == "01-数据统计分析_需求文档":
        data_rows = read_stats_requirement_groups(excel_path, service_dir)
        relation_path = resolve_stats_relation_workbook(template_path, catalog_path, output_path=output_path)
        return fill_stats_requirement_doc(excel_path, data_rows, template_path, output_path, relation_path)
    elif material_type == "02-数据统计分析_设计文档":
        if not catalog_path:
            raise ValueError("02-数据统计分析_设计文档需要 --catalog 参数指定数据目录数据路径")
        if not os.path.exists(catalog_path):
            raise FileNotFoundError(f"数据目录数据文件不存在: {catalog_path}")
        data_rows = read_stats_requirement_groups(excel_path, service_dir)
        relation_path = resolve_stats_relation_workbook(template_path, output_path=output_path)
        return fill_stats_design_doc(excel_path, data_rows, template_path, output_path, catalog_path, relation_path)
    elif material_type == "04-数据统计分析_结果表及使用说明":
        return fill_stats_result_usage_workbook(excel_path, service_dir, template_path, output_path, catalog_path)
    elif material_type == "03-数据统计分析_测试文档":
        return fill_stats_test_docx(excel_path, service_dir, template_path, output_path)
    elif material_type == "03-数据报表_上线记录":
        data_rows = read_excel(excel_path, service_dir)
        return fill_launch_record_doc(excel_path, data_rows, template_path, output_path)
    else:
        raise ValueError(f"不支持的材料类型: {material_type}。当前支持: 01-数据报表_需求文档, 01-数据统计分析_需求文档, 02-数据报表_设计文档, 02-数据统计分析_设计文档, 03-数据报表_上线记录, 03-数据统计分析_测试文档, 04-数据统计分析_结果表及使用说明")


# ══════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="验收文档自动填充工具")
    parser.add_argument("--service-dir", required=True, help="服务目录筛选值，如 N08-数据报表服务")
    parser.add_argument("--material-type", required=True, help="材料类型，如 01-数据报表_需求文档 或 02-数据报表_设计文档")
    parser.add_argument("--excel", required=True, help="台账清单Excel文件路径")
    parser.add_argument("--template", required=True, help="Word模板文件路径")
    parser.add_argument("--output", required=True, help="输出Word文档路径")
    parser.add_argument("--catalog", default=None, help="数据目录数据Excel路径（02-设计文档必填）")
    args = parser.parse_args()

    try:
        result = fill_document(
            excel_path=args.excel,
            service_dir=args.service_dir,
            material_type=args.material_type,
            template_path=args.template,
            output_path=args.output,
            catalog_path=args.catalog,
        )
        print(f"\n✅ 完成: {result}")
    except Exception as e:
        print(f"\n❌ 错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()


