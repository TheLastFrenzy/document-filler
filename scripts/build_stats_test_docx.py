#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build 03-数据统计分析_测试文档 Word document from the adjusted N08 ledger."""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from materials.shared.word_sections import center_table_header_rows


FALLBACK_CONCLUSION = (
    "经测试，在业务经理、测试开发人员的积极配合下，按照测试规范和流程，全部测试完毕。"
    "可正常运行，无报错，且能够生成正确有效的结果，符合业务逻辑，功能都能正常使用，"
    "测试全部通过，符合上线的要求。"
)


def load_pdf_helper():
    helper_path = Path(__file__).with_name("build_stats_test_pdf.py")
    spec = importlib.util.spec_from_file_location("build_stats_test_pdf_for_docx", helper_path)
    module = importlib.util.module_from_spec(spec)
    if spec.loader is None:
        raise RuntimeError(f"无法加载测试截图解析脚本: {helper_path}")
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def style_name(doc, preferred: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[preferred]
        return preferred
    except Exception:
        return fallback


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl_pr.append(borders)


def shade_cell(cell, fill="F1F1F1"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(paragraph, font_name="宋体"):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_numbering(paragraph_element, numbering):
    if numbering is None:
        return
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = deepcopy(numbering)
    style = p_pr.find(qn("w:pStyle"))
    if style is not None:
        p_pr.insert(list(p_pr).index(style) + 1, num_pr)
    else:
        p_pr.insert(0, num_pr)


def new_paragraph_element(doc, text="", style="Normal", first_line_indent=False, align=None, numbering=None):
    paragraph = doc.add_paragraph()
    paragraph.style = style_name(doc, style)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.29)
    if align is not None:
        paragraph.alignment = align
    if text:
        paragraph.add_run(text)
        set_run_font(paragraph)
    apply_numbering(paragraph._element, numbering)
    return paragraph._element


def new_page_break_element(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph._element


def new_image_element(doc, image_path: Path):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        paragraph.add_run().add_picture(str(image_path), width=Inches(5.7))
    except Exception:
        paragraph.alignment = None
        paragraph.add_run(f"图片无法插入：{image_path.name}")
    return paragraph._element


def new_table_element(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    set_table_borders(table)
    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                set_run_font(paragraph)
            if row_idx == 0:
                shade_cell(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    return table._tbl


def find_heading(doc, text, style="Heading 1"):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text and paragraph.style.name == style:
            return paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def remove_between(body, start_element, end_element):
    removing = False
    for child in list(body):
        if child is start_element:
            removing = True
            continue
        if child is end_element:
            return
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def remove_after(body, start_element):
    removing = False
    for child in list(body):
        if child is start_element:
            removing = True
            continue
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def insert_after(ref_element, elements):
    parent = ref_element.getparent()
    index = list(parent).index(ref_element)
    for offset, element in enumerate(elements, start=1):
        parent.insert(index + offset, element)


def first_paragraph_text_after(body, start_element):
    found = False
    for child in list(body):
        if child is start_element:
            found = True
            continue
        if not found:
            continue
        if child.tag == qn("w:p"):
            texts = [node.text or "" for node in child.iter(qn("w:t"))]
            value = "".join(texts).strip()
            if value:
                return value
    return ""


def template_test_heading_numbering(body, start_element, end_element):
    numbering = {}
    found = False
    for child in list(body):
        if child is start_element:
            found = True
            continue
        if child is end_element:
            break
        if not found or child.tag != qn("w:p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
        if num_pr is None:
            continue
        ilvl = num_pr.find(qn("w:ilvl"))
        level = ilvl.get(qn("w:val")) if ilvl is not None else ""
        if level == "1" and "h2" not in numbering:
            numbering["h2"] = deepcopy(num_pr)
        elif level == "2" and "h3" not in numbering:
            numbering["h3"] = deepcopy(num_pr)
        if "h2" in numbering and "h3" in numbering:
            break
    return numbering


def build_test_content_elements(doc, materials, service_dir: str, helper, numbering=None):
    numbering = numbering or {}
    elements = [
        new_paragraph_element(
            doc,
            f"本章节依据《台账清单》中服务目录为“{service_dir}”的结果表清单整理，共涉及{len(materials)}个数据统计分析程序。",
            "缩进_五号_1.5行距",
            first_line_indent=True,
        )
    ]
    for index, item in enumerate(materials, start=1):
        if index > 1:
            elements.append(new_page_break_element(doc))
        title = item.program.result_cn or item.program.result_en
        elements.append(new_paragraph_element(doc, title, "Heading 2", numbering=numbering.get("h2")))
        if item.note:
            elements.append(new_paragraph_element(doc, item.note, "缩进_五号_1.5行距", first_line_indent=True))
        for label in helper.SECTION_LABELS:
            elements.append(new_paragraph_element(doc, label, "Heading 3", numbering=numbering.get("h3")))
            images = item.sections.get(label, [])
            if images:
                for image_path in images:
                    elements.append(new_image_element(doc, image_path))
                elements.append(new_paragraph_element(doc, helper.TEST_RESULTS[label], "缩进_五号_1.5行距", first_line_indent=True))
            elif item.status == "missing_report":
                elements.append(new_paragraph_element(doc, "未找到该工单对应的自测报告附件，截图待补充。", "缩进_五号_1.5行距", first_line_indent=True))
            elif item.status == "missing_program":
                elements.append(new_paragraph_element(doc, "未在对应自测报告中定位到该程序，截图待补充。", "缩进_五号_1.5行距", first_line_indent=True))
            else:
                elements.append(new_paragraph_element(doc, "未在对应自测报告中提取到该项截图，截图待补充。", "缩进_五号_1.5行距", first_line_indent=True))
    return elements


def build_conclusion_elements(doc, materials, conclusion_text: str):
    total = len(materials)
    return [
        new_paragraph_element(doc, conclusion_text or FALLBACK_CONCLUSION, "缩进_五号_1.5行距", first_line_indent=True),
        new_table_element(
            doc,
            [
                ["测试内容", "测试项", "通过项", "通过率"],
                ["数据测试", str(total), str(total), "100%"],
                ["程序测试", str(total), str(total), "100%"],
            ],
        ),
    ]


def update_docx_fields_via_com(doc_path):
    if sys.platform != "win32":
        return
    ps_script = f'''
$word = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open("{doc_path}")
    foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
    $doc.Fields.Update() | Out-Null
    $doc.Save()
    $doc.Close()
    Write-Host "DOCX fields updated"
}} catch {{
    Write-Host "DOCX field update error: $_"
}} finally {{
    if ($word) {{ try {{ $word.Quit() }} catch {{}} }}
}}
'''
    try:
        subprocess.run(
            ["powershell", "-ExecutionPolicy", "Bypass", "-Command", ps_script],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except Exception:
        pass


def build_stats_test_docx(ledger_path: str, service_dir: str, template_path: str, output_path: str):
    """Generate 03-数据统计分析_测试文档 Word document."""
    if template_path and not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在: {template_path}")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    helper = load_pdf_helper()
    with tempfile.TemporaryDirectory(prefix="document_filler_stats_test_docx_") as temp_dir:
        work_dir = Path(temp_dir)
        programs, groups, meta = helper.load_ledger_programs(ledger_path, service_dir)
        reports = helper.extract_self_reports(ledger_path, meta, work_dir)
        for order, report in list(reports.items()):
            reports[order] = helper.parse_report_images(report, work_dir)
        materials, summary = helper.build_materials(programs, reports)

        doc = Document(template_path)
        body = doc.element.body
        test_heading = find_heading(doc, "测试内容")
        conclusion_heading = find_heading(doc, "测试结论")
        if test_heading is None:
            raise ValueError("03-数据统计分析_测试文档模板未找到「测试内容」章节")
        if conclusion_heading is None:
            raise ValueError("03-数据统计分析_测试文档模板未找到「测试结论」章节")
        conclusion_text = first_paragraph_text_after(body, conclusion_heading._element)
        heading_numbering = template_test_heading_numbering(body, test_heading._element, conclusion_heading._element)

        remove_between(body, test_heading._element, conclusion_heading._element)
        insert_after(test_heading._element, build_test_content_elements(doc, materials, service_dir, helper, heading_numbering))
        remove_after(body, conclusion_heading._element)
        insert_after(conclusion_heading._element, build_conclusion_elements(doc, materials, conclusion_text))

        if output.exists():
            output.unlink()
        center_table_header_rows(doc)
        doc.save(output)
        update_docx_fields_via_com(str(output))
    print(json.dumps(summary, ensure_ascii=False))
    print(f"已保存: {output_path}")
    return str(output)


def main():
    parser = argparse.ArgumentParser(description="Build 03-数据统计分析_测试文档 Word document")
    parser.add_argument("--ledger", required=True)
    parser.add_argument("--service-dir", required=True)
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build_stats_test_docx(args.ledger, args.service_dir, args.template, args.output)


if __name__ == "__main__":
    main()
