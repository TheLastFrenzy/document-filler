import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))


class SharedOfficeWordTest(unittest.TestCase):
    def test_docx_template_is_returned_without_conversion(self):
        from materials.shared.office_word import convert_legacy_doc_template

        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.docx"
            template.write_bytes(b"docx")

            with mock.patch("materials.shared.office_word.subprocess.run") as run:
                result = convert_legacy_doc_template(template, Path(temp_dir) / "work")

        self.assertEqual(result, template)
        run.assert_not_called()

    def test_unsupported_template_extension_keeps_existing_error_contract(self):
        from materials.shared.office_word import convert_legacy_doc_template

        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.pdf"

            with self.assertRaisesRegex(
                ValueError,
                r"API接口测试报告模板仅支持 \.doc 或 \.docx",
            ):
                convert_legacy_doc_template(template, Path(temp_dir) / "work")

    def test_legacy_template_conversion_creates_work_directory_and_checks_output(self):
        from materials.shared.office_word import convert_legacy_doc_template

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template = temp / "template.doc"
            template.write_bytes(b"legacy")
            work_dir = temp / "missing" / "work"

            def fake_run(*_args, **_kwargs):
                converted = work_dir / "template.docx"
                self.assertTrue(converted.parent.exists())
                converted.write_bytes(b"converted")
                return SimpleNamespace(stdout="CONVERTED", stderr="", returncode=0)

            with mock.patch("materials.shared.office_word.subprocess.run", side_effect=fake_run):
                result = convert_legacy_doc_template(template, work_dir)

        self.assertEqual(result, work_dir / "template.docx")

    def test_docx_save_uses_injected_toc_updater(self):
        from materials.shared.office_word import save_word_document

        class FakeDocument:
            def save(self, path):
                Path(path).write_bytes(b"docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "result" / "output.docx"
            toc_calls = []
            result = save_word_document(
                FakeDocument(),
                output,
                Path(temp_dir) / "work",
                toc_updater=lambda path: toc_calls.append(Path(path)),
            )

        self.assertEqual(result, str(output))
        self.assertEqual(toc_calls, [output])

    def test_legacy_save_uses_injected_converter(self):
        from materials.shared.office_word import save_word_document

        class FakeDocument:
            def save(self, path):
                Path(path).write_bytes(b"docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            output = temp / "result" / "output.doc"
            work_dir = temp / "missing" / "work"
            converter_calls = []

            def converter(docx_path, output_path):
                converter_calls.append((Path(docx_path), Path(output_path)))
                Path(output_path).write_bytes(b"legacy")
                return Path(output_path)

            result = save_word_document(
                FakeDocument(),
                output,
                work_dir,
                legacy_converter=converter,
            )

        self.assertEqual(result, str(output))
        self.assertEqual(converter_calls, [(work_dir / "output.docx", output)])

    def test_existing_private_compatibility_entrypoints_remain_available(self):
        from materials.n07 import api_launch_record, api_test_report

        self.assertTrue(callable(api_launch_record._convert_docx_to_legacy_doc))
        self.assertTrue(callable(api_launch_record._save_document))
        self.assertTrue(callable(api_test_report._convert_legacy_doc_template))

    def test_save_word_document_centers_table_header_rows_before_saving(self):
        from materials.shared.office_word import save_word_document

        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "out.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    cell.text = f"{row_index}-{col_index}"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            save_word_document(
                document,
                output,
                Path(temp_dir) / "work",
                toc_updater=lambda _path: None,
            )

            result = Document(output)
            self.assertEqual(
                result.tables[0].rows[0].cells[0].paragraphs[0].alignment,
                WD_ALIGN_PARAGRAPH.CENTER,
            )
            self.assertEqual(
                result.tables[0].rows[1].cells[0].paragraphs[0].alignment,
                WD_ALIGN_PARAGRAPH.LEFT,
            )

    def test_clone_table_with_data_centers_generated_header_row(self):
        from materials.shared.word_sections import clone_table_with_data

        document = Document()
        prototype = document.add_table(rows=2, cols=2)
        prototype.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = clone_table_with_data(
            prototype._tbl,
            headers=["header 1", "header 2"],
            rows=[["value 1", "value 2"]],
        )

        header_cell = table.findall(qn("w:tr"))[0].findall(qn("w:tc"))[0]
        alignment = header_cell.find(".//" + qn("w:jc"))
        self.assertIsNotNone(alignment)
        self.assertEqual(alignment.get(qn("w:val")), "center")


if __name__ == "__main__":
    unittest.main()
