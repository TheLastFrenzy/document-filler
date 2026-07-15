import base64
import importlib.util
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

import openpyxl
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
FILL_SCRIPT = SCRIPTS / "fill_document.py"
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document_api_test", FILL_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_api_ledger(path: Path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(
        [
            "服务目录",
            "需求单号",
            "工单号",
            "工单标题",
            "程序数",
            "工单描述",
            "结果表清单",
            "自测报告附件",
        ]
    )
    sheet.append(
        [
            "N07-API接口开发",
            "REQ-1",
            "WO-1",
            "市公安局-出入境证件身份认证-共享接口",
            3,
            "升级离境退税掌上办平台的出入境记录校验能力。",
            "境外人员获取令牌接口 token_api",
            None,
        ]
    )
    sheet.append([None, None, None, None, None, None, "境外人员身份认证申请接口 apply_api", None])
    sheet.append([None, None, None, None, None, None, "境外人员身份认证请求接口 auth_api", None])
    for column in "ABCDEFH":
        sheet.merge_cells(f"{column}2:{column}4")
    workbook.save(path)
    return path


def add_parameter_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    if widths:
        for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
            grid_column.set(qn("w:w"), str(width))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    return table


def add_numbering(paragraph, num_id):
    properties = paragraph._p.get_or_add_pPr()
    existing = properties.find(qn("w:numPr"))
    if existing is not None:
        properties.remove(existing)
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), str(num_id))
    numbering.extend([level, numbering_id])
    properties.append(numbering)


def set_paragraph_format(paragraph, *, font="宋体", size=12, first_line=None, line_spacing=None):
    if first_line is not None:
        paragraph.paragraph_format.first_line_indent = Twips(first_line)
    if line_spacing is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(line_spacing))
        spacing.set(qn("w:lineRule"), "auto")
        paragraph._p.get_or_add_pPr().append(spacing)
    for run in paragraph.runs:
        run.font.name = font
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
        run.font.size = Pt(size)


def set_table_format(table, header_fill):
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            shading = cell_properties.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                cell_properties.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:fill"), header_fill if row_index == 0 else "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9 if row_index == 0 else 10.5)
                    run.font.bold = row_index == 0
                    run.font.color.rgb = RGBColor(0, 0, 0)


def make_self_report(path: Path):
    document = Document()
    document.add_heading("测试目的", level=1)
    document.add_paragraph("验证接口输入输出参数。")
    document.add_heading("测试内容", level=1)

    document.add_heading("境外人员获取令牌接口", level=2)
    document.add_paragraph("输入参数")
    add_parameter_table(
        document,
        ["序号", "参数项", "名称"],
        [["1", "userId", "用户ID"], ["2", "password", "密码"]],
    )
    document.add_paragraph("输出参数")
    add_parameter_table(
        document,
        ["参数", "参数说明", "备注"],
        [["access_token", "Token", "后续接口调用凭证"]],
    )
    document.add_paragraph("测试结果：符合规范要求，测试通过。")

    document.add_heading("境外人员身份认证申请接口", level=2)
    document.add_paragraph("共享接口命名规范性")
    document.add_paragraph("输入参数")
    document.add_paragraph("请求头：")
    add_parameter_table(
        document,
        ["参数", "参数说明", "是否必选"],
        [["access_token", "访问令牌", "是"]],
    )
    document.add_paragraph("请求体：")
    add_parameter_table(
        document,
        ["参数", "参数说明", "是否必选", "类型"],
        [["custNum", "业务站点号", "是", "String"]],
        [1400, 3600, 1200, 1800],
    )
    document.add_paragraph("测试结果：符合规范要求，测试通过。")
    document.add_paragraph("输出参数")
    add_parameter_table(
        document,
        ["参数", "参数说明", "备注"],
        [["bizSerialNum", "业务流水号", "申请成功后返回"]],
    )
    document.add_paragraph("测试结果：符合规范要求，测试通过。")
    document.add_paragraph("共享任务开发代码检查")

    document.add_heading("境外人员身份认证请求接口", level=2)
    document.add_paragraph("输入参数")
    document.add_paragraph("请求头：")
    add_parameter_table(
        document,
        ["参数", "参数说明", "是否必选"],
        [["access_token", "访问令牌", "是"]],
    )
    document.add_paragraph("请求体：")
    add_parameter_table(
        document,
        ["参数", "参数说明", "是否必选", "类型"],
        [["bizSerialNum", "业务流水号", "是", "String"], ["idAuthData", "身份验证数据", "是", "Object"]],
    )
    document.add_paragraph("输出参数")
    add_parameter_table(
        document,
        ["参数", "参数说明", "备注"],
        [["success", "是否成功", ""], ["result", "核验结果", "四位结果数组"]],
    )
    document.add_paragraph("测试结果：符合规范要求，测试通过。")
    document.save(path)
    return path


def add_picture_paragraph(document, image_path: Path):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(1))
    return paragraph


def make_self_report_with_code_images(path: Path):
    image_path = path.with_name("screenshot.png")
    image_path.write_bytes(PNG_1X1)
    document = Document()
    document.add_heading("测试目的", level=1)
    document.add_paragraph("验证接口代码规范性。")
    document.add_heading("4.测试内容", level=1)
    for name in [
        "境外人员获取令牌接口",
        "境外人员身份认证申请接口",
        "境外人员身份认证请求接口",
    ]:
        document.add_heading(name, level=2)
        document.add_paragraph("共享接口命名规范性")
        add_picture_paragraph(document, image_path)
        document.add_paragraph("测试结果：符合规范要求，测试通过。")
        document.add_paragraph("输入参数")
        add_parameter_table(document, ["序号", "参数项", "名称"], [["1", "userId", "用户ID"]])
        document.add_paragraph("输出参数")
        add_parameter_table(document, ["参数", "参数说明"], [["success", "是否成功"]])
        document.add_paragraph("共享任务开发代码检查")
        add_picture_paragraph(document, image_path)
        document.add_paragraph("测试结果：符合规范要求，测试通过。")
        document.add_paragraph("测试结果")
        add_picture_paragraph(document, image_path)
    document.add_heading("测试结论", level=1)
    document.add_paragraph("全部自测通过。")
    document.save(path)
    return path


def make_api_template(path: Path):
    document = Document()
    document.add_heading("需求文档", level=1)
    revision = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["版本", "更新人员", "更新内容", "时间"]):
        revision.rows[0].cells[index].text = value
    document.add_heading("业务场景", level=2)
    document.add_paragraph("上海市大数据中心通过API接口开展公共数据共享开放。")
    business_scene = document.add_paragraph("旧业务场景内容。")
    set_paragraph_format(business_scene, first_line=480, line_spacing=360)
    document.add_paragraph("BUSINESS_STATIC_AFTER")
    document.add_heading("需求说明", level=2)
    document.add_heading("需求清单", level=2)
    demand_summary = document.add_paragraph("服务周期内，共有3张需求单。")
    set_paragraph_format(demand_summary, size=11)
    demand = document.add_table(rows=3, cols=5)
    for index, value in enumerate(["序号", "对应需求单编号", "对应工单编号", "工单内容", "涉及共享任务数量（个）"]):
        demand.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "REQ-OLD", "WO-OLD", "旧工单", "1"]):
        demand.rows[1].cells[index].text = value
    demand.rows[2].cells[0].merge(demand.rows[2].cells[1]).text = "总计"
    demand.rows[2].cells[4].text = "1"
    set_table_format(demand, "B4C6E7")
    work_order_heading = document.add_heading("OLD-WO_旧工单", level=3)
    add_numbering(work_order_heading, 3)
    set_paragraph_format(work_order_heading, font="微软雅黑 Light", size=14)
    demand_body = document.add_paragraph("需求口径：旧内容。")
    add_numbering(demand_body, 0)
    set_paragraph_format(demand_body, first_line=480)
    document.add_heading("共享开放方案", level=2)
    document.add_paragraph("STATIC_SHARED_SOLUTION_MARKER")
    document.add_heading("API接口清单", level=2)
    interface_heading = document.add_heading("企业电子票据证件号码校验接口", level=3)
    add_numbering(interface_heading, 4)
    set_paragraph_format(interface_heading, font="微软雅黑 Light", size=14)
    purpose = document.add_paragraph("旧接口说明。")
    add_numbering(purpose, 0)
    set_paragraph_format(purpose, first_line=480)
    supply = document.add_paragraph("计划供数方式：API接口对外服务")
    add_numbering(supply, 5)
    set_paragraph_format(supply)
    frequency = document.add_paragraph("计划更新频率：无")
    add_numbering(frequency, 5)
    set_paragraph_format(frequency)
    input_label = document.add_paragraph("接口输入参数：")
    set_paragraph_format(input_label, first_line=480)
    input_table = add_parameter_table(document, ["序号", "参数名", "名称"], [["1", "old", "旧参数"]])
    set_table_format(input_table, "D9D9D9")
    output_label = document.add_paragraph("接口输出参数：")
    set_paragraph_format(output_label, first_line=480)
    output_table = add_parameter_table(document, ["序号", "参数项", "名称"], [["1", "oldResult", "旧结果"]])
    set_table_format(output_table, "E7E6E6")
    document.save(path)
    return path


def make_api_code_template(path: Path):
    image_path = path.with_name("template-screenshot.png")
    image_path.write_bytes(PNG_1X1)
    document = Document()
    document.add_paragraph("API接口开发代码")
    document.add_paragraph("开发代码")
    document.add_heading("API接口开发列表", level=1)
    table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["序号", "接口代码", "接口名称", "责任委办"]):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "old_api", "旧接口", "应用开发部"]):
        table.rows[1].cells[index].text = value
    set_table_format(table, "D9EAF7")
    document.add_heading("接口模型明细", level=1)
    interface_heading = document.add_heading("旧接口", level=2)
    set_paragraph_format(interface_heading, font="微软雅黑 Light", size=14)
    add_picture_paragraph(document, image_path)
    add_picture_paragraph(document, image_path)
    document.save(path)
    return path


def make_api_launch_record_template(path: Path):
    image_path = path.with_name("template-launch-screenshot.png")
    image_path.write_bytes(PNG_1X1)
    document = Document()
    document.add_paragraph("API作业上线记录")
    document.add_paragraph("上线记录")
    document.add_heading("API作业上线清单", level=1)
    table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["序号", "接口代码", "接口名称", "责任委办"]):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "old_api", "旧接口", "应用开发部"]):
        table.rows[1].cells[index].text = value
    set_table_format(table, "D9EAF7")
    document.add_heading("API作业上线记录", level=1)
    interface_heading = document.add_heading("旧接口", level=2)
    set_paragraph_format(interface_heading, font="微软雅黑 Light", size=14)
    add_picture_paragraph(document, image_path)
    document.save(path)
    return path


def make_api_test_report_template(path: Path):
    document = Document()
    document.add_paragraph("API接口测试报告")
    document.add_paragraph("测试报告")
    document.add_heading("API测试清单", level=1)
    list_table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["序号", "接口代码", "接口名称", "责任委办"]):
        list_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "old_api", "旧接口", "应用开发部"]):
        list_table.rows[1].cells[index].text = value
    set_table_format(list_table, "D9EAF7")
    document.add_heading("API测试内容", level=1)
    interface_heading = document.add_heading("旧接口 old_api", level=2)
    set_paragraph_format(interface_heading, font="微软雅黑 Light", size=14)
    input_heading = document.add_heading("输入参数", level=3)
    set_paragraph_format(input_heading, font="微软雅黑", size=11)
    input_table = add_parameter_table(
        document,
        ["序号", "参数项", "名称", "测试数据1"],
        [["1", "old", "旧参数", "旧测试数据"]],
    )
    set_table_format(input_table, "D9D9D9")
    output_heading = document.add_heading("输出参数", level=3)
    set_paragraph_format(output_heading, font="微软雅黑", size=11)
    output_table = add_parameter_table(
        document,
        ["序号", "参数项", "名称", "结果数据"],
        [["1", "oldResult", "旧结果", "旧结果数据"]],
    )
    set_table_format(output_table, "E7E6E6")
    result_heading = document.add_heading("测试结果", level=3)
    set_paragraph_format(result_heading, font="微软雅黑", size=11)
    result_body = document.add_paragraph("测试结果与预期结果一致，测试通过。")
    set_paragraph_format(result_body, first_line=480)
    document.add_heading("测试结论", level=1)
    document.add_paragraph("STATIC_TEST_CONCLUSION_MARKER")
    document.save(path)
    return path


def make_api_data_model_template(path: Path):
    document = Document()
    document.add_paragraph("API数据模型设计")
    document.add_heading("目的", level=1)
    body = document.add_paragraph("通过对API数据模型进行设计。")
    set_paragraph_format(body, first_line=480, line_spacing=360)
    document.add_heading("接口配置表设计", level=1)
    document.add_heading("接口定义表", level=2)
    definition = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["字段名称", "字段类型", "默认", "不可为空", "唯一", "字段注释"]):
        definition.rows[0].cells[index].text = value
    for index, value in enumerate(["api_code", "varchar(32)", "NULL", "NO", "YES", "接口编码"]):
        definition.rows[1].cells[index].text = value
    set_table_format(definition, "D9EAF7")
    document.add_heading("接口出入参配置表", level=2)
    work_order_heading = document.add_heading("旧工单标题", level=3)
    set_paragraph_format(work_order_heading, font="微软雅黑 Light", size=14)
    interface_heading = document.add_heading("旧接口", level=4)
    set_paragraph_format(interface_heading, font="微软雅黑 Light", size=12)
    input_heading = document.add_heading("接口入参配置表", level=5)
    set_paragraph_format(input_heading, font="微软雅黑", size=11)
    input_table = add_parameter_table(document, ["序号", "参数名", "名称"], [["1", "old", "旧参数"]])
    set_table_format(input_table, "D9D9D9")
    output_heading = document.add_heading("接口出参配置表", level=5)
    set_paragraph_format(output_heading, font="微软雅黑", size=11)
    output_table = add_parameter_table(document, ["序号", "参数项", "名称"], [["1", "oldResult", "旧结果"]])
    set_table_format(output_table, "E7E6E6")
    document.add_heading("接口组件配置表", level=2)
    document.add_paragraph("STATIC_COMPONENT_MARKER")
    component = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["字段名称", "字段类型", "默认", "不可为空", "唯一", "字段注释"]):
        component.rows[0].cells[index].text = value
    for index, value in enumerate(["api_id", "varchar(32)", "NULL", "NO", "YES", "接口id"]):
        component.rows[1].cells[index].text = value
    set_table_format(component, "D9EAF7")
    document.save(path)
    return path


def paragraph_with_text(document, text):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text)


def xml_or_none(element):
    return None if element is None else element.xml


def make_ole_anchor_package(path: Path):
    worksheet = """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
 <oleObjects><oleObject shapeId="1025" r:id="rId2"/></oleObjects>
 <legacyDrawing r:id="rId1"/>
</worksheet>"""
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
 <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/vmlDrawing" Target="../drawings/vmlDrawing1.vml"/>
 <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="../embeddings/oleObject1.bin"/>
</Relationships>"""
    vml = """<?xml version="1.0" encoding="UTF-8"?>
<xml xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:x="urn:schemas-microsoft-com:office:excel">
 <v:shape id="_x0000_s1025" o:ole="t">
  <x:ClientData><x:Anchor>7, 0, 1, 0, 7, 0, 3, 0</x:Anchor></x:ClientData>
 </v:shape>
</xml>"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("xl/worksheets/sheet1.xml", worksheet)
        archive.writestr("xl/worksheets/_rels/sheet1.xml.rels", relationships)
        archive.writestr("xl/drawings/vmlDrawing1.vml", vml)
        archive.writestr("xl/embeddings/oleObject1.bin", b"embedded-ole-payload")
    return path


class ApiRequirementTest(unittest.TestCase):
    def test_api_requirement_registration_uses_public_output_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "01-API接口开发_需求文档")

        self.assertEqual(Path(output).name, "01-需求文档.docx")

    def test_api_requirement_nonexistent_output_directory_uses_public_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            result_dir = Path(temp_dir) / "结果"
            output = module.resolve_output_path(result_dir, "01-API接口开发_需求文档")

        self.assertEqual(Path(output), result_dir / "01-需求文档.docx")

    def test_api_data_model_registration_uses_public_output_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "02-API接口开发_数据模型设计")

        self.assertEqual(Path(output).name, "02- 数据模型设计（API）.docx")

    def test_api_code_registration_uses_public_output_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "03-API接口开发_接口开发代码")

        self.assertEqual(Path(output).name, "03-接口开发代码.docx")

    def test_api_test_report_registration_uses_public_output_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "04-API接口开发_接口测试报告")

        self.assertEqual(Path(output).name, "04-接口测试报告（含《API接口列表》）.docx")

    def test_api_launch_record_registration_uses_public_output_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "05-API接口开发_作业上线记录")

        self.assertEqual(Path(output).name, "05-作业上线记录（API）.doc")

    def test_read_api_work_orders_groups_merged_rows_and_counts_programs_once(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.shared.ledger import read_api_work_orders

        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = make_api_ledger(Path(temp_dir) / "ledger.xlsx")
            orders = read_api_work_orders(ledger, "N07-API接口开发")

        self.assertEqual(len(orders), 1)
        self.assertEqual(orders[0].program_count, 3)
        self.assertEqual(orders[0].source_rows, (2, 3, 4))
        self.assertEqual(
            [item.chinese_name for item in orders[0].interfaces],
            ["境外人员获取令牌接口", "境外人员身份认证申请接口", "境外人员身份认证请求接口"],
        )

    def test_parse_self_report_preserves_input_groups_and_output_tables(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.api_requirement import parse_api_report
        from materials.shared.ledger import ApiInterface

        with tempfile.TemporaryDirectory() as temp_dir:
            report = make_self_report(Path(temp_dir) / "self-test.docx")
            interfaces = [ApiInterface("境外人员身份认证申请接口", "api_apply", 2)]
            parsed = parse_api_report(report, interfaces)

        item = parsed[0]
        self.assertEqual([group.label for group in item.input_groups], ["请求头", "请求体"])
        self.assertEqual(item.input_groups[0].headers, ["参数", "参数说明", "是否必选"])
        self.assertEqual(item.input_groups[1].rows[0], ["custNum", "业务站点号", "是", "String"])
        self.assertEqual(item.input_groups[1].column_widths, [1400, 3600, 1200, 1800])
        self.assertEqual(item.output_groups[0].rows[0][0], "bizSerialNum")

    def test_read_ole_anchors_resolves_vml_position_and_embedding_payload(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.shared.embedded_docx import read_ole_anchors

        with tempfile.TemporaryDirectory() as temp_dir:
            package = make_ole_anchor_package(Path(temp_dir) / "ledger.xlsx")
            anchors = read_ole_anchors(package)

        self.assertEqual(len(anchors), 1)
        self.assertEqual((anchors[0].row, anchors[0].column), (2, 8))
        self.assertEqual(anchors[0].payload, b"embedded-ole-payload")

    def test_parameter_table_scales_source_widths_and_keeps_data_text_left_aligned(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.shared.word_sections import table_element

        table = table_element(
            ["参数", "参数说明", "是否必选", "类型", "备注"],
            [["phoneNumber", "手机号码", "否", "String", "线上业务可填写手机号码"]],
            [1267, 3117, 518, 913, 2905],
        )
        widths = [int(column.get(qn("w:w"))) for column in table.tblGrid.gridCol_lst]
        data_row = table.findall(qn("w:tr"))[1]
        description_cell = data_row.findall(qn("w:tc"))[1]

        self.assertEqual(sum(widths), 9000)
        self.assertGreater(widths[1], widths[0])
        self.assertGreater(widths[4], widths[0])
        self.assertIsNone(description_cell.find(".//" + qn("w:jc")))

    def test_build_api_requirement_document_replaces_dynamic_sections_and_preserves_static_content(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_requirement

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report(temp / "self-test.docx")
            template = make_api_template(temp / "template.docx")
            output = temp / "01-需求文档.docx"
            with mock.patch.object(
                api_requirement,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
                create=True,
            ):
                with mock.patch.object(api_requirement, "update_toc_via_com", create=True):
                    api_requirement.build_api_requirement_document(
                        excel_path=ledger,
                        service_dir="N07-API接口开发",
                        template_path=template,
                        output_path=output,
                    )

            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("共提供3个API接口服务", text)
        self.assertIn("WO-1_市公安局-出入境证件身份认证-共享接口", text)
        self.assertIn("需求口径：升级离境退税掌上办平台的出入境记录校验能力。", text)
        self.assertIn("STATIC_SHARED_SOLUTION_MARKER", text)
        self.assertIn("境外人员获取令牌接口", text)
        self.assertIn("境外人员身份认证申请接口", text)
        self.assertIn("境外人员身份认证请求接口", text)
        self.assertIn("计划供数方式：API接口对外服务", text)
        self.assertNotIn("企业电子票据证件号码校验接口", text)
        self.assertNotIn("旧接口说明", text)

    def test_build_api_data_model_document_replaces_42_section_and_preserves_static_content(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_data_model

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report(temp / "self-test.docx")
            template = make_api_data_model_template(temp / "template.docx")
            output = temp / "02- 数据模型设计（API）.docx"
            with mock.patch.object(
                api_data_model,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
                create=True,
            ), mock.patch.object(api_data_model, "update_toc_via_com", create=True):
                api_data_model.build_api_data_model_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template,
                    output_path=output,
                )

            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("市公安局-出入境证件身份认证-共享接口", text)
        self.assertIn("境外人员获取令牌接口", text)
        self.assertIn("境外人员身份认证申请接口", text)
        self.assertIn("境外人员身份认证请求接口", text)
        self.assertIn("请求头：", text)
        self.assertIn("请求体：", text)
        self.assertIn("STATIC_COMPONENT_MARKER", text)
        self.assertNotIn("旧工单标题", text)
        self.assertNotIn("旧接口", text)

        self.assertTrue(
            any(
                len(table.columns) == 4
                and [cell.text for cell in table.rows[0].cells]
                == ["参数", "参数说明", "是否必选", "类型"]
                for table in document.tables
            )
        )

    def test_parse_api_code_report_images_keeps_only_requested_screenshot_sections(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.api_code_doc import parse_api_code_report_images
        from materials.shared.ledger import ApiInterface

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            report = make_self_report_with_code_images(temp / "self-test.docx")
            interfaces = [ApiInterface("境外人员获取令牌接口", "token_api", 2)]
            images = parse_api_code_report_images(report, interfaces, temp / "work")

        self.assertEqual(len(images["境外人员获取令牌接口"]["共享接口命名规范性"]), 1)
        self.assertEqual(len(images["境外人员获取令牌接口"]["共享任务开发代码检查"]), 1)

    def test_build_api_code_document_replaces_list_and_model_detail_images(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_code_doc

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report_with_code_images(temp / "self-test.docx")
            template = make_api_code_template(temp / "template.docx")
            output = temp / "03-接口开发代码.docx"
            with mock.patch.object(
                api_code_doc,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
                create=True,
            ), mock.patch.object(api_code_doc, "update_toc_via_com", create=True):
                api_code_doc.build_api_code_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template,
                    output_path=output,
                )

            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("境外人员获取令牌接口", text)
        self.assertIn("境外人员身份认证申请接口", text)
        self.assertIn("境外人员身份认证请求接口", text)
        self.assertNotIn("旧接口", text)
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.tables[0].rows), 4)
        self.assertEqual(document.tables[0].rows[1].cells[1].text, "token_api")
        self.assertEqual(document.tables[0].rows[1].cells[3].text, "上海市大数据中心")
        self.assertEqual(len(document.inline_shapes), 6)

    def test_build_api_launch_record_document_replaces_list_and_launch_images(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_launch_record

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report_with_code_images(temp / "self-test.docx")
            template = make_api_launch_record_template(temp / "template.docx")
            output = temp / "05-作业上线记录（API）.docx"
            with mock.patch.object(
                api_launch_record,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
                create=True,
            ), mock.patch.object(api_launch_record, "update_toc_via_com", create=True):
                api_launch_record.build_api_launch_record_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template,
                    output_path=output,
                )

            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("境外人员获取令牌接口", text)
        self.assertIn("境外人员身份认证申请接口", text)
        self.assertIn("境外人员身份认证请求接口", text)
        self.assertNotIn("旧接口", text)
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.tables[0].rows), 4)
        self.assertEqual(document.tables[0].rows[1].cells[1].text, "token_api")
        self.assertEqual(document.tables[0].rows[1].cells[3].text, "上海市大数据中心")
        self.assertEqual(len(document.inline_shapes), 3)

    def test_api_launch_record_doc_output_creates_missing_work_directory(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_launch_record

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            document = Document()
            document.add_paragraph("API作业上线记录")
            output = temp / "结果" / "05-作业上线记录（API）.doc"
            work_dir = temp / "missing" / "output"

            def fake_convert(docx_path, output_path):
                self.assertTrue(Path(docx_path).exists())
                self.assertTrue(Path(docx_path).parent.exists())
                Path(output_path).parent.mkdir(parents=True, exist_ok=True)
                Path(output_path).write_bytes(b"legacy-doc")
                return Path(output_path)

            with mock.patch.object(
                api_launch_record,
                "_convert_docx_to_legacy_doc",
                side_effect=fake_convert,
            ):
                result = api_launch_record._save_document(document, output, work_dir)

        self.assertEqual(result, str(output))

    def test_build_api_test_report_document_replaces_list_params_and_result_images(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_test_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report_with_code_images(temp / "self-test.docx")
            template = make_api_test_report_template(temp / "template.docx")
            output = temp / "04-接口测试报告（含《API接口列表》）.docx"
            with mock.patch.object(
                api_test_report,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
                create=True,
            ), mock.patch.object(api_test_report, "update_toc_via_com", create=True):
                api_test_report.build_api_test_report_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template,
                    output_path=output,
                )

            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("境外人员获取令牌接口 token_api", text)
        self.assertIn("境外人员身份认证申请接口 apply_api", text)
        self.assertIn("境外人员身份认证请求接口 auth_api", text)
        self.assertIn("STATIC_TEST_CONCLUSION_MARKER", text)
        self.assertNotIn("旧接口 old_api", text)
        self.assertEqual(len(document.tables), 7)
        self.assertEqual(len(document.tables[0].rows), 4)
        self.assertEqual(document.tables[0].rows[1].cells[1].text, "token_api")
        self.assertEqual(document.tables[0].rows[1].cells[3].text, "上海市大数据中心")
        self.assertEqual(
            [cell.text for cell in document.tables[1].rows[0].cells],
            ["序号", "参数项", "名称", "测试数据1"],
        )
        self.assertEqual(
            [cell.text for cell in document.tables[2].rows[0].cells],
            ["序号", "参数项", "名称", "结果数据"],
        )
        self.assertEqual(len(document.inline_shapes), 3)

    def test_api_test_report_doc_conversion_creates_missing_work_directory(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_test_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template = temp / "template.doc"
            template.write_bytes(b"legacy-doc-placeholder")
            work_dir = temp / "missing" / "template"

            def fake_run(*args, **kwargs):
                converted = work_dir / "template.docx"
                self.assertTrue(converted.parent.exists())
                converted.write_bytes(b"converted-docx-placeholder")
                return SimpleNamespace(stdout="CONVERTED", stderr="", returncode=0)

            with mock.patch.object(api_test_report.subprocess, "run", side_effect=fake_run):
                converted = api_test_report._convert_legacy_doc_template(template, work_dir)

        self.assertEqual(converted, work_dir / "template.docx")

    def test_generated_dynamic_blocks_reuse_template_direct_formatting(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_requirement

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report(temp / "self-test.docx")
            template_path = make_api_template(temp / "template.docx")
            output_path = temp / "01-需求文档.docx"
            with mock.patch.object(
                api_requirement,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
            ), mock.patch.object(api_requirement, "update_toc_via_com"):
                api_requirement.build_api_requirement_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template_path,
                    output_path=output_path,
                )

            template = Document(template_path)
            output = Document(output_path)
            output_text = "\n".join(paragraph.text for paragraph in output.paragraphs)
            self.assertIn("BUSINESS_STATIC_AFTER", output_text)
            self.assertNotIn("旧业务场景内容。", output_text)
            paragraph_pairs = [
                ("旧业务场景内容。", "在本服务周期内，围绕市公安局-出入境证件身份认证-共享接口开展API接口开发，共提供3个API接口服务。"),
                ("服务周期内，共有3张需求单。", "服务周期内，共有1张需求单，1张工单涉及3个API接口服务。具体需求单、工单和产出如下表："),
                ("OLD-WO_旧工单", "WO-1_市公安局-出入境证件身份认证-共享接口"),
                ("需求口径：旧内容。", "需求口径：升级离境退税掌上办平台的出入境记录校验能力。"),
                ("企业电子票据证件号码校验接口", "境外人员获取令牌接口"),
                ("计划供数方式：API接口对外服务", "计划供数方式：API接口对外服务"),
                ("计划更新频率：无", "计划更新频率：无"),
                ("接口输入参数：", "接口输入参数："),
                ("接口输出参数：", "接口输出参数："),
            ]
            for template_text, output_text in paragraph_pairs:
                expected = paragraph_with_text(template, template_text)
                actual = paragraph_with_text(output, output_text)
                self.assertEqual(xml_or_none(actual._p.pPr), xml_or_none(expected._p.pPr))
                self.assertEqual(
                    xml_or_none(actual.runs[0]._r.rPr),
                    xml_or_none(expected.runs[0]._r.rPr),
                )

            for table_index in (1, 2, 3):
                expected_table = template.tables[table_index]
                actual_table = output.tables[table_index]
                self.assertEqual(actual_table._tbl.tblPr.xml, expected_table._tbl.tblPr.xml)
                self.assertIsNotNone(
                    actual_table.rows[0]._tr.find(".//" + qn("w:tblHeader"))
                )
                for row_index in (0, 1):
                    expected_cell = expected_table.rows[row_index].cells[0]
                    actual_cell = actual_table.rows[row_index].cells[0]
                    self.assertEqual(actual_cell._tc.tcPr.xml, expected_cell._tc.tcPr.xml)
                    self.assertEqual(
                        xml_or_none(actual_cell.paragraphs[0]._p.pPr),
                        xml_or_none(expected_cell.paragraphs[0]._p.pPr),
                    )
                    self.assertEqual(
                        xml_or_none(actual_cell.paragraphs[0].runs[0]._r.rPr),
                        xml_or_none(expected_cell.paragraphs[0].runs[0]._r.rPr),
                    )

            four_column_table = next(
                table
                for table in output.tables
                if len(table.columns) == 4
                and "类型" in [cell.text for cell in table.rows[0].cells]
            )
            widths = [
                int(column.get(qn("w:w")))
                for column in four_column_table._tbl.tblGrid.gridCol_lst
            ]
            self.assertGreater(widths[1], widths[0])
            self.assertGreater(widths[1], widths[2])

    def test_business_scene_accepts_a_single_template_paragraph(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_requirement

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_api_ledger(temp / "ledger.xlsx")
            report = make_self_report(temp / "self-test.docx")
            template_path = make_api_template(temp / "template.docx")
            template = Document(template_path)
            for paragraph in list(template.paragraphs):
                if paragraph.text in {
                    "上海市大数据中心通过API接口开展公共数据共享开放。",
                    "BUSINESS_STATIC_AFTER",
                }:
                    paragraph._p.getparent().remove(paragraph._p)
            template.save(template_path)
            output_path = temp / "01-需求文档.docx"
            with mock.patch.object(
                api_requirement,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": report},
            ), mock.patch.object(api_requirement, "update_toc_via_com"):
                api_requirement.build_api_requirement_document(
                    excel_path=ledger,
                    service_dir="N07-API接口开发",
                    template_path=template_path,
                    output_path=output_path,
                )

            output = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in output.paragraphs)

        self.assertIn("共提供3个API接口服务", text)
        self.assertNotIn("旧业务场景内容。", text)

    def test_demand_work_order_headings_reuse_template_numbering_for_multiple_orders(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07 import api_requirement
        from materials.shared.ledger import ApiWorkOrder

        orders = [
            ApiWorkOrder("REQ-1", "WO-1", "第一张工单", "第一项需求。", 1, (2,), []),
            ApiWorkOrder("REQ-2", "WO-2", "第二张工单", "第二项需求。", 1, (3,), []),
        ]
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template_path = make_api_template(temp / "template.docx")
            output_path = temp / "01-需求文档.docx"
            with mock.patch.object(api_requirement, "read_api_work_orders", return_value=orders), mock.patch.object(
                api_requirement,
                "extract_embedded_docx_by_work_order",
                return_value={"WO-1": temp / "one.docx", "WO-2": temp / "two.docx"},
            ), mock.patch.object(api_requirement, "parse_api_report"), mock.patch.object(
                api_requirement, "update_toc_via_com"
            ):
                api_requirement.build_api_requirement_document(
                    excel_path=temp / "ledger.xlsx",
                    service_dir="N07-API接口开发",
                    template_path=template_path,
                    output_path=output_path,
                )

            output = Document(output_path)
            headings = [
                paragraph
                for paragraph in output.paragraphs
                if paragraph.text in {"WO-1_第一张工单", "WO-2_第二张工单"}
            ]

        self.assertEqual([paragraph.text for paragraph in headings], ["WO-1_第一张工单", "WO-2_第二张工单"])
        self.assertEqual(
            [paragraph._p.pPr.numPr.numId.val for paragraph in headings],
            [3, 3],
        )
        self.assertEqual(
            [paragraph._p.pPr.numPr.ilvl.val for paragraph in headings],
            [0, 0],
        )

    def test_purpose_text_is_grounded_and_avoids_banned_ai_phrases(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.api_requirement import build_interface_purpose
        from materials.shared.ledger import ApiInterface, ApiWorkOrder, ParameterGroup

        interface = ApiInterface(
            chinese_name="境外人员身份认证申请接口",
            english_name="apply_api",
            source_row=2,
            input_groups=[ParameterGroup("请求体", ["参数", "参数说明"], [["custNum", "业务站点号"]])],
            output_groups=[ParameterGroup("", ["参数", "参数说明"], [["bizSerialNum", "业务流水号"]])],
        )
        order = ApiWorkOrder(
            demand_no="REQ-1",
            work_order_no="WO-1",
            title="出入境证件身份认证",
            description="升级离境退税掌上办平台的出入境记录校验能力。",
            program_count=1,
            source_rows=(2,),
            interfaces=[interface],
        )

        purpose = build_interface_purpose(order, interface)

        self.assertIn("身份认证", purpose)
        self.assertIn("业务流水号", purpose)
        for banned in ("赋能", "至关重要", "确保", "重要支撑"):
            self.assertNotIn(banned, purpose)

    def test_fill_document_dispatches_registered_api_requirement_material(self):
        module = load_fill_document_module()
        builder = mock.Mock(return_value="out.docx")
        spec = mock.Mock(default_filename="01-需求文档.docx")
        with mock.patch.object(module, "get_material_spec", return_value=spec):
            with mock.patch.object(module, "load_material_builder", return_value=builder, create=True):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N07-API接口开发",
                    material_type="01-API接口开发_需求文档",
                    template_path="template.docx",
                    output_path="out.docx",
                )

        self.assertEqual(result, "out.docx")
        builder.assert_called_once_with(
            excel_path="ledger.xlsx",
            service_dir="N07-API接口开发",
            template_path="template.docx",
            output_path="out.docx",
        )

    def test_fill_document_dispatches_registered_api_data_model_material(self):
        module = load_fill_document_module()
        builder = mock.Mock(return_value="out.docx")
        spec = mock.Mock(default_filename="02- 数据模型设计（API）.docx")
        with mock.patch.object(module, "get_material_spec", return_value=spec):
            with mock.patch.object(module, "load_material_builder", return_value=builder, create=True):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N07-API接口开发",
                    material_type="02-API接口开发_数据模型设计",
                    template_path="template.docx",
                    output_path="out.docx",
                )

        self.assertEqual(result, "out.docx")
        builder.assert_called_once_with(
            excel_path="ledger.xlsx",
            service_dir="N07-API接口开发",
            template_path="template.docx",
            output_path="out.docx",
        )

    def test_fill_document_dispatches_registered_api_code_material(self):
        module = load_fill_document_module()
        builder = mock.Mock(return_value="out.docx")
        spec = mock.Mock(default_filename="03-接口开发代码.docx")
        with mock.patch.object(module, "get_material_spec", return_value=spec):
            with mock.patch.object(module, "load_material_builder", return_value=builder, create=True):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N07-API接口开发",
                    material_type="03-API接口开发_接口开发代码",
                    template_path="template.docx",
                    output_path="out.docx",
                )

        self.assertEqual(result, "out.docx")
        builder.assert_called_once_with(
            excel_path="ledger.xlsx",
            service_dir="N07-API接口开发",
            template_path="template.docx",
            output_path="out.docx",
        )

    def test_fill_document_dispatches_registered_api_test_report_material(self):
        module = load_fill_document_module()
        builder = mock.Mock(return_value="out.docx")
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "结果"
            with mock.patch.object(module, "load_material_builder", return_value=builder, create=True):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N07-API接口开发",
                    material_type="04-API接口开发_接口测试报告",
                    template_path="template.docx",
                    output_path=output_dir,
                )

        self.assertEqual(result, "out.docx")
        builder.assert_called_once_with(
            excel_path="ledger.xlsx",
            service_dir="N07-API接口开发",
            template_path="template.docx",
            output_path=str(output_dir / "04-接口测试报告（含《API接口列表》）.docx"),
        )

    def test_fill_document_dispatches_registered_api_launch_record_material(self):
        module = load_fill_document_module()
        builder = mock.Mock(return_value="out.doc")
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "结果"
            with mock.patch.object(module, "load_material_builder", return_value=builder, create=True):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N07-API接口开发",
                    material_type="05-API接口开发_作业上线记录",
                    template_path="template.doc",
                    output_path=output_dir,
                )

        self.assertEqual(result, "out.doc")
        builder.assert_called_once_with(
            excel_path="ledger.xlsx",
            service_dir="N07-API接口开发",
            template_path="template.doc",
            output_path=str(output_dir / "05-作业上线记录（API）.doc"),
        )


if __name__ == "__main__":
    unittest.main()
