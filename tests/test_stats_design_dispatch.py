import importlib.util
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import openpyxl
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "fill_document.py"


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_split_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    headers = [
        "服务目录",
        "需求单号",
        "工单号",
        "工单内容",
        "统计分析结果表清单",
    ]
    ws.append(headers)
    ws.append(["N08-数据统计分析", "REQ-1", "WO-1", "工单一", "结果表一 RESULT_ONE"])
    ws.append([None, None, None, None, "结果表二 RESULT_TWO"])
    ws.append([None, None, None, None, "结果表三 RESULT_THREE"])
    for column in ["A", "B", "C", "D"]:
        ws.merge_cells(f"{column}2:{column}4")
    wb.save(path)


def make_design_template(path: Path):
    doc = Document()
    doc.add_heading("文档介绍", level=1)
    doc.add_heading("需求来源", level=2)
    doc.add_paragraph("服务周期内，共有0张需求单，0张工单涉及0次数据统计分析服务。具体需求单、工单和产出如下表：")
    table = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["序号", "需求单编号", "工单编号", "工单名称", "程序名称"]):
        table.rows[0].cells[idx].text = header
    doc.add_heading("数据统计分析设计", level=1)
    doc.add_heading("旧程序", level=2)
    doc.add_paragraph("旧内容")
    doc.save(path)


def make_usage_workbook(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "1、数据源表list"
    ws.append(["资源名称", "资源编目（非必填）", "资源信息（表名）", "数据融合加工表"])
    ws.append(["源表一", "DIR-SRC-1", "SRC_ONE", "RESULT_ONE"])
    ws.append(["源表二", "DIR-SRC-2", "SRC_TWO", "RESULT_TWO"])

    ws = wb.create_sheet("2、表融合关系")
    row = 1
    for idx, en in enumerate(["RESULT_ONE", "RESULT_TWO", "RESULT_THREE"], start=1):
        ws.cell(row, 1).value = f"2.{idx} {en}"
        ws.cell(row + 1, 2).value = "文字描述"
        ws.cell(row + 1, 3).value = "1\t数据来源准备：读取源表。\n2\t结果输出：写入结果表。"
        row += 3

    ws = wb.create_sheet("4、数据统计分析结果表详情")
    row = 1
    for idx, en in enumerate(["RESULT_ONE", "RESULT_TWO", "RESULT_THREE"], start=1):
        ws.cell(row, 1).value = f"4.{idx}"
        ws.cell(row, 2).value = en
        row += 1
        ws.append(["", "字段中文名", "字段英文名", "字段类型", "默认", "不可为空", "唯一", "字段注释"])
        row += 1
        ws.append(["", "字段一", "FIELD_ONE", "VARCHAR2(50)", "", "不可为空", "唯一", "字段说明"])
        row += 1
    wb.save(path)


def make_catalog(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "关联资源信息"
    ws.append(["资源编码", "数据目录代码", "资源名称", "业务数据更新周期"])
    ws.append(["SRC_ONE", "CAT-SRC-1", "源表一", "每日"])
    ws.append(["SRC_TWO", "CAT-SRC-2", "源表二", "每日"])
    ws.append(["RESULT_ONE", "CAT-R1", "结果表一", "每日"])
    ws.append(["RESULT_TWO", "CAT-R2", "结果表二", "每日"])
    ws.append(["RESULT_THREE", "CAT-R3", "结果表三", "每日"])
    wb.save(path)


class StatsDesignDispatchTest(unittest.TestCase):
    def test_stats_design_requires_catalog(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "02-数据统计分析_设计文档.docx"
            output = temp / "out.docx"
            make_split_ledger(ledger)
            template.write_bytes(b"placeholder")

            with self.assertRaisesRegex(ValueError, "02-数据统计分析_设计文档需要 --catalog"):
                module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="02-数据统计分析_设计文档",
                    template_path=str(template),
                    output_path=str(output),
                )

    def test_dispatches_to_stats_design_word_generator_with_split_programs(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "02-数据统计分析_设计文档.docx"
            relation = temp / "04-数据统计分析_结果表及使用说明.xlsx"
            catalog = temp / "catalog.xlsx"
            output = temp / "out.docx"
            make_split_ledger(ledger)
            template.write_bytes(b"placeholder")
            relation_wb = openpyxl.Workbook()
            relation_wb.active.title = "2、表融合关系"
            relation_wb.save(relation)
            openpyxl.Workbook().save(catalog)

            calls = []

            def fake_fill(excel_path, data_rows, template_path, output_path, catalog_path, relation_path=None):
                calls.append((excel_path, data_rows, template_path, output_path, catalog_path, relation_path))
                Path(output_path).write_bytes(b"generated-docx")
                return output_path

            with mock.patch.object(module, "fill_stats_design_doc", fake_fill, create=True):
                result = module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="02-数据统计分析_设计文档",
                    template_path=str(template),
                    output_path=str(output),
                    catalog_path=str(catalog),
                )

            self.assertEqual(result, str(output))
            self.assertEqual(len(calls), 1)
            self.assertEqual(calls[0][0], str(ledger))
            self.assertEqual(calls[0][2:], (str(template), str(output), str(catalog), str(relation)))
            self.assertEqual(len(calls[0][1]), 1)
            self.assertEqual(
                calls[0][1][0]["results"],
                [("结果表一", "RESULT_ONE"), ("结果表二", "RESULT_TWO"), ("结果表三", "RESULT_THREE")],
            )

    def test_generates_stats_design_word_for_each_split_program(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "02-数据统计分析_设计文档.docx"
            relation = temp / "04-数据统计分析_结果表及使用说明.xlsx"
            catalog = temp / "catalog.xlsx"
            output = temp / "out.docx"
            make_split_ledger(ledger)
            make_design_template(template)
            make_usage_workbook(relation)
            make_catalog(catalog)

            with mock.patch.object(module, "update_toc_via_com", lambda _path: None):
                result = module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="02-数据统计分析_设计文档",
                    template_path=str(template),
                    output_path=str(output),
                    catalog_path=str(catalog),
                )

            self.assertEqual(result, str(output))
            self.assertTrue(output.exists())
            self.assertFalse(output.with_suffix(".pdf").exists())

            doc = Document(output)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("RESULT_ONE", full_text)
            self.assertIn("RESULT_TWO", full_text)
            self.assertIn("RESULT_THREE", full_text)
            self.assertNotIn("旧内容", full_text)
            self.assertEqual(len([p for p in doc.paragraphs if p.style.name == "Heading 2" and p.text.startswith("结果表")]), 3)
            self.assertGreaterEqual(len(doc.tables), 10)


if __name__ == "__main__":
    unittest.main()
