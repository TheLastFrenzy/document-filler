import importlib.util
import importlib
import tempfile
import unittest
import zipfile
import sys
from pathlib import Path
from xml.etree import ElementTree as ET
from unittest import mock

import openpyxl
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

fitz = importlib.import_module("fitz") if importlib.util.find_spec("fitz") else None


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "fill_document.py"
PDF_SCRIPT = ROOT / "scripts" / "build_stats_test_pdf.py"


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def load_pdf_helper_module():
    spec = importlib.util.spec_from_file_location("build_stats_test_pdf", PDF_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def make_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["服务目录", "需求单号", "工单号", "工单内容", "统计分析结果表清单", "03-数据统计分析_测试文档_工单自测报告附件"])
    ws.append(["N08-数据统计分析", "REQ-1", "WO-1", "工单一", "结果表一 RESULT_ONE", None])
    wb.save(path)


def make_new_column_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["服务目录", "需求单号", "工单号", "工单标题", "结果表清单", "03-数据统计分析_测试文档_工单自测报告附件"])
    ws.append(["N08-数据统计分析", "REQ-1", "WO-1", "新版工单标题", "结果表一 RESULT_ONE", None])
    wb.save(path)


def add_direct_numbering(paragraph, ilvl: int, num_id: int = 1):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def make_template_docx(path: Path):
    doc = Document()
    doc.add_paragraph("TEMPLATE_COVER_MARKER")
    doc.add_paragraph("修订记录")
    doc.add_heading("文档说明", level=1)
    doc.add_paragraph("TEMPLATE_STATIC_CHAPTER_MARKER")
    doc.add_heading("项目背景", level=1)
    doc.add_heading("测试环境与配置", level=1)
    doc.add_heading("测试标准", level=1)
    add_direct_numbering(doc.add_heading("测试内容", level=1), 0)
    doc.add_paragraph("旧的测试内容占位说明。")
    add_direct_numbering(doc.add_heading("专题分析_应急处理_应急工单分类_小时", level=2), 1)
    add_direct_numbering(doc.add_heading("程序中英文名称规范性", level=3), 2)
    doc.add_paragraph("OLD_SECTION_SHOULD_NOT_COPY")
    add_direct_numbering(doc.add_heading("测试结论", level=1), 0)
    doc.add_paragraph("TEMPLATE_CONCLUSION_TEXT")
    table = doc.add_table(rows=3, cols=4)
    table.rows[0].cells[0].text = "测试内容"
    table.rows[0].cells[1].text = "测试项"
    table.rows[0].cells[2].text = "通过项"
    table.rows[0].cells[3].text = "通过率"
    table.rows[1].cells[0].text = "数据测试"
    table.rows[1].cells[1].text = "Xx"
    table.rows[1].cells[2].text = "Xx"
    table.rows[1].cells[3].text = "100%"
    table.rows[2].cells[0].text = "程序测试"
    table.rows[2].cells[1].text = "Xx"
    table.rows[2].cells[2].text = "Xx"
    table.rows[2].cells[3].text = "100%"
    doc.save(path)


def paragraph_numbering(docx_path: Path, target_text: str):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    for paragraph in root.findall(".//w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
        if text != target_text:
            continue
        p_pr = paragraph.find("w:pPr", ns)
        num_pr = p_pr.find("w:numPr", ns) if p_pr is not None else None
        if num_pr is None:
            return None
        ilvl = num_pr.find("w:ilvl", ns)
        num_id = num_pr.find("w:numId", ns)
        return (
            ilvl.attrib.get(f"{{{ns['w']}}}val") if ilvl is not None else "",
            num_id.attrib.get(f"{{{ns['w']}}}val") if num_id is not None else "",
        )
    return None


class StatsTestDocDispatchTest(unittest.TestCase):
    def test_dispatches_to_stats_test_docx_generator(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "03-数据统计分析_测试文档.docx"
            output = temp / "out.docx"
            make_ledger(ledger)
            make_template_docx(template)

            calls = []

            def fake_generator(excel_path, service_dir, template_path, output_path):
                calls.append((excel_path, service_dir, template_path, output_path))
                Path(output_path).write_bytes(b"generated docx")
                return output_path

            with mock.patch.object(module, "fill_stats_test_docx", fake_generator, create=True):
                result = module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="03-数据统计分析_测试文档",
                    template_path=str(template),
                    output_path=str(output),
                )

            self.assertEqual(result, str(output))
            self.assertEqual(calls, [(str(ledger), "N08-数据统计分析", str(template), str(output))])
            self.assertEqual(output.read_bytes(), b"generated docx")

    def test_stats_test_parser_accepts_new_ledger_column_names(self):
        helper = load_pdf_helper_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = Path(temp_dir) / "ledger.xlsx"
            make_new_column_ledger(ledger)

            programs, _groups, _meta = helper.load_ledger_programs(str(ledger), "N08-数据统计分析")

        self.assertEqual(len(programs), 1)
        self.assertEqual(programs[0].work_name, "新版工单标题")
        self.assertEqual(programs[0].result_cn, "结果表一")
        self.assertEqual(programs[0].result_en, "RESULT_ONE")

    def test_generates_docx_with_template_static_sections_and_conclusion_totals(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "03-数据统计分析_测试文档.docx"
            output = temp / "out.docx"

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["服务目录", "需求单号", "工单号", "工单内容", "统计分析结果表清单", "03-数据统计分析_测试文档_工单自测报告附件"])
            ws.append(["N08-数据统计分析", "REQ-1", "WO-1", "工单一", "结果表一 RESULT_ONE", None])
            ws.append([None, None, None, None, "结果表二 RESULT_TWO", None])
            for column in ["A", "B", "C", "D"]:
                ws.merge_cells(f"{column}2:{column}3")
            wb.save(ledger)
            make_template_docx(template)

            result = module.fill_document(
                excel_path=str(ledger),
                service_dir="N08-数据统计分析",
                material_type="03-数据统计分析_测试文档",
                template_path=str(template),
                output_path=str(output),
            )

            self.assertEqual(result, str(output))
            self.assertTrue(output.exists())
            doc = Document(output)
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("TEMPLATE_COVER_MARKER", full_text)
            self.assertIn("TEMPLATE_STATIC_CHAPTER_MARKER", full_text)
            self.assertIn("TEMPLATE_CONCLUSION_TEXT", full_text)
            self.assertIn("本章节依据《台账清单》中服务目录为“N08-数据统计分析”的结果表清单整理，共涉及2个数据统计分析程序。", full_text)
            self.assertIn("结果表一", full_text)
            self.assertIn("结果表二", full_text)
            self.assertNotIn("专题分析_应急处理_应急工单分类_小时", full_text)
            self.assertNotIn("OLD_SECTION_SHOULD_NOT_COPY", full_text)
            table_texts = [
                [cell.text for cell in row.cells]
                for table in doc.tables
                for row in table.rows
            ]
            self.assertIn(["数据测试", "2", "2", "100%"], table_texts)
            self.assertIn(["程序测试", "2", "2", "100%"], table_texts)

    def test_generated_test_content_reuses_template_heading_numbering(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "03-数据统计分析_测试文档.docx"
            output = temp / "out.docx"
            make_ledger(ledger)
            make_template_docx(template)

            module.fill_document(
                excel_path=str(ledger),
                service_dir="N08-数据统计分析",
                material_type="03-数据统计分析_测试文档",
                template_path=str(template),
                output_path=str(output),
            )

            self.assertEqual(paragraph_numbering(output, "结果表一"), ("1", "1"))
            self.assertEqual(paragraph_numbering(output, "程序中英文名称规范性"), ("2", "1"))


if __name__ == "__main__":
    unittest.main()
