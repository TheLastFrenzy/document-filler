import datetime as dt
import importlib.util
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import openpyxl
from PIL import Image, ImageChops
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
FILL_SCRIPT = ROOT / "scripts" / "fill_document.py"
N02_SCRIPT = ROOT / "scripts" / "materials" / "n02" / "test_document.py"


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document_n02", FILL_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def load_n02_module():
    spec = importlib.util.spec_from_file_location("n02_test_document", N02_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_ledger(path: Path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(["任务名称", "执行策略", "成功/失败", "上线时间", "来源委办", "表名"])
    sheet.append(
        ["task_a", "08:00", "3/0", dt.datetime(2025, 8, 22, 16, 7, 5), "上海市财政局", "DWD_TASK_A"]
    )
    sheet.append(
        [
            "task_b",
            "10,00:00",
            "4/0",
            dt.datetime(2025, 9, 22, 16, 56, 5),
            "上海市规划和自然资源局",
            "DWD_TASK_B",
        ]
    )
    sheet.append(
        ["task_c", "8hour", "5/0", dt.datetime(2025, 4, 22, 3, 7, 5), "上海市规划和自然资源局", "DWD_TASK_C"]
    )
    sheet.append(
        ["task_d", "30min", "7/2", dt.datetime(2026, 1, 25, 12, 7, 5), "上海市财政局", "DWD_TASK_D"]
    )
    workbook.save(path)


def make_template_image(path: Path):
    image = Image.new("RGB", (1723, 90), (249, 249, 249))
    image.save(path)


def add_case_table(document: Document, image_path: Path):
    table = document.add_table(rows=8, cols=4)
    table.rows[0].cells[0].text = "测试用例："
    table.rows[1].cells[0].text = "测试目的"
    table.rows[1].cells[1].text = "测试归集程序是否成功运行"
    table.rows[2].cells[0].text = "前置条件"
    table.rows[2].cells[1].text = "归集程序开发完成"
    table.rows[3].cells[0].text = "步    骤"
    table.rows[3].cells[1].text = "具体的操作步骤\n第一步：执行归集程序\n第二步：查看日志结果是否运行成功"
    table.rows[4].cells[0].text = "参数化变量"
    table.rows[4].cells[1].text = "无"
    table.rows[5].cells[0].text = "运行设置"
    table.rows[5].cells[1].text = "无"
    table.rows[6].cells[0].text = "测试结果"
    table.rows[6].cells[1].text = "归集任务执行正常，无报错，测试通过"
    picture_paragraph = table.rows[6].cells[1].add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(4.72))
    table.rows[7].cells[0].text = "测试人"
    table.rows[7].cells[1].text = "占位测试人"
    table.rows[7].cells[2].text = "测试时间"
    table.rows[7].cells[3].text = "20250822"
    return table


def make_template_docx(path: Path, image_path: Path):
    document = Document()
    document.add_paragraph("批数据、实时数据、流式、空间数据的抽取")
    document.add_paragraph("测试报告")
    document.add_paragraph("（V1.0）")
    document.add_heading("测试目的", level=1)
    document.add_paragraph("为了确保数据批抽取归集程序能够正常执行，进行归集程序运行测试。")
    document.add_heading("测试人员", level=1)
    document.add_paragraph("栾希旺")
    document.add_heading("测试范围", level=1)
    document.add_paragraph("针对服务期内新增的24个归集任务进行测试，涉及3个委办。")
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "任务名"
    table.rows[0].cells[2].text = "来源委办"
    table.rows[0].cells[3].text = "表名"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "template_task"
    table.rows[1].cells[2].text = "上海市财政局"
    table.rows[1].cells[3].text = "DWD_TEMPLATE_TASK"
    document.add_heading("测试环境", level=1)
    document.add_paragraph("数据开发平台。")
    document.add_heading("测试方法", level=1)
    document.add_paragraph("执行归集程序，查看是否成功运行。")
    document.add_heading("测试案例", level=1)
    document.add_paragraph("template_task")
    add_case_table(document, image_path)
    document.add_heading("测试结论", level=1)
    document.add_paragraph("通过对上述24个归集任务测试，任务均能正常运行，测试通过。")
    document.add_heading("附件：测试结果截图", level=1)
    document.add_paragraph("见第6章节测试案例表格，测试结果中截图。")
    document.save(path)


class N02TestDocumentTest(unittest.TestCase):
    def test_dispatches_registered_n02_material(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            output_dir = temp / "outputs"
            output_dir.mkdir()
            expected_output = output_dir / "03-测试文档.docx"
            calls = []

            def fake_builder(excel_path, service_dir, template_path, output_path):
                calls.append((excel_path, service_dir, template_path, output_path))
                Path(output_path).write_bytes(b"generated n02")
                return output_path

            with mock.patch.object(module, "load_material_builder", return_value=fake_builder):
                result = module.fill_document(
                    excel_path="ledger.xlsx",
                    service_dir="N02材料",
                    material_type="03-测试文档",
                    template_path="template.docx",
                    output_path=str(output_dir),
                )

            self.assertEqual(result, str(expected_output))
            self.assertEqual(calls, [("ledger.xlsx", "N02材料", "template.docx", str(expected_output))])
            self.assertEqual(expected_output.read_bytes(), b"generated n02")

    def test_normalizes_strategy_text(self):
        module = load_n02_module()

        self.assertEqual(module.normalize_execution_strategy("08:00"), "每天08点")
        self.assertEqual(module.normalize_execution_strategy("10,00:00"), "每月10号的00点")
        self.assertEqual(module.normalize_execution_strategy("8hour"), "每8hour")
        self.assertEqual(module.normalize_execution_strategy("30min"), "每30min")

    def test_reads_ledger_tasks_and_formats_display_fields(self):
        module = load_n02_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = Path(temp_dir) / "N02材料清单.xlsx"
            make_ledger(ledger)

            tasks = module.read_n02_tasks(str(ledger))

        self.assertEqual(len(tasks), 4)
        self.assertEqual(tasks[0].strategy_text, "每天08点")
        self.assertEqual(tasks[1].strategy_text, "每月10号的00点")
        self.assertEqual(tasks[2].strategy_text, "每8hour")
        self.assertEqual(tasks[3].strategy_text, "每30min")
        self.assertEqual(tasks[0].test_time_text, "20250822")
        self.assertEqual(tasks[3].test_time_text, "20260125")

    def test_renders_screenshot_from_template_image(self):
        module = load_n02_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template = temp / "图片1.png"
            output = temp / "rendered.png"
            make_template_image(template)
            task = module.N02Task(
                task_name="caizhengju_db_12_20_shfs_t_czpjff_a",
                strategy="08:00",
                strategy_text="每天08点",
                success_failure="3/0",
                launch_time=dt.datetime(2025, 8, 22, 16, 7, 5),
                launch_time_text="2025-08-22 16:07:05",
                test_time_text="20250822",
                source_department="上海市财政局",
                table_name="DWD_T_CZPJFF_A",
            )

            module.render_task_screenshot(task, template, output)

            original = Image.open(template)
            rendered = Image.open(output)
            self.assertEqual(original.size, rendered.size)
            self.assertIsNotNone(ImageChops.difference(original, rendered).getbbox())

    def test_builds_full_document_from_template_sections(self):
        module = load_n02_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "N02材料清单.xlsx"
            template_dir = temp / "template"
            template_dir.mkdir()
            template_docx = template_dir / "03-测试文档.docx"
            template_image = template_dir / "图片1.png"
            output_dir = temp / "result"
            output_dir.mkdir()
            output_path = output_dir / "03-测试文档.docx"
            make_ledger(ledger)
            make_template_image(template_image)
            make_template_docx(template_docx, template_image)

            result = module.build_n02_test_document(str(ledger), "N02材料", str(template_docx), str(output_path))

            self.assertEqual(result, str(output_path))
            self.assertTrue(output_path.exists())
            document = Document(output_path)
            full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("针对服务期内新增的4个归集任务进行测试，涉及2个委办。", full_text)
            self.assertIn("通过对上述4个归集任务测试，任务均能正常运行，测试通过。", full_text)
            for task_name in ("task_a", "task_b", "task_c", "task_d"):
                self.assertIn(task_name, full_text)
            self.assertEqual(len(document.inline_shapes), 4)
            self.assertIn("栾希旺", document.tables[1].cell(7, 1).text)
            self.assertIn("20250822", document.tables[1].cell(7, 3).text)


if __name__ == "__main__":
    unittest.main()
