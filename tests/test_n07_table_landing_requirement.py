import importlib.util
import base64
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

import openpyxl
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
FILL_SCRIPT = SCRIPTS / "fill_document.py"
PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lvE1"
    "8wAAAABJRU5ErkJggg=="
)


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document_table_landing_test", FILL_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def add_wps_cellimages(path, image_names):
    with zipfile.ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    cell_images = [
        (
            f'<etc:cellImage><xdr:pic><xdr:nvPicPr><xdr:cNvPr id="{index}" '
            f'name="{name}"/></xdr:nvPicPr><xdr:blipFill><a:blip r:embed="rId{index}"/>'
            f"</xdr:blipFill></xdr:pic></etc:cellImage>"
        )
        for index, name in enumerate(image_names, start=1)
    ]
    entries["xl/cellimages.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<etc:cellImages xmlns:etc="http://www.wps.cn/officeDocument/2017/etCustomData" '
        'xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        + "".join(cell_images)
        + "</etc:cellImages>"
    ).encode("utf-8")
    rels = [
        (
            f'<Relationship Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            f'Target="media/{name}.png"/>'
        )
        for index, name in enumerate(image_names, start=1)
    ]
    entries["xl/_rels/cellimages.xml.rels"] = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(rels)
        + "</Relationships>"
    ).encode("utf-8")
    for name in image_names:
        entries[f"xl/media/{name}.png"] = PNG_BYTES

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for name, data in entries.items():
            target.writestr(name, data)


def make_attachment_bytes(rows):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "上线记录库表落地"
    sheet.append(
        [
            "工单号",
            "落地库名",
            "落地表名",
            "表中文名称",
            "落地数据量",
            "落地表简单说明（50字以内）",
            "上线时间",
            "下发任务简单说明（50字以内）",
        ]
    )
    evidence = workbook.create_sheet("佐证截图")
    evidence.append(["任务名", "任务中文名", "落地库名", "落地表名", "表中文名", "源表数据量", "目标表数据量"])
    image_names = []
    for index, row in enumerate(rows, start=1):
        padded = list(row) + [""] * (9 - len(row))
        sheet.append(padded[:8])
        source_image_name = f"SOURCE_IMAGE_{index}"
        target_image_name = f"TARGET_IMAGE_{index}"
        image_names.extend([source_image_name, target_image_name])
        source_table = padded[8] or padded[2]
        evidence.append(
            [
                source_table,
                padded[7],
                padded[1],
                padded[2],
                padded[3],
                f'=_xlfn.DISPIMG("{source_image_name}",1)',
                f'=_xlfn.DISPIMG("{target_image_name}",1)',
            ]
        )
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_file:
        path = Path(temp_file.name)
    try:
        workbook.save(path)
        add_wps_cellimages(path, image_names)
        return path.read_bytes()
    finally:
        path.unlink(missing_ok=True)


def inject_workbook_embeddings(path, embeddings):
    with zipfile.ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    sheet_xml = entries["xl/worksheets/sheet1.xml"].decode("utf-8")
    if "xmlns:r=" not in sheet_xml:
        sheet_xml = sheet_xml.replace(
            "<worksheet ",
            '<worksheet xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" ',
            1,
        )
    ole_objects = (
        '<oleObjects>'
        '<oleObject progId="Excel.Sheet.12" shapeId="1025" r:id="rId2"/>'
        '<oleObject progId="Excel.Sheet.12" shapeId="1026" r:id="rId3"/>'
        "</oleObjects>"
        '<legacyDrawing r:id="rId1"/>'
    )
    sheet_xml = sheet_xml.replace("</worksheet>", ole_objects + "</worksheet>")
    entries["xl/worksheets/sheet1.xml"] = sheet_xml.encode("utf-8")
    entries["xl/worksheets/_rels/sheet1.xml.rels"] = b"""<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
 <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/vmlDrawing" Target="../drawings/vmlDrawing1.vml"/>
 <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="../embeddings/Workbook1.xlsx"/>
 <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="../embeddings/Workbook2.xlsx"/>
</Relationships>"""
    entries["xl/drawings/vmlDrawing1.vml"] = b"""<?xml version="1.0" encoding="UTF-8"?>
<xml xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:x="urn:schemas-microsoft-com:office:excel">
 <v:shape id="_x0000_s1025" o:ole="t">
  <x:ClientData><x:Anchor>9, 0, 1, 0, 9, 0, 1, 0</x:Anchor></x:ClientData>
 </v:shape>
 <v:shape id="_x0000_s1026" o:ole="t">
  <x:ClientData><x:Anchor>9, 0, 2, 0, 9, 0, 3, 0</x:Anchor></x:ClientData>
 </v:shape>
</xml>"""
    entries["xl/embeddings/Workbook1.xlsx"] = embeddings[0]
    entries["xl/embeddings/Workbook2.xlsx"] = embeddings[1]

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for name, data in entries.items():
            target.writestr(name, data)


def make_table_landing_ledger(path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(
        [
            "服务目录",
            "需求单号",
            "工单号",
            "工单标题",
            "程序数",
            "工单描述",
            "结果表清单",
            "数据统计分析执行周期",
            "数据更新要求",
            "自测报告附件",
            "下发前置机中文名",
            "下发前置机数据库类型",
            "上线交付截图1",
            "上线交付截图2",
        ]
    )
    sheet.append(
        [
            "N07-库表落地方式",
            "REQ-1",
            "WO-1",
            "工单一",
            1,
            "工单一描述",
            "PRE_SRC_ONE",
            "每日",
            "每日凌晨",
            None,
            "前置机A",
            "MySQL",
            '=DISPIMG("LAUNCH_ONE_A",1)',
            '=DISPIMG("LAUNCH_ONE_B",1)',
        ]
    )
    sheet.append(
        [
            "N07-库表落地方式",
            "REQ-2",
            "WO-2",
            "工单二",
            2,
            "工单二描述",
            "PRE_SRC_TWO_A",
            "每月",
            "每月15日",
            None,
            "前置机B",
            "Oracle",
            '=DISPIMG("LAUNCH_TWO_A_1",1)',
            '=DISPIMG("LAUNCH_TWO_A_2",1)',
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            "PRE_SRC_TWO_B",
            None,
            None,
            None,
            None,
            None,
            '=DISPIMG("LAUNCH_TWO_B_1",1)',
            '=DISPIMG("LAUNCH_TWO_B_2",1)',
        ]
    )
    for column in "ABCDEFHIJKL":
        sheet.merge_cells(f"{column}3:{column}4")
    workbook.save(path)

    inject_workbook_embeddings(
        path,
        [
            make_attachment_bytes(
                [
                    [
                        "WO-1",
                        "DB_ONE",
                        "TABLE_ONE",
                        "业务一",
                        "10",
                        "落地表一说明",
                        "2025-01-02 09:00:00",
                        "调度一说明",
                        "PRE_SRC_ONE",
                    ]
                ]
            ),
            make_attachment_bytes(
                [
                    [
                        "WO-2",
                        "DB_THREE",
                        "TABLE_THREE",
                        "业务三",
                        "30",
                        "落地表三说明",
                        "2025-03-04 11:00:00",
                        "调度三说明",
                        "PRE_SRC_TWO_B",
                    ],
                    [
                        "WO-2",
                        "DB_TWO",
                        "TABLE_TWO",
                        "业务二",
                        "20",
                        "落地表二说明",
                        "2025-02-03 10:00:00",
                        "调度二说明",
                        "PRE_SRC_TWO_A",
                    ],
                ]
            ),
        ],
    )
    add_wps_cellimages(
        path,
        [
            "LAUNCH_ONE_A",
            "LAUNCH_ONE_B",
            "LAUNCH_TWO_A_1",
            "LAUNCH_TWO_A_2",
            "LAUNCH_TWO_B_1",
            "LAUNCH_TWO_B_2",
        ],
    )
    return path


def make_requirement_template(path):
    document = Document()
    document.add_heading("需求文档", level=1)
    document.add_heading("业务场景", level=2)
    document.add_paragraph("上海市大数据中心通过库表落地方式开展公共数据共享。")
    document.add_paragraph("旧需求来源说明。")
    source_table = document.add_table(rows=3, cols=5)
    for index, value in enumerate(["序号", "对应需求单编号", "对应工单编号", "工单内容", "涉及共享任务数量（个）"]):
        source_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "REQ-OLD", "WO-OLD", "旧工单", "1"]):
        source_table.rows[1].cells[index].text = value
    source_table.rows[2].cells[0].text = "总计"
    source_table.rows[2].cells[4].text = "1"
    document.add_heading("需求说明", level=2)
    document.add_heading("WO-OLD_旧工单", level=2)
    document.add_paragraph("需求口径：旧内容")
    document.add_paragraph("涉及共享任务数量：1")
    detail_table = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["序号", "落地库", "落地表", "对象用户", "库表来源", "业务场景"]):
        detail_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_DB", "OLD_TABLE", "旧前置机", "OLD_SRC", "旧场景"]):
        detail_table.rows[1].cells[index].text = value
    document.add_paragraph("计划供数方式：库表下发")
    document.add_paragraph("计划更新频率：旧频率")
    document.add_paragraph("更新时间：旧时间")
    document.save(path)
    return path


def make_design_template(path):
    document = Document()
    document.add_heading("数据模型设计", level=1)
    document.add_heading("业务逻辑说明", level=1)
    logic_table = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["序号", "目标库", "目标库类型", "目标表", "频率", "规格（记录数）"]):
        logic_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_DB", "OLD_TYPE", "OLD_TABLE", "旧频率", "0"]):
        logic_table.rows[1].cells[index].text = value
    document.add_paragraph("OLD_TABLE：")
    image_path = Path(path).with_suffix(".png")
    image_path.write_bytes(PNG_BYTES)
    document.add_paragraph().add_run().add_picture(str(image_path))
    image_path.unlink(missing_ok=True)

    document.add_heading("库表说明", level=1)
    detail_table = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["序号", "落地库", "落地表", "对象用户", "库表来源", "业务场景"]):
        detail_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_DB", "OLD_TABLE", "旧前置机", "OLD_SRC", "旧场景"]):
        detail_table.rows[1].cells[index].text = value

    document.add_heading("推送方式", level=1)
    document.add_paragraph("固定说明保留")
    document.add_paragraph("推送频率")
    push_table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["序号", "调度名", "调度中文名", "推送频率"]):
        push_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_JOB", "旧调度", "旧频率"]):
        push_table.rows[1].cells[index].text = value
    document.add_paragraph("验证方式")
    document.add_paragraph("固定验证说明保留")
    document.save(path)
    return path


def make_test_report_template(path):
    document = Document()
    document.add_heading("测试报告", level=1)
    document.add_heading("库表落地测试清单", level=1)
    list_table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(["序号", "调度名", "调度中文名"]):
        list_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_JOB", "旧调度说明"]):
        list_table.rows[1].cells[index].text = value

    document.add_heading("库表落地测试内容", level=1)
    document.add_heading("OLD_JOB", level=2)
    document.add_paragraph("OLD_JOB")
    image_path = Path(path).with_suffix(".png")
    image_path.write_bytes(PNG_BYTES)
    document.add_paragraph().add_run().add_picture(str(image_path))
    image_path.unlink(missing_ok=True)
    document.add_paragraph("测试结论：旧结论")

    document.add_heading("测试结论", level=1)
    conclusion_table = document.add_table(rows=2, cols=6)
    for index, value in enumerate(["序号", "测试调度清单", "业务场景", "测试方法", "测试工具", "测试结论"]):
        conclusion_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_JOB", "旧场景", "旧方法", "旧工具", "旧结论"]):
        conclusion_table.rows[1].cells[index].text = value
    document.save(path)
    return path


def make_launch_record_template(path):
    document = Document()
    document.add_heading("作业上线记录", level=1)
    document.add_heading("库表落地上线清单", level=1)
    launch_table = document.add_table(rows=2, cols=5)
    for index, value in enumerate(["序号", "调度名", "上线时间", "更新频率", "场景名称"]):
        launch_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_JOB", "旧上线时间", "旧频率", "旧场景"]):
        launch_table.rows[1].cells[index].text = value
    document.add_heading("库表落地上线记录", level=1)
    document.add_paragraph("固定上线说明保留")
    document.save(path)
    return path


def make_share_record_template(path):
    document = Document()
    document.add_heading("共享记录", level=1)
    document.add_heading("数据库表落地验证清单", level=1)
    list_table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(["序号", "调度名", "调度中文名", "场景名称"]):
        list_table.rows[0].cells[index].text = value
    for index, value in enumerate(["1", "OLD_JOB", "旧调度中文名", "旧场景"]):
        list_table.rows[1].cells[index].text = value

    document.add_heading("数据库表落地验证记录", level=1)
    document.add_heading("OLD_JOB", level=2)
    document.add_paragraph("期望记录条数：旧期望")
    image_path = Path(path).with_suffix(".png")
    image_path.write_bytes(PNG_BYTES)
    document.add_paragraph().add_run().add_picture(str(image_path))
    document.add_paragraph("验证记录条数：旧验证")
    document.add_paragraph().add_run().add_picture(str(image_path))
    image_path.unlink(missing_ok=True)
    document.add_paragraph("验证结果：旧结果")
    document.save(path)
    return path


class N07TableLandingRequirementTest(unittest.TestCase):
    def test_registration_uses_public_requirement_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "01-库表落地方式_需求文档")

        self.assertEqual(Path(output).name, "01-需求文档.docx")

    def test_read_table_landing_work_orders_groups_merged_rows_and_reads_workbook_tasks(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing import read_table_landing_work_orders

        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = make_table_landing_ledger(Path(temp_dir) / "ledger.xlsx")
            orders = read_table_landing_work_orders(ledger, "N07-库表落地方式")

        self.assertEqual([order.work_order_no for order in orders], ["WO-1", "WO-2"])
        self.assertEqual(orders[0].source_rows, (2,))
        self.assertEqual(orders[1].source_rows, (3, 4))
        self.assertEqual(orders[1].program_count, 2)
        self.assertEqual(orders[1].source_tables, ["PRE_SRC_TWO_A", "PRE_SRC_TWO_B"])
        self.assertEqual(orders[1].database_type, "Oracle")
        self.assertEqual(
            [(task.landing_database, task.landing_table, task.business_scene) for task in orders[1].tasks],
            [("DB_TWO", "TABLE_TWO", "业务二"), ("DB_THREE", "TABLE_THREE", "业务三")],
        )
        self.assertEqual([task.landing_data_count for task in orders[1].tasks], ["20", "30"])
        self.assertEqual([task.dispatch_description for task in orders[1].tasks], ["调度二说明", "调度三说明"])
        self.assertEqual([task.launch_time for task in orders[1].tasks], ["2025-02-03 10:00:00", "2025-03-04 11:00:00"])
        self.assertEqual([task.source_volume_image for task in orders[1].tasks], [PNG_BYTES, PNG_BYTES])
        self.assertEqual([task.target_volume_image for task in orders[1].tasks], [PNG_BYTES, PNG_BYTES])

    def test_build_table_landing_requirement_document_replaces_business_and_detail_sections(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing_requirement import build_table_landing_requirement_document

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_table_landing_ledger(temp / "ledger.xlsx")
            template = make_requirement_template(temp / "template.docx")
            output = temp / "out.docx"

            result = build_table_landing_requirement_document(
                excel_path=str(ledger),
                service_dir="N07-库表落地方式",
                template_path=str(template),
                output_path=str(output),
            )

            document = Document(result)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("WO-1_工单一", paragraphs)
        self.assertIn("WO-2_工单二", paragraphs)
        self.assertIn("需求口径：工单二描述", paragraphs)
        self.assertIn("计划更新频率：每月", paragraphs)
        self.assertIn("更新时间：每月15日", paragraphs)
        self.assertNotIn("WO-OLD_旧工单", paragraphs)

        source_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        self.assertEqual(source_rows[1], ["1", "REQ-1", "WO-1", "工单一", "1"])
        self.assertEqual(source_rows[2], ["2", "REQ-2", "WO-2", "工单二", "2"])
        self.assertEqual(source_rows[-1][-1], "3")

        first_detail = [[cell.text.strip() for cell in row.cells] for row in document.tables[1].rows]
        second_detail = [[cell.text.strip() for cell in row.cells] for row in document.tables[2].rows]
        self.assertEqual(first_detail[1], ["1", "DB_ONE", "TABLE_ONE", "前置机A", "PRE_SRC_ONE", "业务一"])
        self.assertEqual(second_detail[1], ["1", "DB_TWO", "TABLE_TWO", "前置机B", "PRE_SRC_TWO_A", "业务二"])
        self.assertEqual(second_detail[2], ["2", "DB_THREE", "TABLE_THREE", "前置机B", "PRE_SRC_TWO_B", "业务三"])

    def test_registration_uses_public_design_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "02-库表落地方式_数据模型设计")

        self.assertEqual(Path(output).name, "02-数据模型设计.doc")

    def test_build_table_landing_design_document_replaces_tables_and_volume_images(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing_design import build_table_landing_design_document

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_table_landing_ledger(temp / "ledger.xlsx")
            template = make_design_template(temp / "template.docx")
            output = temp / "out.docx"

            result = build_table_landing_design_document(
                excel_path=str(ledger),
                service_dir="N07-库表落地方式",
                template_path=str(template),
                output_path=str(output),
            )

            document = Document(result)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("目标表名：TABLE_ONE", paragraphs)
        self.assertIn("目标表名：TABLE_TWO", paragraphs)
        self.assertIn("目标表名：TABLE_THREE", paragraphs)
        self.assertIn("固定说明保留", paragraphs)
        self.assertIn("固定验证说明保留", paragraphs)
        self.assertNotIn("OLD_TABLE：", paragraphs)

        logic_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        detail_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[1].rows]
        push_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[2].rows]

        self.assertEqual(logic_rows[1], ["1", "DB_ONE", "MySQL", "TABLE_ONE", "每日", "10"])
        self.assertEqual(logic_rows[2], ["2", "DB_TWO", "Oracle", "TABLE_TWO", "每月", "20"])
        self.assertEqual(logic_rows[3], ["3", "DB_THREE", "Oracle", "TABLE_THREE", "每月", "30"])
        self.assertEqual(detail_rows[2], ["2", "DB_TWO", "TABLE_TWO", "前置机B", "PRE_SRC_TWO_A", "业务二"])
        self.assertEqual(push_rows[3], ["3", "PRE_SRC_TWO_B", "调度三说明", "每月"])
        self.assertEqual(len(document.inline_shapes), 3)

    def test_registration_uses_public_test_report_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "04-库表落地方式_测试报告")

        self.assertEqual(Path(output).name, "04- 测试报告（库表落地）.doc")

    def test_build_table_landing_test_report_document_replaces_sections_and_launch_images(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing_test_report import build_table_landing_test_report_document

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_table_landing_ledger(temp / "ledger.xlsx")
            template = make_test_report_template(temp / "template.docx")
            output = temp / "out.docx"

            result = build_table_landing_test_report_document(
                excel_path=str(ledger),
                service_dir="N07-库表落地方式",
                template_path=str(template),
                output_path=str(output),
            )

            document = Document(result)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("PRE_SRC_ONE", paragraphs)
        self.assertIn("PRE_SRC_TWO_A", paragraphs)
        self.assertIn("PRE_SRC_TWO_B", paragraphs)
        self.assertIn("测试结论：测试结果与预期结果一致，测试通过。", paragraphs)
        self.assertNotIn("OLD_JOB", paragraphs)

        list_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        conclusion_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[1].rows]
        self.assertEqual(list_rows[1], ["1", "PRE_SRC_ONE", "调度一说明"])
        self.assertEqual(list_rows[2], ["2", "PRE_SRC_TWO_A", "调度二说明"])
        self.assertEqual(list_rows[3], ["3", "PRE_SRC_TWO_B", "调度三说明"])
        self.assertEqual(conclusion_rows[1], ["1", "PRE_SRC_ONE", "业务一", "程序调度", "DACP", "通过"])
        self.assertEqual(conclusion_rows[3], ["3", "PRE_SRC_TWO_B", "业务三", "程序调度", "DACP", "通过"])
        self.assertEqual(len(document.inline_shapes), 6)

    def test_table_landing_test_items_fall_back_when_attachment_task_is_missing(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing import TableLandingTask, TableLandingWorkOrder
        from materials.n07.table_landing_test_report import _build_test_items

        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = make_table_landing_ledger(Path(temp_dir) / "ledger.xlsx")
            order = TableLandingWorkOrder(
                demand_no="REQ-1",
                work_order_no="WO-1",
                title="工单一",
                description="工单描述",
                program_count=2,
                source_rows=(2, 3),
                source_tables=["PRE_SRC_ONE", "PRE_SRC_TWO_A"],
                target_user="前置机",
                database_type="MySQL",
                update_cycle="每日",
                update_requirement="每日凌晨",
                tasks=[
                    TableLandingTask(
                        landing_database="DB_ONE",
                        landing_table="TABLE_ONE",
                        business_scene="业务一",
                        dispatch_description="调度一说明",
                    )
                ],
            )

            items = _build_test_items(str(ledger), [order])

        self.assertEqual(items[0].dispatch_description, "调度一说明")
        self.assertEqual(items[0].business_scene, "业务一")
        self.assertEqual(items[1].dispatch_description, "PRE_SRC_TWO_A")
        self.assertEqual(items[1].business_scene, "PRE_SRC_TWO_A")

    def test_registration_uses_public_launch_record_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "05-库表落地方式_作业上线记录")

        self.assertEqual(Path(output).name, "05-作业上线记录.doc")

    def test_build_table_landing_launch_record_document_replaces_launch_table(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing_launch_record import build_table_landing_launch_record_document

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_table_landing_ledger(temp / "ledger.xlsx")
            template = make_launch_record_template(temp / "template.docx")
            output = temp / "out.docx"

            result = build_table_landing_launch_record_document(
                excel_path=str(ledger),
                service_dir="N07-库表落地方式",
                template_path=str(template),
                output_path=str(output),
            )

            document = Document(result)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("固定上线说明保留", paragraphs)
        self.assertNotIn("OLD_JOB", paragraphs)

        launch_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        self.assertEqual(launch_rows[1], ["1", "PRE_SRC_ONE", "2025-01-02 09:00:00", "每日", "业务一"])
        self.assertEqual(launch_rows[2], ["2", "PRE_SRC_TWO_A", "2025-02-03 10:00:00", "每月", "业务二"])
        self.assertEqual(launch_rows[3], ["3", "PRE_SRC_TWO_B", "2025-03-04 11:00:00", "每月", "业务三"])

    def test_registration_uses_public_share_record_filename(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            output = module.resolve_output_path(temp_dir, "06-库表落地方式_共享记录")

        self.assertEqual(Path(output).name, "06-共享记录.doc")

    def test_build_table_landing_share_record_document_replaces_list_and_record_sections(self):
        if str(SCRIPTS) not in sys.path:
            sys.path.insert(0, str(SCRIPTS))
        from materials.n07.table_landing_share_record import build_table_landing_share_record_document

        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = make_table_landing_ledger(temp / "ledger.xlsx")
            template = make_share_record_template(temp / "template.docx")
            output = temp / "out.docx"

            result = build_table_landing_share_record_document(
                excel_path=str(ledger),
                service_dir="N07-库表落地方式",
                template_path=str(template),
                output_path=str(output),
            )

            document = Document(result)

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("PRE_SRC_ONE", paragraphs)
        self.assertIn("PRE_SRC_TWO_A", paragraphs)
        self.assertIn("PRE_SRC_TWO_B", paragraphs)
        self.assertIn("期望记录条数：20", paragraphs)
        self.assertIn("验证记录条数：30", paragraphs)
        self.assertIn("验证结果：期望记录条数与验证记录条数一致，验证通过；", paragraphs)
        self.assertNotIn("OLD_JOB", paragraphs)

        list_rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        self.assertEqual(list_rows[1], ["1", "PRE_SRC_ONE", "调度一说明", "业务一"])
        self.assertEqual(list_rows[2], ["2", "PRE_SRC_TWO_A", "调度二说明", "业务二"])
        self.assertEqual(list_rows[3], ["3", "PRE_SRC_TWO_B", "调度三说明", "业务三"])
        self.assertEqual(len(document.inline_shapes), 6)


if __name__ == "__main__":
    unittest.main()
