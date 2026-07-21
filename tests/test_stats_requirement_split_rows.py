import importlib.util
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import openpyxl
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "fill_document.py"


def load_fill_document_module():
    spec = importlib.util.spec_from_file_location("fill_document", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_split_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    headers = [
        "服务目录",
        "需求单号",
        "工单号",
        "工单内容",
        "报表统计次数",
        "业务描述",
        "统计分析结果表清单",
        "程序XML文本",
        "执行周期",
        "数据更新要求",
        "数据量对后续运维的特殊要求",
    ]
    ws.append(headers)
    ws.append(
        [
            "N08-数据统计分析",
            "REQ-1",
            "WO-1",
            "工单一",
            "3",
            "业务描述一",
            "结果表一 RESULT_ONE",
            "<mxGraphModel />",
            "按日更新",
            "每日更新",
            "无",
        ]
    )
    ws.append([None, None, None, None, None, None, "结果表二 RESULT_TWO", "<mxGraphModel />", None, None, None])
    ws.append([None, None, None, None, None, None, "结果表三 RESULT_THREE", "<mxGraphModel />", None, None, None])
    for column in ["A", "B", "C", "D", "E", "F", "I", "J", "K"]:
        ws.merge_cells(f"{column}2:{column}4")
    wb.save(path)


def make_new_column_split_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    headers = [
        "服务目录",
        "需求单号",
        "工单号",
        "工单标题",
        "程序数",
        "工单描述",
        "结果表清单",
        "程序XML文本",
        "数据统计分析执行周期",
        "数据更新要求",
        "数据量对后续运维的特殊要求",
    ]
    ws.append(headers)
    ws.append(
        [
            "N08-数据统计分析",
            "REQ-1",
            "WO-1",
            "新版工单标题",
            "3",
            "新版工单描述",
            "结果表一 RESULT_ONE",
            "<mxGraphModel />",
            "按日更新",
            "每日更新",
            "无",
        ]
    )
    ws.append([None, None, None, None, None, None, "结果表二 RESULT_TWO", "<mxGraphModel />", None, None, None])
    ws.append([None, None, None, None, None, None, "结果表三 RESULT_THREE", "<mxGraphModel />", None, None, None])
    for column in ["A", "B", "C", "D", "E", "F", "I", "J", "K"]:
        ws.merge_cells(f"{column}2:{column}4")
    wb.save(path)


def make_relation_workbook(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "2、表融合关系"
    ws["A1"] = "2.1 RESULT_ONE"
    ws["B2"] = "文字描述"
    ws["C2"] = "1\t数据来源准备：读取源表A。\n2\t结果输出：写入结果表一。"
    ws["A4"] = "2.2 RESULT_TWO"
    ws["B5"] = "文字描述"
    ws["C5"] = "1\t数据来源准备：读取源表B。\n2\t结果输出：写入结果表二。"
    wb.save(path)


def make_generation_ledger(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(
        [
            "服务目录",
            "需求单号",
            "工单号",
            "工单标题",
            "程序数",
            "工单描述",
            "结果表清单",
            "执行周期",
            "数据更新要求",
            "数据量对后续运维的特殊要求",
        ]
    )
    ws.append(
        [
            "N08-数据统计分析",
            "REQ-1",
            "WO-1",
            "工单一",
            "2",
            "业务描述一",
            "任务中文名一 RESULT_ONE",
            "按日更新",
            "每日更新",
            "常规运维",
        ]
    )
    ws.append([None, None, None, None, None, None, "任务中文名二 RESULT_TWO", None, None, None])
    for column in ["A", "B", "C", "D", "E", "F", "H", "I", "J"]:
        ws.merge_cells(f"{column}2:{column}3")
    ws.append(
        [
            "N08-数据统计分析",
            "REQ-2",
            "WO-2",
            "工单二",
            "1",
            "业务描述二",
            "任务中文名三 RESULT_THREE",
            "按月更新",
            "每月更新",
            "重点巡检",
        ]
    )
    wb.save(path)


def make_new_requirement_template(path: Path, include_task_table: bool = True):
    doc = Document()
    doc.add_heading("模板说明", level=1)
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "无关表格"
    doc.add_heading("需求来源", level=2)
    doc.add_paragraph("旧需求来源说明。")
    doc.add_heading("需求单、工单产出对应列表", level=3)
    first_table = doc.add_table(rows=1, cols=5)
    for cell, value in zip(
        first_table.rows[0].cells,
        ["序号", "对应需求单编号", "对应工单编号", "工单标题", "数据统计分析次数"],
    ):
        cell.text = value
    first_table.style = "Table Grid"
    first_table.add_row()
    first_table.add_row()
    _format_mapping_table_prototypes(first_table, has_total=True)
    doc.add_heading("工单与任务的对应关系", level=3)
    if include_task_table:
        second_table = doc.add_table(rows=1, cols=4)
        for cell, value in zip(
            second_table.rows[0].cells,
            ["序号", "对应需求单编号", "对应工单编号", "任务中文名"],
        ):
            cell.text = value
        second_table.style = "Table Grid"
        second_table.add_row()
        _format_mapping_table_prototypes(second_table, has_total=False)
    doc.add_heading("需求内容", level=1)
    doc.add_heading("旧工单", level=2)
    doc.add_paragraph("旧内容")
    doc.add_heading("其他要求", level=1)
    doc.add_paragraph("保持不变")
    doc.save(path)


def _format_mapping_table_prototypes(table, has_total):
    header = table.rows[0]
    header.height = Pt(20)
    header.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    tr_properties = header._tr.get_or_add_trPr()
    tr_properties.append(OxmlElement("w:tblHeader"))
    for cell in header.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(192, 0, 0)

    data = table.rows[1]
    data.height = Pt(28)
    data.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for index, cell in enumerate(data.cells):
        cell.text = f"数据原型{index + 1}"
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.italic = True
        run.font.name = "Courier New"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 102, 204)

    if not has_total:
        return
    total = table.rows[2]
    total.height = Pt(18)
    total.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    total._tr.remove(total._tr.tc_lst[1])
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "2")
    total._tr.tc_lst[0].get_or_add_tcPr().append(grid_span)
    for cell in total.cells:
        cell.text = "合计原型"
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 128, 0)


def table_rows(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


class StatsRequirementSplitRowsTest(unittest.TestCase):
    def test_generates_both_new_template_mapping_tables_and_updated_work_order_content(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "template.docx"
            output = temp / "output.docx"
            relation = temp / "relation.xlsx"
            make_generation_ledger(ledger)
            make_new_requirement_template(template)
            make_relation_workbook(relation)
            groups = module.read_stats_requirement_groups(str(ledger), "N08-数据统计分析")

            with mock.patch.object(module, "update_toc_via_com"):
                module.fill_stats_requirement_doc(
                    str(ledger), groups, str(template), str(output), str(relation)
                )

            generated = Document(output)

        self.assertIn(
            "服务周期内，共有2张需求单，2张工单涉及3次数据统计分析",
            [paragraph.text for paragraph in generated.paragraphs],
        )
        self.assertNotIn("具体需求单、工单和产出如下表", "\n".join(p.text for p in generated.paragraphs))
        self.assertEqual(
            table_rows(generated.tables[1]),
            [
                ["序号", "对应需求单编号", "对应工单编号", "工单标题", "数据统计分析次数"],
                ["1", "REQ-1", "WO-1", "工单一", "2"],
                ["2", "REQ-2", "WO-2", "工单二", "1"],
                ["合计", "合计", "", "", "3"],
            ],
        )
        self.assertEqual(
            table_rows(generated.tables[2]),
            [
                ["序号", "对应需求单编号", "对应工单编号", "任务中文名"],
                ["1", "REQ-1", "WO-1", "任务中文名一"],
                ["2", "REQ-1", "WO-1", "任务中文名二"],
                ["3", "REQ-2", "WO-2", "任务中文名三"],
            ],
        )
        source_table = generated.tables[1]
        source_header = source_table.rows[0]
        self.assertIsNotNone(source_header._tr.find(".//" + qn("w:tblHeader")))
        self.assertEqual(source_table.style.name, "Table Grid")
        self.assertEqual(source_header._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")), "400")
        self.assertEqual(source_header.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(source_header.cells[0].paragraphs[0].runs[0].bold)
        self.assertEqual(source_header.cells[0].paragraphs[0].runs[0].font.name, "Arial")
        self.assertEqual(source_header.cells[0].paragraphs[0].runs[0].font.size.pt, 14)
        self.assertEqual(source_header.cells[0].paragraphs[0].runs[0].font.color.rgb, RGBColor(192, 0, 0))

        source_data = source_table.rows[1]
        self.assertEqual(source_data._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")), "560")
        self.assertEqual(source_data.cells[0].vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(source_data.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertTrue(source_data.cells[0].paragraphs[0].runs[0].italic)
        self.assertEqual(source_data.cells[0].paragraphs[0].runs[0].font.name, "Courier New")
        self.assertEqual(source_data.cells[0].paragraphs[0].runs[0].font.size.pt, 11)
        self.assertEqual(source_data.cells[0].paragraphs[0].runs[0].font.color.rgb, RGBColor(0, 102, 204))

        source_total = source_table.rows[-1]
        self.assertEqual(source_total._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")), "360")
        self.assertEqual(source_total.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertTrue(source_total.cells[0].paragraphs[0].runs[0].bold)
        self.assertEqual(source_total.cells[0].paragraphs[0].runs[0].font.name, "Times New Roman")
        self.assertEqual(source_total.cells[0].paragraphs[0].runs[0].font.size.pt, 12)
        self.assertEqual(source_total.cells[0].paragraphs[0].runs[0].font.color.rgb, RGBColor(0, 128, 0))

        task_table = generated.tables[2]
        self.assertIsNotNone(task_table.rows[0]._tr.find(".//" + qn("w:tblHeader")))
        self.assertEqual(task_table.style.name, "Table Grid")
        self.assertEqual(task_table.rows[1]._tr.get_or_add_trPr().find(qn("w:trHeight")).get(qn("w:val")), "560")
        self.assertEqual(task_table.rows[1].cells[0].paragraphs[0].runs[0].font.name, "Courier New")
        paragraphs = [paragraph.text for paragraph in generated.paragraphs]
        self.assertIn("数据加工要求", paragraphs)
        self.assertNotIn("数据加工周期", paragraphs)
        self.assertIn("数据统计分析执行周期：按日更新。", paragraphs)
        self.assertIn("数据更新要求：每日更新。", paragraphs)
        self.assertIn("对运维的工作要求：常规运维。", paragraphs)
        self.assertNotIn("数据量对后续运维的特殊要求：常规运维。", paragraphs)

    def test_missing_task_mapping_table_has_clear_error(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "template.docx"
            output = temp / "output.docx"
            relation = temp / "relation.xlsx"
            make_generation_ledger(ledger)
            make_new_requirement_template(template, include_task_table=False)
            make_relation_workbook(relation)
            groups = module.read_stats_requirement_groups(str(ledger), "N08-数据统计分析")

            with self.assertRaisesRegex(
                ValueError,
                "工单与任务的对应关系.*关联表格",
            ):
                module.fill_stats_requirement_doc(
                    str(ledger), groups, str(template), str(output), str(relation)
                )

    def test_groups_split_result_rows_by_work_order(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = Path(temp_dir) / "ledger.xlsx"
            make_split_ledger(ledger)

            groups = module.read_stats_requirement_groups(str(ledger), "N08-数据统计分析")

        self.assertEqual(len(groups), 1)
        self.assertEqual(groups[0]["工单号"], "WO-1")
        self.assertEqual(groups[0]["报表统计次数"], "3")
        self.assertEqual(
            groups[0]["results"],
            [
                ("结果表一", "RESULT_ONE"),
                ("结果表二", "RESULT_TWO"),
                ("结果表三", "RESULT_THREE"),
            ],
        )

    def test_groups_accept_new_ledger_column_names(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            ledger = Path(temp_dir) / "ledger.xlsx"
            make_new_column_split_ledger(ledger)

            groups = module.read_stats_requirement_groups(str(ledger), "N08-数据统计分析")

        self.assertEqual(len(groups), 1)
        self.assertEqual(groups[0]["工单内容"], "新版工单标题")
        self.assertEqual(groups[0]["报表统计次数"], "3")
        self.assertEqual(groups[0]["业务描述"], "新版工单描述")
        self.assertEqual(groups[0]["业务说明"], "新版工单描述")
        self.assertEqual(groups[0]["数据统计分析执行周期"], "按日更新")
        self.assertEqual(groups[0]["统计分析结果表清单"], "结果表一 RESULT_ONE\n结果表二 RESULT_TWO\n结果表三 RESULT_THREE")

    def test_loads_relation_descriptions_by_result_table_name(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            relation = Path(temp_dir) / "relation.xlsx"
            make_relation_workbook(relation)

            descriptions = module.load_stats_relation_descriptions(str(relation))

        self.assertIn("RESULT_ONE", descriptions)
        self.assertIn("RESULT_TWO", descriptions)
        self.assertIn("读取源表A", descriptions["RESULT_ONE"])

    def test_single_result_description_is_split_into_multiple_logic_steps(self):
        module = load_fill_document_module()

        steps = module.build_stats_requirement_business_logic_steps(
            [("结果表一", "RESULT_ONE")],
            {
                "RESULT_ONE": (
                    "1\t数据来源准备：读取源表A作为基础数据。\n"
                    "2\t数据筛选清洗：过滤无效记录并统一统计口径。\n"
                    "3\t结果输出：写入结果表一供统计分析使用。"
                )
            },
        )

        self.assertEqual([step for step, _ in steps], ["1", "2", "3"])
        self.assertIn("读取源表A", steps[0][1])
        self.assertIn("过滤无效记录", steps[1][1])
        self.assertNotIn("清洗", steps[1][1])
        self.assertIn("数据范围确认", steps[1][1])
        self.assertIn("写入结果表一", steps[2][1])

    def test_dispatch_uses_split_groups_and_sibling_relation_workbook(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            ledger = temp / "ledger.xlsx"
            template = temp / "01-数据统计分析_需求文档.docx"
            relation = temp / "04-数据统计分析_结果表及使用说明.xlsx"
            output = temp / "out.docx"
            make_split_ledger(ledger)
            make_relation_workbook(relation)
            template.write_bytes(b"placeholder")

            calls = []

            def fake_fill(excel_path, data_rows, template_path, output_path, relation_path=None):
                calls.append((excel_path, data_rows, template_path, output_path, relation_path))
                Path(output_path).write_bytes(b"generated")
                return output_path

            with mock.patch.object(module, "fill_stats_requirement_doc", fake_fill):
                result = module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="01-数据统计分析_需求文档",
                    template_path=str(template),
                    output_path=str(output),
                )

            self.assertEqual(result, str(output))
            self.assertEqual(len(calls), 1)
            self.assertEqual(len(calls[0][1]), 1)
            self.assertEqual(len(calls[0][1][0]["results"]), 3)
            self.assertEqual(calls[0][4], str(relation))

    def test_dispatch_prefers_generated_relation_workbook_next_to_output(self):
        module = load_fill_document_module()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template_dir = temp / "template"
            output_dir = temp / "output"
            template_dir.mkdir()
            output_dir.mkdir()
            ledger = temp / "ledger.xlsx"
            template = template_dir / "01-数据统计分析_需求文档.docx"
            stale_relation = template_dir / "04-数据统计分析_结果表及使用说明.xlsx"
            fresh_relation = output_dir / "04-数据统计分析_结果表及使用说明.xlsx"
            output = output_dir / "01-数据统计分析_需求文档.docx"
            make_split_ledger(ledger)
            make_relation_workbook(stale_relation)
            make_relation_workbook(fresh_relation)
            template.write_bytes(b"placeholder")

            calls = []

            def fake_fill(excel_path, data_rows, template_path, output_path, relation_path=None):
                calls.append((excel_path, data_rows, template_path, output_path, relation_path))
                Path(output_path).write_bytes(b"generated")
                return output_path

            with mock.patch.object(module, "fill_stats_requirement_doc", fake_fill):
                result = module.fill_document(
                    excel_path=str(ledger),
                    service_dir="N08-数据统计分析",
                    material_type="01-数据统计分析_需求文档",
                    template_path=str(template),
                    output_path=str(output),
                )

            self.assertEqual(result, str(output))
            self.assertEqual(calls[0][4], str(fresh_relation))


if __name__ == "__main__":
    unittest.main()
